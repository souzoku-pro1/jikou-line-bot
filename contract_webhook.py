"""委任契約書（時効援用）自動生成 Webhook — CONTRACT-GEN-1（第1版）＋fix1

状態機械（fix1[02]・CAS＝$revision 楽観ロック）:
  契約書作成 --（CAS 勝者: 作成→作成中)--> 契約書作成中
    --（生成+検証+upload+添付 PUT〔revision=claim+1〕）--> 契約書作成済
  契約書作成(必須欠落)      → 変更なし（不足フィールド名のみ通知・作用は通知だけ）
  契約書作成中 + 添付なし   → 回収: CAS 再claim → 生成/添付 → 契約書作成済
  契約書作成中 + 添付あり   → 整合確認不能＝自動上書きせず CAS で「要確認」へ
                              ＋管理者通知（fix1[02] reconcile 規則）
  契約書作成済              → already_done skip（再配送・ACK 喪失の冪等化）
  要確認/空/他値            → stale_status skip（fix1[01] 正本の完全一致検証）
  CAS 敗者（409）           → 作用 0 で skip（並行 2 本でも生成/upload は 1 回）

入口ガード（fix1[01]）:
  - webhook 本文 app.id が App21 の実 app ID と完全一致（欠落・非数字・別 App
    は get_record 含め作用 0 で skip）
  - 本文ステータスが「契約書作成」のとき以外は skip（自 update の echo
    〔作成中/作成済〕もここで落ちる＝安価）

凍結文言の構造保証（fix1[03]）:
  - テンプレートは人承認済み現物の SHA-256 と完全一致（実行時+テスト両方）
  - 生成物の報酬条項（第2条全体）がテンプレートと正規化後逐語一致すること
    を実行時検証（不一致は添付せず 500＝差し込み事故の構造検知）

差し込み仕様（CONTRACT-GEN-1・8 キー=一意プレースホルダ数）:
  - {{依頼者氏名}}=顧客名（2 箇所）・{{依頼者住所}}=住所
  - {{対象債権者1}}=問い合わせ業者名・{{対象債権者2}}/{{対象債権者3}}=新設 field
  - 空き枠・契約年月日は全角空白（原本の体裁維持・契約日は締結時確定が既定）

第2版（CONTRACT-GEN-2・PDF 化+CloudSign 自動登録）:
  クラウドサイン登録 --（CAS: 登録→登録中）--> クラウドサイン登録中
    --（PDF 生成+凍結検証 → CloudSign 書類作成+PDF 添付+宛先追加
        → doc id 書き戻し PUT〔revision=claim+1〕）--> クラウドサイン登録済
  前提未充足（docx 未添付/メールアドレス欠落・形式不正/v1 必須欠落）
      → 変更なし（不足フィールド名のみ通知・値は非搭載）
  CloudSign 途中失敗 → 下書き削除（部分状態を残さない）＋掃除成功時のみ
      ステータス巻き戻し（登録中→登録）→ 500 で kintone 再配送=自動再試行。
      掃除失敗時は巻き戻さず 500 → 再配送は下記 reconcile で「要確認」へ
  書類作成 POST の結果不明（fix1[01] ACK 喪失窓: POST 開始後の transport
      例外・5xx・確定拒否 allowlist 外の 4xx〔408/429 等・fix2[04]〕・
      id 欠落）→ 下書きが存在し得るため巻き戻し・再作成とも禁止
      （CloudSignResultUnknown）。「登録中」維持 → reconcile で「要確認」。
      「POST 到達前の失敗（token 取得等）」と「確定拒否 status
      （_CS_DEFINITE_REJECTION）」のみ従来のクリーン巻き戻しを許可
      （_cs_request の unknown_window が実装上の区別）
  クラウドサイン登録中（reconcile） → CloudSign 側に下書きが残り得る（外部
      状態）ため自動再実行はせず常に CAS で「要確認」+管理者通知（fail-closed・
      v1 の添付有無分岐と異なる点は二重下書き防止のため）
  クラウドサイン登録済 → already_done skip（冪等化）

CloudSign 連携の一線（裁定済み方針）:
  - 呼ぶのは 書類作成（POST /documents）・PDF 添付（POST .../files）・
    宛先追加（POST .../participants）・下書き削除（DELETE、掃除時のみ）。
  - 送信 API（PUT /documents/{id}）は呼ばない。送信操作は大野が CloudSign
    画面で行う（対外効果の一線）。テストが source 走査で PUT 不在を pin。

PDF 化の設計（CONTRACT-GEN-2 条件）:
  - テンプレ docx を単一の正とし、fill_template 済み docx の段落を
    contract_pdf が描画（本文をコードに二重管理しない）。
  - 生成 PDF の抽出テキストに対し第1版同水準の凍結検証を実行時に行う
    （verify_frozen_pdf: 第2条逐語一致〔全空白除去後の連続部分列〕+
    差し込みキー残存なし。不一致は登録せず 500）。

スコープ外: CloudSign 送信 API の呼び出し・締結後処理（既存
cloudsign_webhook が cloudsign_document_id で照合して担う）。

特約（JIKOU-CONTRACT-TOKUYAKU）:
  - App 21 の 特約（MULTI_LINE_TEXT）を、雛形の「特約事項」見出し+{{特約}} 本文
    （最後の条文の後・締結文の前）へ差し込む。空（空白のみ）なら 2 段落とも
    削除。非空なら改行ごとに段落を分け、本文段落の書式（pPr・rPr）を複製する
    （apply_tokuyaku・fill_template の run 潰しの影響を受けない専用処理）。
    特約本文の {{ }} は展開しない。TOKUYAKU_MAX_CHARS 超は生成せず「要確認」。
  - docx と CloudSign 用 PDF の両方がテンプレ+レコード値から生成されるため、
    特約を反映する正規の手段は「契約書ステータスを『契約書作成』に戻して
    再生成」（添付は置換される）。ただし cloudsign_document_id が非空
    （CloudSign 下書き作成済み）の再生成要求は「要確認」へ倒す（下書きとの齟齬
    防止・下書きを削除してから再実行）。
  - 運用上の注意: 添付 docx を人が差し替えても、CloudSign 用 PDF はテンプレ+
    レコード値から再生成されるため**差し替え内容は CloudSign に反映されない**。
    文面の変更は雛形（テンプレ）か 特約 欄で行うこと。
  - fix1（CT-01/CT-02）: ガードは _regeneration_guard 1 関数に集約し、契約書作成
    トリガ／作成中 reconcile／クラウドサイン登録 の 3 経路が同じ判定順
    （cs_registered → tokuyaku_too_long → tokuyaku_invalid）で「要確認」へ倒す。
    特約に {{ }} を含む入力は展開せず入力不正として拒否する。

ChannelConfig 化（HOUKI-CONTRACT-GEN）:
  App 固定部分（app・雛形パス/SHA・凍結条項・出力名・CloudSign タイトル・差し込みの
  準備〔prepare〕・通知）を ChannelConfig に切り出し、jikou（本モジュールの定数を
  呼出し時に解決＝既存テストの patch がそのまま効く）と houki（hub/houki_contract）の
  2 設定を持つ。状態機械・CAS・reconcile・_regeneration_guard・CloudSign 呼出し・
  凍結検証の枠組みは 1 経路を共用し、cfg=None は jikou（時効側の挙動不変）。
  houki route: POST /souzoku-houki/contract/{HOUKI_WEBHOOK_TOKEN}（相談カード読取と
  同じトークン・同じ検証方式）。CloudSign 宛先は prepare が返す署名者（全員/代表者）。
"""

import copy
import hashlib
import hmac
import io
import logging
import os
import re
from dataclasses import dataclass, field
from typing import Awaitable, Callable

from docx import Document
from docx.text.paragraph import Paragraph
from fastapi import APIRouter, Request
from fastapi.responses import JSONResponse

import contract_pdf
from hub import kintone as hub_kintone
from hub.docx_builder import fill_template
from hub.redact import emit
from hub.webhook_auth import extract_record_id, verify_token

logger = logging.getLogger("contract")

_APP = hub_kintone.KintoneApp(
    "App 21 (案件)", "KINTONE_APP_ID", "KINTONE_API_TOKEN")

# ── 状態機械（fix1[02]・閉集合 4 値） ────────────────────────────────────────
FIELD_STATUS     = "契約書ステータス"
STATUS_TRIGGER   = "契約書作成"
STATUS_WORKING   = "契約書作成中"
STATUS_DONE      = "契約書作成済"
STATUS_REVIEW    = "要確認"
FIELD_ATTACHMENT = "委任契約書"
TEMPLATE_PATH    = "docx_templates/jikou/委任契約書.docx"
OUTPUT_FILENAME  = "委任契約書_時効援用.docx"
_BLANK           = "　"

# ── 特約（JIKOU-CONTRACT-TOKUYAKU） ──────────────────────────────────────────
FIELD_TOKUYAKU     = "特約"
TOKUYAKU_HEADING   = "特約事項"          # 雛形の見出し段落（逐語）
TOKUYAKU_KEY       = "{{特約}}"          # 雛形の本文段落（単一 run）
TOKUYAKU_MAX_CHARS = 600                 # 超過は生成せず「要確認」（pin）

# ── 第2版（CONTRACT-GEN-2）: CloudSign 登録の状態 3 値+フィールド ──────────
STATUS_CS_TRIGGER = "クラウドサイン登録"
STATUS_CS_WORKING = "クラウドサイン登録中"
STATUS_CS_DONE    = "クラウドサイン登録済"
FIELD_EMAIL       = "メールアドレス"
FIELD_CS_DOC_ID   = "cloudsign_document_id"   # cloudsign_webhook の照合キー
OUTPUT_PDF_NAME   = "委任契約書_時効援用.pdf"
# 簡易 grammar（fail-closed）: ASCII のローカル部@ドメイン.TLD のみ許可。
# 全角・空白・@ 二重などは形式不正として登録しない
_EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")

# fix1[03]: 人承認済み現物テンプレートの SHA-256。
# JIKOU-CONTRACT-TOKUYAKU: 2026-08-22 収載現物（7cc168a1…・scripts/fixtures に
# 保存）へ scripts/add_contract_tokuyaku.py で「特約事項」+{{特約}} の 2 段落を
# 追加した現物（決定的 zip・再現性は test_contract_tokuyaku が pin）
TEMPLATE_SHA256 = (
    "ead90bb8154f64318cc80ee7d6a2dda129192936ca29e32746348ac6145856fd")
# fix1[03]: 報酬条項（第2条全体）の正規化済み逐語（弁護士凍結事項）
FROZEN_CLAUSE = (
    "第2条（弁護士報酬）",
    "1　本件の弁護士報酬（手数料）は、対象債権者1社につき金44,000円"
    "（消費税込み）とする。",
    "2　対象債権者が複数の場合の報酬は、前項の金額に社数を乗じた額とする"
    "（割引は行わない。）。",
    "3　報酬は前払いとし、分割払いはできない。",
)

_REQUIRED_NAME   = "顧客名"
_REQUIRED_ADDR   = "住所"
_CREDITOR_FIELDS = ("問い合わせ業者名", "対象債権者2", "対象債権者3")

router = APIRouter()


class ContractIntegrityError(RuntimeError):
    """fix1[03]: テンプレート/生成物の凍結文言検証に失敗（添付しない）。"""


class CloudSignResultUnknown(RuntimeError):
    """CONTRACT-GEN-2-fix1[01]: 書類作成 POST の結果不明（ACK 喪失窓）。

    POST 開始後に doc id を得られなかった場合（transport 例外・5xx 応答・
    2xx だが id 欠落）は、CloudSign 側に下書きが作成済みの可能性があり id も
    不明。「未作成」と同一視した自動巻き戻し・自動再作成は二重下書きを作る
    ため禁止し、「登録中」維持 → 再配送の reconcile が「要確認」+通知へ倒す
    （人が CloudSign 画面で下書きの有無を確認して整理する運用）。"""


def _fv(record: dict, code: str) -> str:
    return str((record.get(code) or {}).get("value") or "").strip()


# ── ChannelConfig（HOUKI-CONTRACT-GEN） ──────────────────────────────────────
@dataclass
class Prepared:
    """prepare の結果（生成/登録の前提と差し込み済み生成関数）。
    missing: 必須欠落（通知のみ・状態不変。jikou の _missing_fields）
    review: (skip 分類, 通知文)＝「要確認」へ CAS 遷移する理由（houki の前提未充足）
    render: 凍結検証込みの docx bytes を返す関数（docx 添付と PDF の単一経路）
    participants: CloudSign 宛先 [(email, name)]（jikou=1 名・houki=署名者）
    fingerprint: 送信直前の再取得で比較する指紋（jikou は "" =比較しない）
    summary: 通知本文に添える要約（値を含まない）"""
    missing: list = field(default_factory=list)
    review: tuple | None = None
    render: Callable[[], bytes] = lambda: b""
    participants: list = field(default_factory=list)
    fingerprint: str = ""
    summary: str = ""


@dataclass(frozen=True)
class ChannelConfig:
    name: str
    app: hub_kintone.KintoneApp
    template_path: str
    template_sha256: str
    frozen_clause: tuple
    output_filename: str
    output_pdf_name: str
    cs_title: Callable[[str], str]
    prepare: Callable[[dict, str, str], Awaitable[Prepared]]   # (record, record_id, context)
    notify: Callable[[str, str, str], Awaitable[None]]         # (kind, record_id, text)


async def _jikou_prepare(record: dict, record_id: str, context: str) -> Prepared:
    return Prepared(missing=_missing_fields(record),
                    render=lambda: render_contract_docx(record),
                    participants=[(_fv(record, FIELD_EMAIL), _fv(record, _REQUIRED_NAME))])


async def _jikou_notify(kind: str, record_id: str, text: str) -> None:
    """時効側は従来どおり「通知のみ」の経路（missing/要確認/前提未充足）だけ管理者
    LINE へ出す。created/failed は従来どおり通知しない（挙動不変）。"""
    if kind == "review":
        await _notify(text)


def _jikou_cfg() -> ChannelConfig:
    """時効版の設定。定数は呼出し時に解決する（テストの patch.object が効く）。"""
    return ChannelConfig(
        name="jikou", app=_APP, template_path=TEMPLATE_PATH,
        template_sha256=TEMPLATE_SHA256, frozen_clause=FROZEN_CLAUSE,
        output_filename=OUTPUT_FILENAME, output_pdf_name=OUTPUT_PDF_NAME,
        cs_title=lambda rid: f"委任契約書_案件No.{rid}",
        prepare=_jikou_prepare, notify=_jikou_notify)


_HOUKI_NOTIFY_KINDS = {"created": "houki_contract_created",
                       "review": "houki_contract_needs_review",
                       "failed": "houki_contract_failed"}


async def _houki_notify(kind: str, record_id: str, text: str) -> None:
    key = _HOUKI_NOTIFY_KINDS.get(kind, "houki_contract_needs_review")
    try:
        from hub.notify import notify_admin_line
        await notify_admin_line(text, throttle_key=f"{key}:{record_id}",
                                throttle_on_success_only=True)
    except Exception:
        logger.error("[HOUKI_CONTRACT] admin notify failed (fixed text)")


async def _houki_prepare(record: dict, record_id: str, context: str) -> Prepared:
    """申述人集合の取得（被相続人グループID）・費用計算・署名者決定。前提未充足は
    review（要確認へ CAS）。context="register" ではメールアドレスも検査する。"""
    from hub import houki_contract as hc
    p = await hc.plan(record)
    problems = list(p.applicant_problems)
    if context == "register":
        problems += p.email_problems
    review = None
    if problems:
        review = ("houki_preconditions",
                  f"【相続放棄 委任契約書・要確認】案件レコードNo.{record_id} は前提を"
                  "満たしていないため処理を中止しました。\n・" + "\n・".join(problems)
                  + "\n" + p.summary())
    return Prepared(review=review,
                    render=lambda: hc.render(p.records, p.mode, p.fees),
                    participants=list(p.participants), fingerprint=p.fingerprint,
                    summary=p.summary())


def _houki_cfg() -> ChannelConfig:
    from hub import houki_contract as hc
    from hub.houki_case_store import APP_HOUKI_CASE
    return ChannelConfig(
        name="houki", app=APP_HOUKI_CASE, template_path=hc.TEMPLATE_PATH,
        template_sha256=hc.TEMPLATE_SHA256, frozen_clause=hc.FROZEN_CLAUSE,
        output_filename=hc.OUTPUT_FILENAME, output_pdf_name=hc.OUTPUT_PDF_NAME,
        cs_title=lambda rid: f"相続放棄_委任契約書_案件No.{rid}",
        prepare=_houki_prepare, notify=_houki_notify)


def _resolve(cfg: ChannelConfig | None) -> ChannelConfig:
    return cfg if cfg is not None else _jikou_cfg()


def _missing_fields(record: dict) -> list[str]:
    missing = []
    if not _fv(record, _REQUIRED_NAME):
        missing.append(_REQUIRED_NAME)
    if not _fv(record, _REQUIRED_ADDR):
        missing.append(_REQUIRED_ADDR)
    if not any(_fv(record, c) for c in _CREDITOR_FIELDS):
        missing.append("債権者（問い合わせ業者名/対象債権者2/対象債権者3 の"
                       "いずれか1つ以上）")
    return missing


def build_fill_data(record: dict) -> dict:
    """fill_template 用 8 キー+{{特約}}（生値）。render_contract_docx は
    {{特約}} を fill_template に渡さず apply_tokuyaku で専用処理する（run 潰しの
    影響を受けず・改行ごとに段落化）。fill_template に直接渡した場合は単一段落
    への素朴な置換になる（互換のため値は同じ）。"""
    return {
        "{{依頼者氏名}}":  _fv(record, _REQUIRED_NAME),
        "{{依頼者住所}}":  _fv(record, _REQUIRED_ADDR),
        "{{対象債権者1}}": _fv(record, "問い合わせ業者名") or _BLANK,
        "{{対象債権者2}}": _fv(record, "対象債権者2") or _BLANK,
        "{{対象債権者3}}": _fv(record, "対象債権者3") or _BLANK,
        "{{契約年}}": _BLANK, "{{契約月}}": _BLANK, "{{契約日}}": _BLANK,
        TOKUYAKU_KEY: str((record.get(FIELD_TOKUYAKU) or {}).get("value") or ""),
    }


def _placeholder_data(data: dict) -> dict:
    """{{特約}} を除いた 8 キー（apply_tokuyaku が専用処理するため）。"""
    return {k: v for k, v in data.items() if k != TOKUYAKU_KEY}


def tokuyaku_lines(text: str) -> list[str]:
    """特約本文を段落へ分ける（改行ごと・空行は除く）。空/空白のみは []。"""
    return [ln.strip() for ln in str(text or "").splitlines() if ln.strip()]


def tokuyaku_problem(record: dict) -> str | None:
    """特約の入力検査（fail-closed）。判定順: 上限超（tokuyaku_too_long）→
    使用できない記号 {{ }}（tokuyaku_invalid・fix1 CT-02: 仕様変更=展開しない
    ではなく入力不正として拒否）。該当なら skip 分類、なければ None。"""
    text = str((record.get(FIELD_TOKUYAKU) or {}).get("value") or "")
    if len(text.strip()) > TOKUYAKU_MAX_CHARS:
        return "tokuyaku_too_long"
    if "{{" in text or "}}" in text:
        return "tokuyaku_invalid"
    return None


def apply_tokuyaku(docx_bytes: bytes, text: str) -> bytes:
    """雛形の「特約事項」見出し+{{特約}} 本文を特約の実値で置き換える専用処理
    （fill_template の「全 run を先頭 run に潰す」既知問題の影響を受けない）。
    - 空（空白のみ）: 見出し・本文の 2 段落を削除（空の見出しを残さない）
    - 非空: 改行ごとに段落を分け、本文段落の書式（pPr・rPr）を複製して差し込む
      （1 行なら 1 段落）。特約本文の {{ }} は展開しない（文字として出す）——
      fix1 CT-02 で {{ }} を含む特約は tokuyaku_problem が入力不正として拒否する
      ため、本関数には到達しない（防御的に「展開しない」実装は維持）"""
    doc = Document(io.BytesIO(docx_bytes))
    body = next((p for p in doc.paragraphs if p.text == TOKUYAKU_KEY), None)
    if body is None:
        raise ContractIntegrityError("tokuyaku placeholder missing")
    heading = body._p.getprevious()
    if heading is None or Paragraph(heading, body._parent).text != TOKUYAKU_HEADING:
        raise ContractIntegrityError("tokuyaku heading missing")
    lines = tokuyaku_lines(text)
    if not lines:
        parent = body._p.getparent()
        parent.remove(heading)
        parent.remove(body._p)
    else:
        template_p = copy.deepcopy(body._p)      # 書式（pPr・rPr）の原型
        _set_paragraph_text(body, lines[0])
        prev = body._p
        for line in lines[1:]:
            new_p = copy.deepcopy(template_p)
            prev.addnext(new_p)
            _set_paragraph_text(Paragraph(new_p, body._parent), line)
            prev = new_p
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _set_paragraph_text(para, line: str) -> None:
    """先頭 run（rPr 保持）にテキストを置き、残りの run を空にする。"""
    para.runs[0].text = line
    for r in para.runs[1:]:
        r.text = ""


def render_contract_docx(record: dict) -> bytes:
    """雛形の整合検証 → 8 キー差し込み → 特約の専用処理 → 凍結条項検証。
    docx 添付と CloudSign 用 PDF の両方がこの単一の生成経路を使う。"""
    verify_template_integrity()
    data = build_fill_data(record)
    docx_bytes = fill_template(TEMPLATE_PATH, _placeholder_data(data))
    docx_bytes = apply_tokuyaku(docx_bytes, data[TOKUYAKU_KEY])
    verify_frozen_clause(docx_bytes)
    return docx_bytes


def _clause_of(docx_bytes: bytes, frozen: tuple | None = None) -> tuple:
    """docx から凍結ブロック（見出し+各項・正規化=前後空白除去）を抽出。
    frozen 省略時は時効版の第2条（FROZEN_CLAUSE）。"""
    frozen = FROZEN_CLAUSE if frozen is None else frozen
    doc = Document(io.BytesIO(docx_bytes))
    paras = [p.text.strip() for p in doc.paragraphs]
    try:
        i = paras.index(frozen[0])
    except ValueError:
        return ()
    return tuple(paras[i:i + len(frozen)])


def verify_template_integrity() -> None:
    """fix1[03]: 収載テンプレートが人承認済み現物（SHA-256 完全一致）である
    ことの実行時検証。不一致は生成しない。"""
    data = open(TEMPLATE_PATH, "rb").read()
    if hashlib.sha256(data).hexdigest() != TEMPLATE_SHA256:
        raise ContractIntegrityError("template hash mismatch")


def verify_frozen_clause(docx_bytes: bytes, frozen: tuple | None = None) -> None:
    """fix1[03]: 生成物の報酬条項（第2条全体）がテンプレートと逐語一致する
    ことの実行時検証。不一致は添付しない（差し込み事故の構造検知）。
    frozen 省略時は時効版（houki は hub/houki_contract.FROZEN_CLAUSE を渡す）。"""
    frozen = FROZEN_CLAUSE if frozen is None else frozen
    if _clause_of(docx_bytes, frozen) != frozen:
        raise ContractIntegrityError("frozen clause mismatch")


def verify_frozen_pdf(pdf_bytes: bytes, cfg: ChannelConfig | None = None) -> None:
    """CONTRACT-GEN-2: 生成 PDF の抽出テキストへの凍結検証（第1版同水準）。

    PDF は折返しで改行位置が変わるため、全空白（全角空白含む）除去後の
    連結文字列に対して第2条（見出し+3 項）が連続部分列として文字単位で
    逐語一致すること、および差し込みキー（{{ / }}）が残存しないことを検証。
    不一致は CloudSign へ登録しない（500）。"""
    flat = "".join(contract_pdf.pdf_text(pdf_bytes).split())
    frozen = "".join("".join(part.split()) for part in _resolve(cfg).frozen_clause)
    if frozen not in flat:
        raise ContractIntegrityError("frozen clause missing in pdf")
    if "{{" in flat or "}}" in flat:
        raise ContractIntegrityError("unfilled key in pdf")


def verify_pdf_full_text(docx_bytes: bytes, pdf_bytes: bytes) -> None:
    """fix1[02]: PDF 抽出全文と docx 正規化全文の完全一致 pin（第2条限定の
    verify_frozen_pdf を包含する全文保証）。

    正規化は改行・空白（全角空白含む）の除去のみ＝文字列内容そのものは
    docx と PDF で 1 文字も違わないことを要求する。表は docx 側・PDF 側の
    どちらの走査にも乗らず本 pin では欠落を検知できないため、contract_pdf
    が入口で拒否する（PdfUnsupportedStructure・表は非対応）。"""
    doc = Document(io.BytesIO(docx_bytes))
    doc_flat = "".join("".join(p.text.split()) for p in doc.paragraphs)
    pdf_flat = "".join(contract_pdf.pdf_text(pdf_bytes).split())
    if doc_flat != pdf_flat:
        raise ContractIntegrityError("pdf full text mismatch")


async def _notify(text: str) -> None:
    """管理者 LINE 通知（best-effort・固定文言+レコード番号のみ）。"""
    try:
        from hub.notify import notify_admin_line
        await notify_admin_line(text)
    except Exception:
        logger.error("[CONTRACT] admin notify failed (fixed text)")


async def _generate_and_attach(record_id: str, prepared: Prepared,
                               final_revision: str,
                               cfg: ChannelConfig | None = None) -> None:
    """生成 → 凍結文言検証 → upload → 添付+作成済（revision CAS つき PUT）。"""
    cfg = _resolve(cfg)
    docx_bytes = prepared.render()
    logger.info("[CONTRACT] generated record_id=%s bytes=%d",
                emit(record_id, "record_id", "log", "operator"),
                len(docx_bytes))
    mime = ("application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.document")
    file_key = await hub_kintone.upload_file(
        cfg.app, cfg.output_filename, docx_bytes, mime)
    await hub_kintone.update_record(cfg.app, record_id, {
        FIELD_ATTACHMENT: [{"fileKey": file_key}],
        FIELD_STATUS: STATUS_DONE,
    }, revision=final_revision)
    logger.info("[CONTRACT] attached record_id=%s",
                emit(record_id, "record_id", "log", "operator"))
    await cfg.notify("created", record_id,
                     f"【相続放棄 委任契約書】案件レコードNo.{record_id} の委任契約書を"
                     f"作成し添付しました。\n{prepared.summary}")


# ── CloudSign API（CONTRACT-GEN-2）───────────────────────────────────────────
# 呼ぶのは POST（作成/添付/宛先）と DELETE（掃除）のみ。送信 API（PUT）は
# 呼ばない＝送信操作は大野が CloudSign 画面で行う（テストが source pin）。

# fix2[04]: unknown_window で「書類が作成されないまま拒否された」と確定判定
# してよい status の明示 allowlist（閉集合）。HTTP 意味論上、要求を受理せず
# 拒否を応答したことが確定するもののみ:
#   400（要求不正）/ 401（認証拒否・token 再取得の対象）/ 403（権限拒否）/
#   404（endpoint/資源なし）/ 409（競合拒否）/ 422（検証拒否）
# allowlist 外の 4xx は結果不明へ倒す——特に 408（タイムアウト応答＝処理有無
# 不明）・429（レート制限＝処理有無をこの応答からは断定できない）。
_CS_DEFINITE_REJECTION = frozenset({400, 401, 403, 404, 409, 422})


def _cs_request(method: str, path: str, *, unknown_window: bool = False,
                **kwargs):
    """CloudSign API 呼び出し。token 管理は cloudsign_webhook._token（本番
    稼働中の単一の正）を共用し、401 は取り直して 1 回だけ再試行。

    fix1[01]: unknown_window=True（書類作成 POST 専用）は「POST 到達前の
    失敗」と「POST 開始後の結果不明」を実装上区別する。
      - token 取得の失敗＝POST 未実行（下書き未作成が確定）→ 通常伝播
        （呼び出し側のクリーン巻き戻しを許可）
      - request 実行中の transport 例外・5xx 応答・確定拒否 allowlist
        （_CS_DEFINITE_REJECTION）外の 4xx（408/429 等）＝結果不明 →
        CloudSignResultUnknown（巻き戻し・再作成禁止の経路へ）
      - allowlist 内の 4xx 応答＝CloudSign が作成しないまま拒否を応答済み
        （未作成確定）→ 通常の HTTPError。401 も未作成確定なので token を
        取り直して 1 回再試行し、再試行の POST は再び結果不明窓に入る。"""
    import requests

    import cloudsign_webhook as cs
    url = f"{cs.CLOUDSIGN_API_BASE}{path}"

    def _attempt(token: str):
        try:
            return requests.request(
                method, url, headers={"Authorization": f"Bearer {token}"},
                timeout=30, **kwargs)
        except Exception as e:
            if unknown_window:
                raise CloudSignResultUnknown(
                    "cloudsign request outcome unknown") from e
            raise

    resp = _attempt(cs._token.get())
    if resp.status_code == 401:
        cs._token.invalidate()
        resp = _attempt(cs._token.get())
    if unknown_window and resp.status_code >= 400 \
            and resp.status_code not in _CS_DEFINITE_REJECTION:
        raise CloudSignResultUnknown(
            "cloudsign request outcome unknown (non-definite status)")
    resp.raise_for_status()
    return resp


def _cs_create_document(record_id: str, title: str | None = None) -> str:
    """書類作成（下書き）。タイトルは案件 No のみ（氏名等の PII は載せない）。

    fix1[01]: 唯一の unknown_window 呼出し。2xx 応答から id を取り出せない
    場合（本文不正・id 欠落）も「作成された可能性があるが特定できない」＝
    結果不明として扱う（ContractIntegrityError から変更・票 [01] 由来）。"""
    if title is None:
        title = f"委任契約書_案件No.{record_id}"
    resp = _cs_request("POST", "/documents", unknown_window=True,
                       data={"title": title})
    try:
        doc_id = str((resp.json() or {}).get("id") or "")
    except Exception as e:
        raise CloudSignResultUnknown(
            "cloudsign create response unreadable") from e
    if not doc_id:
        raise CloudSignResultUnknown("cloudsign document id missing")
    return doc_id


def _cs_attach_pdf(doc_id: str, pdf_bytes: bytes, name: str | None = None) -> None:
    if name is None:
        name = OUTPUT_PDF_NAME
    _cs_request("POST", f"/documents/{doc_id}/files",
                files={"uploadfile":
                       (name, pdf_bytes, "application/pdf")},
                data={"name": name})


def _cs_add_participant(doc_id: str, email: str, name: str) -> None:
    _cs_request("POST", f"/documents/{doc_id}/participants",
                data={"email": email, "name": name})


def _cs_delete_draft(doc_id: str) -> bool:
    """途中失敗時の下書き掃除（部分状態を残さない）。成功=True。失敗しても
    例外は伝播させない（元の失敗を 500 で報告するのが主）。"""
    try:
        _cs_request("DELETE", f"/documents/{doc_id}")
        return True
    except Exception:
        logger.error("[CONTRACT] cloudsign draft cleanup failed")
        return False


async def _claim(record_id: str, revision: str, to_status: str,
                 app: hub_kintone.KintoneApp | None = None) -> str | None:
    """CAS: $revision 一致時のみステータス遷移。勝者は次 revision（claim+1）を
    返す。

    fix2（CONTRACT-GEN-04）: cas_lost（None・HTTP 200）に落とすのは
    **revision 競合（409）のみ**。通信障害（transport_error）・認証障害
    （401/403）・5xx 等それ以外の KintoneError は再送出し、外側の except が
    HTTP 500 へ落として kintone Webhook の再配送に委ねる（沈黙させない）。
    本関数は「作成→作成中」claim・「作成中」再claim・「要確認」遷移の
    3 箇所すべてで共用される単一の正。"""
    try:
        await hub_kintone.update_record(
            app if app is not None else _APP, record_id,
            {FIELD_STATUS: to_status}, revision=revision)
    except hub_kintone.KintoneError as e:
        if getattr(e, "status", None) == 409:
            return None                  # CAS 敗者（競合）のみ 200/作用 0
        raise                            # 障害系は外側で 500 → 再配送へ
    return str(int(revision) + 1)


async def _to_review(record_id: str, revision: str, text: str, skip: str,
                     cfg: ChannelConfig | None = None):
    """生成せず「要確認」へ CAS 遷移+管理者通知（reconcile と同型）。"""
    cfg = _resolve(cfg)
    next_rev = await _claim(record_id, revision, STATUS_REVIEW, cfg.app)
    if next_rev is None:
        return JSONResponse(status_code=200,
                            content={"ok": True, "skip": "cas_lost"})
    await cfg.notify("review", record_id, text)
    return JSONResponse(status_code=200, content={"ok": True, "skip": skip})


def _regeneration_guard(record_id: str, record: dict,
                        context: str = "regenerate") -> tuple[str, str] | None:
    """再生成/再登録の運用ガード（3 経路共通の単一判定関数。判定順:
    ①CloudSign 下書き作成済み（cs_registered）→ ②特約上限（tokuyaku_too_long）
    → ③特約の使用不可記号（tokuyaku_invalid）。該当なら (skip 分類, 固定文言)。
    context="regenerate"（契約書作成トリガ・作成中 reconcile）／"register"
    （クラウドサイン登録）で cs_registered の文言を分ける。
    houki も同じ判定（欄コードは App 40 も同名・特約の上限は人が書いた欄のみ）。"""
    if _fv(record, FIELD_CS_DOC_ID):
        if context == "register":
            return ("cs_registered",
                    f"【委任契約書・要確認】CloudSign 登録済みのため再登録を中止しま"
                    f"した（レコード番号 {record_id}）。既存の下書きを確認してください。")
        return ("cs_registered",
                f"【委任契約書・要確認】CloudSign 登録済みのため再生成を中止しました"
                f"（レコード番号 {record_id}）。特約を反映する場合は CloudSign の"
                "下書きを削除してから再度お試しください。")
    tk = tokuyaku_problem(record)
    if tk == "tokuyaku_too_long":
        return (tk,
                f"【委任契約書・要確認】特約が {TOKUYAKU_MAX_CHARS} 字を超えている"
                f"ため生成を中止しました（レコード番号 {record_id}）。特約を短くして"
                f"から契約書ステータスを「{STATUS_TRIGGER}」に設定し直してください。")
    if tk == "tokuyaku_invalid":
        return (tk,
                "【委任契約書・要確認】特約欄に使用できない記号（{{ }}）が含まれて"
                f"います（レコード番号 {record_id}）。特約欄を修正して再度お試し"
                "ください。")
    return None


async def _reconcile_working(record_id: str, record: dict,
                             revision: str, cfg: ChannelConfig | None = None):
    """fix1[02] reconcile: 「作成中」で停止した行の回収。
    添付なし=前回 run が upload/添付前に停止 → CAS 再claim して再生成（回収）。
    添付あり=起動内容との整合を機械確認できない → 自動上書きせず CAS で
    「要確認」へ倒し管理者通知。"""
    cfg = _resolve(cfg)
    attachment = (record.get(FIELD_ATTACHMENT) or {}).get("value") or []
    if attachment:
        next_rev = await _claim(record_id, revision, STATUS_REVIEW, cfg.app)
        if next_rev is None:
            return JSONResponse(status_code=200, content={
                "ok": True, "skip": "cas_lost"})
        await cfg.notify(
            "review", record_id,
            f"【委任契約書】案件 No.{record_id} は生成が中断した状態で既に"
            "添付ファイルが存在するため、自動では上書きせず"
            f"「{STATUS_REVIEW}」にしました。添付内容を確認のうえ、再生成する"
            f"場合は添付を削除してステータスを「{STATUS_TRIGGER}」に設定して"
            "ください")
        return JSONResponse(status_code=200, content={
            "ok": True, "skip": "needs_review"})
    guard = _regeneration_guard(record_id, record, "regenerate")
    if guard:
        return await _to_review(record_id, revision, guard[1], guard[0], cfg)
    prepared = await cfg.prepare(record, record_id, "generate")
    if prepared.review:
        return await _to_review(record_id, revision, prepared.review[1],
                                prepared.review[0], cfg)
    next_rev = await _claim(record_id, revision, STATUS_WORKING, cfg.app)
    if next_rev is None:
        return JSONResponse(status_code=200,
                            content={"ok": True, "skip": "cas_lost"})
    await _generate_and_attach(record_id, prepared, next_rev, cfg)
    return JSONResponse(status_code=200, content={
        "ok": True, "record_id": record_id, "recovered": True})


async def _cloudsign_flow(record_id: str, record: dict, revision: str,
                          cfg: ChannelConfig | None = None):
    """CONTRACT-GEN-2: PDF 生成 → CloudSign 書類作成+PDF 添付+宛先追加 →
    doc id 書き戻し（登録済）。送信 API は呼ばない（対外効果の一線）。"""
    cfg = _resolve(cfg)
    prepared = await cfg.prepare(record, record_id, "register")
    # houki: 前提未充足（申述人集合・メール）は「要確認」へ CAS 遷移（状態を動かす）
    if prepared.review:
        return await _to_review(record_id, revision, prepared.review[1],
                                prepared.review[0], cfg)
    # fail-closed: 前提未充足は登録しない（状態も動かさない・値は非搭載）
    problems = list(prepared.missing)
    if not (record.get(FIELD_ATTACHMENT) or {}).get("value"):
        problems.append(f"{FIELD_ATTACHMENT}（docx 未添付＝先に"
                        f"「{STATUS_TRIGGER}」を実行してください）")
    for email, _name in prepared.participants:
        if not email:
            problems.append(FIELD_EMAIL)
        elif not _EMAIL_RE.fullmatch(email):
            problems.append(f"{FIELD_EMAIL}（形式不正）")
    # 特約の検査（上限/使用不可記号）と CloudSign 登録済みは dispatcher の
    # _regeneration_guard（3 経路共通）で「要確認」へ倒し済み（fix1 CT-01/02）
    if problems:
        logger.info("[CONTRACT] cloudsign preconditions unmet record_id=%s "
                    "count=%d",
                    emit(record_id, "record_id", "log", "operator"),
                    len(problems))
        await cfg.notify(
            "review", record_id,
            f"【委任契約書】案件 No.{record_id} はクラウドサイン登録の前提を"
            f"満たしていないため登録しませんでした。不足: "
            f"{'・'.join(problems)}。kintone で解消後、契約書ステータスを"
            f"「{STATUS_CS_TRIGGER}」に設定し直してください")
        return JSONResponse(status_code=200, content={
            "ok": True, "skip": "cs_preconditions", "missing": problems})

    # CAS（登録→登録中）勝者のみ実行（並行 2 本でも CloudSign 作成は 1 回）
    next_rev = await _claim(record_id, revision, STATUS_CS_WORKING, cfg.app)
    if next_rev is None:
        logger.info("[CONTRACT] cs cas lost record_id=%s",
                    emit(record_id, "record_id", "log", "operator"))
        return JSONResponse(status_code=200,
                            content={"ok": True, "skip": "cas_lost"})

    # HOUKI-CONTRACT-GEN: 送信直前に申述人集合を再取得（TOCTOU）。変化していれば
    # CloudSign を呼ばず「要確認」へ（登録中→要確認・CAS）
    if prepared.fingerprint:
        again = await cfg.prepare(record, record_id, "register")
        if again.review or again.fingerprint != prepared.fingerprint:
            return await _to_review(
                record_id, next_rev,
                f"【相続放棄 委任契約書・要確認】案件レコードNo.{record_id}: "
                "登録直前に申述人集合または署名者が変化したため CloudSign 登録を"
                f"中止しました。\n{again.summary}", "applicants_changed", cfg)

    # PDF 生成（テンプレ=単一の正）+ 凍結検証（docx 段・PDF 段の二層）。
    # 検証失敗はここで 500（系統的エラー→再配送が reconcile で「要確認」へ）
    docx_bytes = prepared.render()               # 特約込み（単一の生成経路）
    pdf_bytes = contract_pdf.docx_to_pdf_bytes(docx_bytes)
    verify_frozen_pdf(pdf_bytes, cfg)
    verify_pdf_full_text(docx_bytes, pdf_bytes)   # fix1[02] 全文 pin
    logger.info("[CONTRACT] pdf generated record_id=%s bytes=%d",
                emit(record_id, "record_id", "log", "operator"),
                len(pdf_bytes))

    # CloudSign 作成→添付→宛先。途中失敗は下書き削除（部分状態を残さない）
    # ＋掃除成功時のみ巻き戻し（登録中→登録）→ raise → 500 → 再配送で
    # 自動再試行。掃除失敗時は巻き戻さない（reconcile が「要確認」へ倒す）
    doc_id = None
    try:
        if cfg.name == "jikou":
            # 時効側は呼出し形も従来どおり（既存テストが引数形を pin）
            doc_id = _cs_create_document(record_id)
            _cs_attach_pdf(doc_id, pdf_bytes)
        else:
            doc_id = _cs_create_document(record_id, cfg.cs_title(record_id))
            _cs_attach_pdf(doc_id, pdf_bytes, cfg.output_pdf_name)
        for email, name in prepared.participants:
            _cs_add_participant(doc_id, email, name)
    except CloudSignResultUnknown:
        # fix1[01]: 作成結果不明＝下書きが存在し得るが id 不明。掃除も巻き
        # 戻しもせず「登録中」維持で 500 → 再配送は reconcile が「要確認」
        # +通知へ倒す（自動再作成の禁止＝二重下書き防止）
        logger.error("[CONTRACT] cloudsign create outcome unknown "
                     "record_id=%s",
                     emit(record_id, "record_id", "log", "operator"))
        raise
    except Exception:
        cleaned = _cs_delete_draft(doc_id) if doc_id else True
        if cleaned:
            try:
                await hub_kintone.update_record(
                    cfg.app, record_id, {FIELD_STATUS: STATUS_CS_TRIGGER},
                    revision=next_rev)
            except Exception:
                logger.error("[CONTRACT] cs status rollback failed "
                             "record_id=%s",
                             emit(record_id, "record_id", "log", "operator"))
        raise

    await hub_kintone.update_record(cfg.app, record_id, {
        FIELD_CS_DOC_ID: doc_id,
        FIELD_STATUS: STATUS_CS_DONE,
    }, revision=next_rev)
    logger.info("[CONTRACT] cloudsign registered record_id=%s",
                emit(record_id, "record_id", "log", "operator"))
    await cfg.notify("created", record_id,
                     f"【相続放棄 委任契約書】案件レコードNo.{record_id} を CloudSign に"
                     "下書き登録しました（送信は CloudSign 画面で行ってください）。"
                     f"\n{prepared.summary}")
    return JSONResponse(status_code=200, content={
        "ok": True, "record_id": record_id, "cloudsign": True})


async def _reconcile_cs_working(record_id: str, revision: str,
                                cfg: ChannelConfig | None = None):
    """CONTRACT-GEN-2 reconcile: 「クラウドサイン登録中」で停止した行の回収。

    CloudSign 側に下書きが残っている可能性がある（外部状態・kintone からは
    機械確認できない）ため、v1 の回収と異なり自動再実行はせず、常に CAS で
    「要確認」へ倒して管理者通知（二重下書き防止の fail-closed）。"""
    cfg = _resolve(cfg)
    next_rev = await _claim(record_id, revision, STATUS_REVIEW, cfg.app)
    if next_rev is None:
        return JSONResponse(status_code=200,
                            content={"ok": True, "skip": "cas_lost"})
    await cfg.notify(
        "review", record_id,
        f"【委任契約書】案件 No.{record_id} はクラウドサイン登録が中断した"
        "状態のため「要確認」にしました。CloudSign 画面で下書きの有無を確認し"
        "（重複下書きがあれば削除）、再実行する場合は契約書ステータスを"
        f"「{STATUS_CS_TRIGGER}」に設定し直してください")
    return JSONResponse(status_code=200,
                        content={"ok": True, "skip": "cs_needs_review"})


async def _dispatch(body, cfg: ChannelConfig):
    """token 検証後の共通経路（app 同一性 → record id → 本文ステータス gate →
    正本の完全一致検証 → 各経路）。"""
    # fix1[01]: app 同一性——実 app ID と完全一致（欠落・非数字・別 App は
    # get_record 含め作用 0）
    app_in_body = str(((body.get("app") or {}) if isinstance(body, dict)
                       else {}).get("id") or "")
    if not app_in_body.isdigit() or app_in_body != str(cfg.app.app_id()):
        logger.warning("[CONTRACT] app mismatch in webhook body")
        return JSONResponse(status_code=200,
                            content={"ok": True, "skip": "app_mismatch"})

    record_id = extract_record_id(body)
    if not record_id:
        logger.warning("[CONTRACT] record id missing in webhook body")
        return JSONResponse(status_code=200,
                            content={"ok": True, "skip": "no_record_id"})

    # 本文ステータス gate（自 update の echo=作成中/作成済/登録中/登録済 も
    # ここで落ちる）。通過はトリガ 2 値（契約書作成/クラウドサイン登録）のみ
    try:
        status_in_webhook = body["record"][FIELD_STATUS]["value"]
    except (KeyError, TypeError):
        status_in_webhook = None
    if status_in_webhook not in (STATUS_TRIGGER, STATUS_CS_TRIGGER):
        logger.info("[CONTRACT] not triggered record_id=%s",
                    emit(record_id, "record_id", "log", "operator"))
        return JSONResponse(status_code=200,
                            content={"ok": True, "skip": "not_triggered"})

    try:
        record = await hub_kintone.get_record(cfg.app, record_id)
        current = _fv(record, FIELD_STATUS)
        revision = _fv(record, "$revision")

        # fix1[01]: 正本の完全一致検証（stale 本文は作用 0 で skip）。
        # dispatch は本文でなく正本ステータスに対して行う
        if current in (STATUS_DONE, STATUS_CS_DONE):
            logger.info("[CONTRACT] already done record_id=%s",
                        emit(record_id, "record_id", "log", "operator"))
            return JSONResponse(status_code=200,
                                content={"ok": True, "skip": "already_done"})
        if not revision.isdigit():
            logger.info("[CONTRACT] stale status record_id=%s",
                        emit(record_id, "record_id", "log", "operator"))
            return JSONResponse(status_code=200,
                                content={"ok": True, "skip": "stale_status"})
        if current == STATUS_WORKING:
            return await _reconcile_working(record_id, record, revision, cfg)
        if current == STATUS_CS_WORKING:
            return await _reconcile_cs_working(record_id, revision, cfg)
        if current == STATUS_CS_TRIGGER:
            # fix1 CT-01: 登録経路にも共通ガード（cs_registered=再登録の文言）。
            # 該当は CloudSign API 作用 0 で「要確認」へ CAS 遷移+通知
            guard = _regeneration_guard(record_id, record, "register")
            if guard:
                return await _to_review(record_id, revision, guard[1], guard[0], cfg)
            return await _cloudsign_flow(record_id, record, revision, cfg)
        if current != STATUS_TRIGGER:
            logger.info("[CONTRACT] stale status record_id=%s",
                        emit(record_id, "record_id", "log", "operator"))
            return JSONResponse(status_code=200,
                                content={"ok": True, "skip": "stale_status"})

        prepared = await cfg.prepare(record, record_id, "generate")
        # fail-closed: 必須欠落は生成しない（状態も動かさない）
        missing = prepared.missing
        if missing:
            logger.info("[CONTRACT] missing required fields record_id=%s "
                        "count=%d",
                        emit(record_id, "record_id", "log", "operator"),
                        len(missing))
            await cfg.notify(
                "review", record_id,
                f"【委任契約書】案件 No.{record_id} は必須項目が未入力のため"
                f"生成しませんでした。不足: {'・'.join(missing)}。"
                "kintone で入力後、契約書ステータスを"
                f"「{STATUS_TRIGGER}」に設定し直してください")
            return JSONResponse(status_code=200, content={
                "ok": True, "skip": "missing_fields", "missing": missing})

        # JIKOU-CONTRACT-TOKUYAKU: 再生成の運用ガード（①CloudSign 下書き作成済み
        # → ②特約上限 → ③使用不可記号）。該当は生成せず「要確認」へ CAS 遷移+通知
        guard = _regeneration_guard(record_id, record, "regenerate")
        if guard:
            return await _to_review(record_id, revision, guard[1], guard[0], cfg)
        # houki: 申述人集合・被相続人氏名・顧客名/住所の前提未充足は「要確認」へ
        if prepared.review:
            return await _to_review(record_id, revision, prepared.review[1],
                                    prepared.review[0], cfg)

        # fix1[02]: CAS（作成→作成中）勝者のみ生成（並行 2 本でも upload 1 回）
        next_rev = await _claim(record_id, revision, STATUS_WORKING, cfg.app)
        if next_rev is None:
            logger.info("[CONTRACT] cas lost record_id=%s",
                        emit(record_id, "record_id", "log", "operator"))
            return JSONResponse(status_code=200,
                                content={"ok": True, "skip": "cas_lost"})
        await _generate_and_attach(record_id, prepared, next_rev, cfg)
    except Exception as e:
        logger.error("[CONTRACT] error record_id=%s cls=%s: %s",
                     emit(record_id, "record_id", "log", "operator"),
                     type(e).__name__,
                     emit(str(e), "vendor_raw", "log", "operator"))
        await cfg.notify("failed", record_id,
                         f"【相続放棄 委任契約書・失敗】案件レコードNo.{record_id} の処理で"
                         "内部エラーが発生しました（kintone の再配送で再試行されます。"
                         "契約書ステータスを確認してください）。")
        return JSONResponse(status_code=500,
                            content={"error": "internal_error"})

    return JSONResponse(status_code=200,
                        content={"ok": True, "record_id": record_id})


@router.post("/contract/{secret}")
async def contract_webhook(secret: str, request: Request):
    if not verify_token(secret or "", "DOCUMENT_WEBHOOK_SECRET"):
        return JSONResponse(status_code=403, content={"error": "forbidden"})

    try:
        body = await request.json()
    except Exception:
        return JSONResponse(status_code=400, content={"error": "invalid json"})
    return await _dispatch(body, _jikou_cfg())


@router.post("/souzoku-houki/contract/{secret}")
async def houki_contract_webhook(secret: str, request: Request):
    """HOUKI-CONTRACT-GEN: App 40 の受け口。token は相談カード読取・申述書と同じ
    HOUKI_WEBHOOK_TOKEN・同じ検証方式（未設定/時効側同値=404・不一致=403）。"""
    from shinjutsu_webhook import _TOKEN_ENV, houki_webhook_disabled_reason
    reason = houki_webhook_disabled_reason()
    if reason is not None:
        if reason != "token_unset":
            logger.warning("[HOUKI_CONTRACT] endpoint disabled (token misconfig)")
        return JSONResponse(status_code=404, content={"error": "not found"})
    if not hmac.compare_digest(secret or "", os.environ.get(_TOKEN_ENV, "")):
        return JSONResponse(status_code=403, content={"error": "forbidden"})

    try:
        body = await request.json()
    except Exception:
        return JSONResponse(status_code=400, content={"error": "invalid json"})
    return await _dispatch(body, _houki_cfg())
