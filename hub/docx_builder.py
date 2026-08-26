"""docx 生成の共通実装（hub/docx_builder）

設計: docs/architecture/03-common-components.md §6

- fill_template / to_wareki: document_webhook.py から移設（T0-3・実装変更なし）。
  既存の import 経路は document_webhook 側の re-export で維持される。
- resolve_template: docx_templates/<ユニットの template_dir>/<種別>.docx の規約解決
- validate_template: テンプレート内に必要プレースホルダ（{{...}}）が揃っているか検査
  （daily_healthcheck に登録し、人がテンプレートを編集して差込キーを消した事故を検知）
"""

import io
import re
from datetime import date
from pathlib import Path

from docx import Document

from config import UNIT_CONFIG

# テンプレート規約のルートディレクトリ（リポジトリ相対・既存の配置と同じ）
TEMPLATE_ROOT = "docx_templates"

_PLACEHOLDER_RE = re.compile(r"\{\{[^{}]+\}\}")


class TemplateNotFound(Exception):
    """テンプレートファイルが規約の場所に存在しない / ユニットが未定義"""


def to_wareki(d: date) -> str:
    """西暦→和暦表記（document_webhook から移設・実装不変）"""
    if d >= date(2019, 5, 1):
        n, era = d.year - 2018, "令和"
    elif d >= date(1989, 1, 8):
        n, era = d.year - 1988, "平成"
    else:
        return d.strftime("%Y年%m月%d日")
    y = "元" if n == 1 else str(n)
    return f"{era}{y}年{d.month}月{d.day}日"


def fill_runs(para, mapping: dict) -> None:
    """run 単位でプレースホルダを差し込む（段落の run 構成・各 run の rPr を
    保持する。SOUZOKU-HOUKI-H7A で notice_webhook._fill_runs〔JIKOU-NOTICE-1-
    fix2 で確立〕から逐語昇格・実装不変）。

    段落全体を先頭 run へ潰す方式（fill_template）は、ラベル run の書式
    （均等割り付け w:fitText 等）が行全体へ及び崩れる（時効援用通知書の
    ふりがな・生年月日行で実機発見）。本関数は各 run の rPr を保ったまま、
    プレースホルダを含む run の中だけで置換する。

    契約: テンプレートはプレースホルダを**単一 run に収める**こと（収載時に
    検査する・run 跨ぎのプレースホルダは置換されず残るため、呼び出し側の
    残存検査〔fail-closed〕で拒否される）。新規テンプレ（相続放棄 H-7(b)(c)
    以降）は本関数を使う。
    """
    for r in para.runs:
        if "{{" in r.text:
            text = r.text
            for k, v in mapping.items():
                text = text.replace(k, v)
            r.text = text


def fill_template(template_path: str, data: dict) -> bytes:
    """テンプレートを差し込み置換して docx の bytes を返す
    （document_webhook から移設・実装不変。run 分割されたプレースホルダにも対応）

    既知問題（JIKOU-NOTICE-1-fix2 で実機発見・H7A で明文化）: 段落の全 run を
    先頭 run へ潰すため、run 単位の書式（均等割り付け w:fitText・font 差等）が
    行全体へ波及して崩れる。既存利用者（契約書・送付状等＝この挙動で SHA/凍結
    pin 済み）の互換のため実装は変更しない。**新規テンプレートは run 保持形の
    fill_runs を使うこと**（プレースホルダは単一 run に収める契約）。"""
    doc = Document(template_path)

    def replace_in_paragraph(para):
        full = "".join(run.text for run in para.runs)
        if not any(k in full for k in data):
            return
        for k, v in data.items():
            full = full.replace(k, v)
        if para.runs:
            para.runs[0].text = full
            for run in para.runs[1:]:
                run.text = ""

    for para in doc.paragraphs:
        replace_in_paragraph(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _replace_multiline_in_paragraph(para, data: dict) -> None:
    full = "".join(run.text for run in para.runs)
    if not any(k in full for k in data):
        return
    for k, v in data.items():
        full = full.replace(k, str(v))
    if not para.runs:
        return
    lines = full.split("\n")
    para.runs[0].text = lines[0]
    for run in para.runs[1:]:
        run.text = ""
    for line in lines[1:]:
        br = para.add_run()
        br.add_break()
        para.add_run(line)


def _apply_multiline(doc: Document, data: dict) -> None:
    for para in doc.paragraphs:
        _replace_multiline_in_paragraph(para, data)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_multiline_in_paragraph(para, data)


def fill_template_multiline(template_path: str, data: dict) -> bytes:
    """fill_template の複数行対応版（T2-1 で追加。既存 fill_template は不変）。

    値に改行（\\n）を含むプレースホルダを Word の改行（<w:br/>）として差し込む。
    同封物一覧のような「複数行を1プレースホルダで渡す」用途（設計 07 §2）に使う。
    ※ 2行目以降の run は段落既定の書式になる（雛形は段落書式で整えること）
    """
    doc = Document(template_path)
    _apply_multiline(doc, data)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def fill_table_rows(doc: Document, rows: list[dict], marker_prefix: str = "{{行:",
                    *, row_marker: str | None = None,
                    empty_text: str | None = None) -> None:
    """{{行:列名}} を含むテンプレート行を rows の件数分複製して差し込む
    （souzoku-shorui 02 §2・S2。T2-2 テンプレ統合〔送付案内の書類表〕で使用中）。

    - テンプレート行の書式・罫線は行の複製（deepcopy）で継承される
    - row_marker（例 "{{行:財産}}"）を指定すると、そのマーカーを含む行だけを対象にする。
      可変表が複数ある帳票（財産目録の種別セクション等）はセクションごとに
      マーカーを変えて本関数を呼び分ける。複製行ではマーカー自体は空文字になる
      （rows の列名と一致する場合はその値が優先）
    - row_marker 未指定時は最初に見つかったテンプレート行1つが対象（従来互換）
    - rows が空のときはテンプレート行を削除する。empty_text を指定した場合は
      書式を継承した1行を残し、row_marker の位置（未指定時は行内の最初の
      プレースホルダ）に empty_text を差し込み、他のプレースホルダは空欄にする
    """
    import copy as _copy

    from docx.table import _Row

    needle = row_marker if row_marker is not None else marker_prefix
    for table in doc.tables:
        for row in table.rows:
            if not any(needle in cell.text for cell in row.cells):
                continue
            template_tr = row._tr

            def _stamp(data: dict) -> None:
                new_tr = _copy.deepcopy(template_tr)
                template_tr.addprevious(new_tr)
                new_row = _Row(new_tr, table)
                for cell in new_row.cells:
                    for para in cell.paragraphs:
                        _replace_multiline_in_paragraph(para, data)

            if not rows and empty_text is not None:
                row_text = "\n".join(cell.text for cell in row.cells)
                placeholders = [ph for ph in _PLACEHOLDER_RE.findall(row_text)
                                if ph.startswith(marker_prefix)]
                data = {ph: "" for ph in placeholders}
                anchor = row_marker if row_marker is not None else (
                    placeholders[0] if placeholders else None)
                if anchor is not None:
                    data[anchor] = empty_text
                _stamp(data)
            for item in rows:
                data = {marker_prefix + k + "}}": str(v) for k, v in item.items()}
                if row_marker is not None:
                    data.setdefault(row_marker, "")
                _stamp(data)
            template_tr.getparent().remove(template_tr)
            return
    # テンプレート行が見つからない場合は何もしない（validate_template で検知する）


def fill_template_with_table(template_path: str, data: dict,
                             table_rows: list[dict],
                             marker_prefix: str = "{{行:") -> bytes:
    """スカラー差込（複数行対応）＋表の行複製をまとめて行う（送付案内の正式書式用）"""
    doc = Document(template_path)
    _apply_multiline(doc, data)
    fill_table_rows(doc, table_rows, marker_prefix=marker_prefix)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def resolve_template(unit: str, doc_type: str, base_dir: str = TEMPLATE_ROOT) -> Path:
    """規約ベースのテンプレート解決: <base_dir>/<UNIT_CONFIG[unit].template_dir>/<doc_type>.docx
    ユニット未定義・ファイル不存在は TemplateNotFound"""
    conf = UNIT_CONFIG.get(unit)
    if conf is None:
        raise TemplateNotFound(f"ユニット未定義: {unit}（config.UNIT_CONFIG に登録してください）")
    path = Path(base_dir) / conf["template_dir"] / f"{doc_type}.docx"
    if not path.is_file():
        raise TemplateNotFound(f"テンプレートがありません: {path}")
    return path


def _extract_text(path) -> str:
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return "\n".join(parts)


def validate_template(path, required_keys: list[str]) -> list[str]:
    """テンプレートに必要プレースホルダが揃っているか検査し、欠けているキーの一覧を返す
    （空リスト = 問題なし）。ファイル不存在は TemplateNotFound。"""
    p = Path(path)
    if not p.is_file():
        raise TemplateNotFound(f"テンプレートがありません: {p}")
    text = _extract_text(p)
    return [k for k in required_keys if k not in text]


def list_placeholders(path) -> list[str]:
    """テンプレート内の {{...}} プレースホルダを列挙する（レジストリ整備の補助）"""
    return sorted(set(_PLACEHOLDER_RE.findall(_extract_text(path))))
