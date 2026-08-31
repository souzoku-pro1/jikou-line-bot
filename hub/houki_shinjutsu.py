"""相続放棄申述書の生成（SOUZOKU-HOUKI-H7C・hub 側モジュール）

テンプレ=docx_templates/houki/相続放棄申述書.docx（公式様式を
scripts/make_shinjutsu_template.py でプレースホルダ化・SHA-256 pin）。
差し込みは hub/docx_builder.fill_runs（run 単位置換・rPr 保持）一本。
fill_template（段落潰し）は使用禁止（docx_builder の docstring どおり）。

弁護士決定（凍結・改変禁止）:
- マル囲み選択（番号系）は丸数字置換（１→①〜７→⑦の閉じたマッピング・
  1 グループにつき丸数字は必ず 1 つ=複数該当は構造的に生じない〔単一値
  DROP_DOWN 由来〕・値が空のグループは無印のまま=手書き運用）
- 語句系の選択（元号・都道府県併記）は触らない（印刷後の手書き運用）
- 「相続財産の概略」欄なしの様式のまま提出運用

凍結検証（JIKOU-NOTICE-1 確立形の同型）:
- テンプレは収載現物の SHA-256 と完全一致（TEMPLATE_SHA256）
- 生成物の全段落（本文+全テーブルセル）が「テンプレへ同一の差し込み+
  sanctioned substitution（許可した丸数字置換のみ）を適用した期待列」と
  完全一致・プレースホルダ残存なし。不一致は ShinjutsuIntegrityError

fail-closed（票の逐語）:
- 必須差し込み値（申述人氏名・被相続人氏名・死亡日）欠落=生成拒否
- 相続の開始を知った日が導出不能（App 40 フィールド空/不正）=生成拒否
- マッピング不能（閉集合外の非空値・元号が様式の不動文字と両立しない
  日付・管轄家庭裁判所が「家庭裁判所」を含まない非空値）=生成拒否
拒否理由は closed vocabulary（ShinjutsuRejection.reasons）で返し、通知は
受け口側が行う（フィールド名のみ・PII 非搭載）。
値が空なだけの任意項目は空欄のまま生成する（生成は止めない・手書き運用）。
"""

import datetime
import hashlib
import io
import re

from docx import Document
from docx.table import _Cell

from hub.docx_builder import fill_runs

TEMPLATE_PATH = "docx_templates/houki/相続放棄申述書.docx"
# 収載現物（scripts/make_shinjutsu_template.py 生成・コミット済み artifact）の
# pin。再収載は zip タイムスタンプで SHA が変わるため、pin はコミット現物。
# 役割は「コミット物の同一性」（fix1[H7C-01]: 正本性の防壁は canonical 側）
TEMPLATE_SHA256 = (
    "a6bbc4d12da1699adff39daf0612e3b24a80ee81061a90e826b254a69d996c46")
# fix1[H7C-01]: 正本性の防壁=正規化済み内容の pin（収載スクリプト由来の独立
# 定数・zip タイムスタンプ非依存）。canonical manifest は全段落テキスト全量+
# run 分割（プレースホルダの位置・個数を含む）+rPr 要点（fitText/sz/u）を
# 決定的に直列化した SHA-256——不動文字の 1 文字改変・プレースホルダの増減・
# 位置替え・ラベル書式の変更はいずれか（多くは複数）の行を変え、必ず検出
# される。「テンプレ改変+TEMPLATE_SHA256 追随」の同時変更はこの定数を
# 変えない限り通らず、この定数の変更は票由来のレビュー対象になる。
# 意図的更新の手順は scripts/make_shinjutsu_template.py の docstring を参照
TEMPLATE_CANONICAL_SHA256 = (
    "c16dfb76e9d7e4be3bd843d3877df209a222a816258f9b04df85d9b8f9e4f88f")

_FW = "　"
_JST = datetime.timezone(datetime.timedelta(hours=9))


class ShinjutsuIntegrityError(RuntimeError):
    """テンプレ pin 不一致・生成物の凍結検証失敗（添付しない）。"""


class ShinjutsuRejection(RuntimeError):
    """fail-closed の生成拒否（理由は閉集合語彙のリスト）。"""

    def __init__(self, reasons: list[str]):
        super().__init__(" / ".join(reasons))
        self.reasons = reasons


# ── 丸数字置換（sanctioned substitution・弁護士決定の閉じたマッピング） ───────────
_CIRCLED = {"１": "①", "２": "②", "３": "③", "４": "④",
            "５": "⑤", "６": "⑥", "７": "⑦"}

# App 40 実選択肢（form fields API 実測・逐語）→ 様式の番号
KANKEI_MAP = {"子": "１", "孫": "２", "配偶者": "３",
              "直系尊属（父母・祖父母）": "４", "兄弟姉妹": "５",
              "おいめい": "６", "その他": "７"}
SHITTA_MAP = {"被相続人死亡の当日": "１", "死亡の通知をうけた日": "２",
              "先順位者の相続放棄を知った日": "３", "その他": "４"}
RIYU_MAP = {"被相続人から生前に贈与を受けている。": "１",
            "生活が安定している。": "２", "遺産が少ない。": "３",
            "遺産を分散させたくない。": "４", "債務超過のため。": "５",
            "その他": "６"}

# 丸数字を打つセルの番地（テンプレ構造は SHA pin で固定・(table, row, tc)）
_CIRCLE_CELLS = {"kankei": (4, 3, 2), "shitta": (6, 1, 0), "riyu": (6, 3, 0)}

# 値が空のときにプレースホルダへ戻す既定文字列（様式の元の空白・□。
# 空欄=元の見た目のまま印刷される）
PLACEHOLDER_DEFAULTS = {
    "{{裁判所前}}": "", "{{裁判所後}}": "",
    "{{提出年}}": _FW * 2, "{{提出月}}": _FW * 2, "{{提出日}}": _FW * 2,
    "{{記名}}": _FW * 14,
    "{{チェック戸籍}}": "□", "{{戸籍通数}}": _FW * 2, "{{チェック除票}}": "□",
    "{{郵便前}}": _FW * 2, "{{郵便後}}": _FW * 2,
    "{{電話1}}": _FW, "{{電話2}}": _FW * 4, "{{電話3}}": " " * 4,
    "{{申述人住所}}": _FW * 23,
    "{{申述人フリガナ}}": _FW, "{{申述人氏名}}": _FW * 2,
    "{{生年}}": _FW, "{{生月}}": _FW, "{{生日}}": _FW, "{{年齢}}": _FW * 4,
    "{{続柄その他}}": _FW * 9,
    "{{被相続人本籍}}": "", "{{被相続人最後の住所}}": "",
    "{{被相続人フリガナ}}": "", "{{被相続人氏名}}": "",
    "{{死亡年}}": _FW * 3, "{{死亡月}}": _FW * 3, "{{死亡日}}": _FW * 3,
    "{{知年}}": _FW * 3, "{{知月}}": _FW * 3, "{{知日}}": _FW * 3,
    "{{知った日区分その他}}": _FW * 10,
}

# 拒否理由の閉集合（通知にはこの語彙+レコード番号のみ・PII 非搭載）
REJECT_MISSING_NAME = "申述人氏名（顧客名）未入力"
REJECT_MISSING_DECEASED = "被相続人氏名 未入力"
REJECT_MISSING_DEATH = "死亡日_申告 未入力または形式不正"
REJECT_SHITTA_MISSING = "相続の開始を知った日 未入力（導出不能）"
REJECT_SHITTA_ERA = "相続の開始を知った日が令和より前（様式の不動文字と両立しない）"
REJECT_DEATH_ERA = "死亡日_申告が平成より前（様式の不動文字と両立しない）"
REJECT_KANKEI_UNMAPPED = "続柄が様式の選択肢へ対応付けできない"
REJECT_SHITTA_KUBUN_UNMAPPED = "知った日の区分が様式の選択肢へ対応付けできない"
REJECT_RIYU_UNMAPPED = "放棄の理由が様式の選択肢へ対応付けできない"
REJECT_COURT_UNMAPPED = "管轄家庭裁判所に「家庭裁判所」が含まれない"


def _fv(record: dict, code: str) -> str:
    return str(((record or {}).get(code) or {}).get("value") or "").strip()


def _zen(n) -> str:
    return str(n).translate(str.maketrans("0123456789", "０１２３４５６７８９"))


def _parse_date(raw: str) -> datetime.date | None:
    try:
        return datetime.date.fromisoformat(raw)
    except ValueError:
        return None


def _wareki(d: datetime.date) -> tuple[str, int] | None:
    """(元号, 年数)。昭和より前は None（様式に元号がない）。"""
    if d >= datetime.date(2019, 5, 1):
        return "令和", d.year - 2018
    if d >= datetime.date(1989, 1, 8):
        return "平成", d.year - 1988
    if d >= datetime.date(1926, 12, 25):
        return "昭和", d.year - 1925
    return None


def _today_jst() -> datetime.date:
    return datetime.datetime.now(_JST).date()


def _split_zip(address: str) -> tuple[str, str, str]:
    """住所文字列から郵便番号を決定的に分離（先頭の 〒?123-4567 形のみ）。
    分離できなければ (空, 空, 全文)。"""
    m = re.match(r"^〒?\s*(\d{3})[-－ー]\s*(\d{4})\s*(.*)$", address)
    if m:
        return _zen(m.group(1)), _zen(m.group(2)), m.group(3)
    return "", "", address


def _split_phone(phone: str) -> tuple[str, str, str]:
    """電話番号をハイフンで 3 分割（できなければ全文を第 1 枠へ）。"""
    parts = [p for p in re.split(r"[-－ー]", phone) if p]
    if len(parts) == 3:
        return _zen(parts[0]), _zen(parts[1]), _zen(parts[2])
    return _zen(phone), "", ""


def _split_court(court: str) -> tuple[str, str]:
    """管轄家庭裁判所を（前, 後）へ分割（「家庭裁判所」の初出で分ける）。
    含まない非空値は ShinjutsuRejection（デタラメを刷らない）。"""
    if not court:
        return "", ""
    if "家庭裁判所" not in court:
        raise ShinjutsuRejection([REJECT_COURT_UNMAPPED])
    before, after = court.split("家庭裁判所", 1)
    return before, after


def _koseki_attachment(record: dict) -> tuple[str, str, str]:
    """書類チェック SUBTABLE から添付書類欄の値を決定的に導出する。

    採用規則（実測選択肢の閉集合・書類状態=受領 の行のみ数える。手元にない
    書類を添付宣言しない fail-safe。同封対象は発送管理〔M4〕の意味論のため
    使わない）:
      - 戸籍チェック+通数: 書類名 ∈ {被相続人死亡記載戸籍, 申述人戸籍}
      - 除票チェック: 書類名 = 被相続人住民票除票
    該当 0 件は □・空欄のまま（手書き運用）。"""
    rows = ((record.get("書類チェック") or {}).get("value")) or []
    koseki = 0
    johyo = 0
    for row in rows:
        v = row.get("value") or {}
        name = str(((v.get("書類名") or {}).get("value")) or "")
        state = str(((v.get("書類状態") or {}).get("value")) or "")
        if state != "受領":
            continue
        if name in ("被相続人死亡記載戸籍", "申述人戸籍"):
            koseki += 1
        elif name == "被相続人住民票除票":
            johyo += 1
    check_k = "■" if koseki else "□"
    count_k = f"{_FW}{_zen(koseki)}{_FW}" if koseki else _FW * 2
    check_j = "■" if johyo else "□"
    return check_k, count_k, check_j


def _circles(record: dict) -> tuple[dict, list[str]]:
    """レコード実値 → 各グループの丸数字対象（{group: 数字}・空値グループは
    含めない=無印）と、マッピング不能理由の一覧。"""
    out: dict = {}
    reasons: list[str] = []
    for group, code, mapping, reject in (
            ("kankei", "続柄", KANKEI_MAP, REJECT_KANKEI_UNMAPPED),
            ("shitta", "知った日の区分", SHITTA_MAP,
             REJECT_SHITTA_KUBUN_UNMAPPED),
            ("riyu", "放棄の理由", RIYU_MAP, REJECT_RIYU_UNMAPPED)):
        value = _fv(record, code)
        if not value:
            continue
        digit = mapping.get(value)
        if digit is None:
            reasons.append(reject)
        else:
            out[group] = digit
    return out, reasons


def build_fill_data(record: dict,
                    today: datetime.date | None = None) -> dict:
    """App 40 レコード → プレースホルダ全 33 種の値（空は様式の既定空白）。
    fail-closed 条件はここで集約して ShinjutsuRejection にまとめて送出する。"""
    today = today or _today_jst()
    reasons: list[str] = []

    name = _fv(record, "顧客名")
    deceased = _fv(record, "被相続人氏名")
    death = _parse_date(_fv(record, "死亡日_申告"))
    if not name:
        reasons.append(REJECT_MISSING_NAME)
    if not deceased:
        reasons.append(REJECT_MISSING_DECEASED)
    if death is None:
        reasons.append(REJECT_MISSING_DEATH)

    shitta = _parse_date(_fv(record, "相続の開始を知った日"))
    if shitta is None:
        reasons.append(REJECT_SHITTA_MISSING)
    elif shitta < datetime.date(2019, 5, 1):
        reasons.append(REJECT_SHITTA_ERA)     # 様式は「令和」不動文字

    death_w = _wareki(death) if death else None
    if death is not None and (death_w is None or death_w[0] == "昭和"):
        reasons.append(REJECT_DEATH_ERA)      # 様式は「平成・令和」不動文字

    circles, circle_reasons = _circles(record)
    reasons.extend(circle_reasons)

    try:
        court_pre, court_post = _split_court(_fv(record, "管轄家庭裁判所"))
    except ShinjutsuRejection as e:
        reasons.extend(e.reasons)
        court_pre = court_post = ""

    if reasons:
        raise ShinjutsuRejection(reasons)

    fill = dict(PLACEHOLDER_DEFAULTS)

    submit_w = _wareki(today)
    fill["{{提出年}}"] = _zen(submit_w[1])
    fill["{{提出月}}"] = _zen(today.month)
    fill["{{提出日}}"] = _zen(today.day)
    fill["{{裁判所前}}"] = court_pre
    fill["{{裁判所後}}"] = court_post
    fill["{{記名}}"] = name

    zip3, zip4, addr_rest = _split_zip(_fv(record, "住所"))
    if zip3:
        fill["{{郵便前}}"] = zip3
        fill["{{郵便後}}"] = zip4
    fill["{{申述人住所}}"] = addr_rest or PLACEHOLDER_DEFAULTS["{{申述人住所}}"]
    tel = _fv(record, "電話番号")
    if tel:
        t1, t2, t3 = _split_phone(tel)
        fill["{{電話1}}"] = t1 or PLACEHOLDER_DEFAULTS["{{電話1}}"]
        if t2:
            fill["{{電話2}}"] = t2
        if t3:
            fill["{{電話3}}"] = t3
    furigana = _fv(record, "furigana")
    if furigana:
        fill["{{申述人フリガナ}}"] = furigana
    fill["{{申述人氏名}}"] = name

    birth = _parse_date(_fv(record, "生年月日"))
    if birth:
        birth_w = _wareki(birth)
        if birth_w:
            fill["{{生年}}"] = _zen(birth_w[1])
            fill["{{生月}}"] = _zen(birth.month)
            fill["{{生日}}"] = _zen(birth.day)
            age = (today.year - birth.year
                   - ((today.month, today.day) < (birth.month, birth.day)))
            fill["{{年齢}}"] = _zen(age)

    zoku_sonota = _fv(record, "続柄その他")
    if circles.get("kankei") == "７" and zoku_sonota:
        fill["{{続柄その他}}"] = zoku_sonota

    honseki = _fv(record, "被相続人本籍")
    if honseki:
        fill["{{被相続人本籍}}"] = honseki
    last_addr = _fv(record, "被相続人最後の住所")
    if last_addr:
        fill["{{被相続人最後の住所}}"] = last_addr
    d_furi = _fv(record, "被相続人ふりがな")
    if d_furi:
        fill["{{被相続人フリガナ}}"] = d_furi
    fill["{{被相続人氏名}}"] = deceased

    fill["{{死亡年}}"] = _zen(death_w[1])
    fill["{{死亡月}}"] = _zen(death.month)
    fill["{{死亡日}}"] = _zen(death.day)

    shitta_w = _wareki(shitta)
    fill["{{知年}}"] = _zen(shitta_w[1])
    fill["{{知月}}"] = _zen(shitta.month)
    fill["{{知日}}"] = _zen(shitta.day)
    kubun_sonota = _fv(record, "知った日の区分その他")
    if circles.get("shitta") == "４" and kubun_sonota:
        fill["{{知った日区分その他}}"] = kubun_sonota

    check_k, count_k, check_j = _koseki_attachment(record)
    fill["{{チェック戸籍}}"] = check_k
    fill["{{戸籍通数}}"] = count_k
    fill["{{チェック除票}}"] = check_j

    return fill


# ── docx の組み立て・凍結検証 ────────────────────────────────────────────────
_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _walk_paragraphs(doc):
    yield from doc.paragraphs
    for tb in doc.tables:
        for row in tb.rows:
            for tc in row._tr.tc_lst:
                yield from _Cell(tc, tb).paragraphs


def canonical_manifest_text(doc) -> str:
    """テンプレ内容の決定的な直列化（fix1[H7C-01]・zip タイムスタンプ非依存）。

    粒度: 段落ごとに (a) 段落テキスト全量、(b) 各 run のテキスト（＝
    プレースホルダの個数・位置・run 分割）、(c) rPr 要点（fitText の有無・
    sz・下線）。不動文字の 1 文字改変は (a)(b) を、プレースホルダの増減・
    位置替えは (b) を、ラベル書式（均等割り付け等）の変更は (c) を変える。"""
    lines = []
    for i, p in enumerate(_walk_paragraphs(doc)):
        lines.append(f"P{i}\t{p.text}")
        for j, r in enumerate(p.runs):
            rpr = r._r.find(_W_NS + "rPr")
            fit = sz = und = "-"
            if rpr is not None:
                if rpr.find(_W_NS + "fitText") is not None:
                    fit = "fit"
                sz_el = rpr.find(_W_NS + "sz")
                if sz_el is not None:
                    sz = sz_el.get(_W_NS + "val") or "-"
                if rpr.find(_W_NS + "u") is not None:
                    und = "u"
            lines.append(f"R{i}.{j}\t{fit}\t{sz}\t{und}\t{r.text}")
    return "\n".join(lines)


def canonical_sha256(doc) -> str:
    return hashlib.sha256(
        canonical_manifest_text(doc).encode("utf-8")).hexdigest()


def verify_template_integrity() -> None:
    """コミット物の同一性（バイナリ SHA）+正本性（canonical pin・fix1）。"""
    data = open(TEMPLATE_PATH, "rb").read()
    if hashlib.sha256(data).hexdigest() != TEMPLATE_SHA256:
        raise ShinjutsuIntegrityError("template hash mismatch")
    if canonical_sha256(Document(TEMPLATE_PATH)) != TEMPLATE_CANONICAL_SHA256:
        raise ShinjutsuIntegrityError("template canonical mismatch")


def _circle_cell(doc, addr: tuple, digit: str) -> None:
    """指定セル内の全角数字 digit（唯一）を丸数字へ置換（run 内文字置換・
    rPr 不変）。唯一でなければ IntegrityError（様式構造の想定崩れ）。"""
    ti, ri, ci = addr
    cell = _Cell(doc.tables[ti].rows[ri]._tr.tc_lst[ci], doc.tables[ti])
    hits = [r for p in cell.paragraphs for r in p.runs if digit in r.text]
    total = sum(r.text.count(digit) for r in hits)
    if total != 1:
        raise ShinjutsuIntegrityError("circle target not unique")
    hits[0].text = hits[0].text.replace(digit, _CIRCLED[digit])


def build_shinjutsu_docx(fill: dict, circles: dict) -> bytes:
    """テンプレへ丸数字置換→fill_runs（run 単位・rPr 保持）の順で適用。

    丸数字が先: 対象数字の一意性は SHA pin 済みテンプレに対して検査する
    （差し込み値に全角数字が含まれてもセル内一意性が壊れない順序）。"""
    doc = Document(TEMPLATE_PATH)
    for group, digit in circles.items():
        _circle_cell(doc, _CIRCLE_CELLS[group], digit)
    for p in _walk_paragraphs(doc):
        if "{{" in p.text:
            fill_runs(p, fill)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def expected_paragraph_texts(fill: dict, circles: dict) -> list[str]:
    """凍結検証の期待列: テンプレ全段落（本文+テーブルセル）へ同一の差し込み+
    sanctioned substitution（許可グループの許可数字のみ丸数字化）を適用。
    それ以外の本文はテンプレ逐語のまま＝完全一致検査で凍結を保証する。"""
    tdoc = Document(TEMPLATE_PATH)
    # fix1[H7C-01]: 期待列の基準は canonical 定数に錨づける——テンプレが
    # canonical pin と一致することを確認してから導出する（コミット物からの
    # 自己生成に正本性を委ねない。バイナリ SHA を追随更新する同時改変は
    # ここで拒否される）
    if canonical_sha256(tdoc) != TEMPLATE_CANONICAL_SHA256:
        raise ShinjutsuIntegrityError("template canonical mismatch")
    # 期待列は「canonical 検証済みテンプレ+承認済み変換（丸数字→差し込みの
    # 同順）」からのみ導出される（sanctioned substitution 以外の差分は
    # 完全一致検査で検出される）
    for group, digit in circles.items():
        _circle_cell(tdoc, _CIRCLE_CELLS[group], digit)
    for p in _walk_paragraphs(tdoc):
        for r in p.runs:
            if "{{" in r.text:
                text = r.text
                for k, v in fill.items():
                    text = text.replace(k, v)
                r.text = text
    return [p.text for p in _walk_paragraphs(tdoc)]


def verify_shinjutsu_docx(docx_bytes: bytes, fill: dict,
                          circles: dict) -> None:
    """生成物の全段落がテンプレ由来の期待列と完全一致・プレースホルダ残存
    なし・丸数字は 1 グループ 1 個のみ（弁護士決定の構造検査）。"""
    got = [p.text for p in _walk_paragraphs(Document(io.BytesIO(docx_bytes)))]
    if got != expected_paragraph_texts(fill, circles):
        raise ShinjutsuIntegrityError("body mismatch against template")
    joined = "\n".join(got)
    if "{{" in joined or "}}" in joined:
        raise ShinjutsuIntegrityError("unfilled placeholder")
    n_circled = sum(joined.count(c) for c in _CIRCLED.values())
    if n_circled != len(circles):
        raise ShinjutsuIntegrityError("circled digit count mismatch")


def generate(record: dict, today: datetime.date | None = None) -> bytes:
    """検証済み申述書 docx を生成（拒否は ShinjutsuRejection・完全性は
    ShinjutsuIntegrityError）。kintone への I/O は受け口側の責務。"""
    verify_template_integrity()
    circles, _reasons = _circles(record)   # 拒否判定は build_fill_data に集約
    fill = build_fill_data(record, today=today)
    docx_bytes = build_shinjutsu_docx(fill, circles)
    verify_shinjutsu_docx(docx_bytes, fill, circles)
    return docx_bytes
