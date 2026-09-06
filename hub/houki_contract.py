"""相続放棄 委任契約書の生成（HOUKI-CONTRACT-GEN・houki 固有部分）

contract_webhook.py（時効版）を ChannelConfig 化した上で、相続放棄（App 40）向けの
差し込み・申述人の束ね・費用計算・署名者決定を本モジュールが担う。状態機械・CAS・
CloudSign 呼出し・凍結検証の枠組みは contract_webhook の共通経路（cfg 経由）。

申述人の束ね（1 レコード=1 申述人・fix1 HCG-01「管理レコード」方式）:
  管理レコード = 被相続人グループID の値が自身のレコード番号と一致するレコード。
  被相続人グループID が空のレコードは自分自身が管理レコード（申述人 1 名）。
  契約書作成/クラウドサイン登録の起点が管理レコードでなければ生成も CloudSign も
  行わず起点を 要確認（not_manager: 通知に代表者のレコード番号）。グループID が数値
  でない（group_invalid）／その番号のレコードが無い（group_missing）も 要確認。
  申述人集合 = 被相続人グループID が管理レコードの番号と一致する全レコード（管理
  レコードを含む）。代表者（署名者・定型文の氏名）= 管理レコード。並び順は管理
  レコード先頭、以下レコード番号昇順。
  規則 4（二重防御）: 集合内の他レコードに 契約書ステータス 非空または
  cloudsign_document_id 非空があれば 要確認（other_member_has_contract）。
  排他の本体は管理レコードの $revision CAS（契約書作成/登録の経路に入れるのは
  管理レコードのみ＝グループ内の CAS 対象は 1 レコードに収束）。
  集合は送信直前にも「ID 指定で再取得した管理レコード」から再構成し、指紋
  （JSON 直列化 → sha256・fix1 HCG-03）が変化していれば 要確認（HCG-02）。

費用（弁護士凍結事項・定数は本モジュールの 1 か所のみ・テストで pin）:
  報酬合計 = FEE_BASE + FEE_ADDITIONAL × (n − 1)
  実費合計 = DEPOSIT_PER_APPLICANT × n
  送付先数 = 全員の 債権者一覧 の正規化後和集合の件数（空行は数えない）
  追加送付件数 = max(0, 送付先数 − INCLUDED_DESTINATIONS)
  追加送付料合計 = 追加送付件数 × EXTRA_SEND_FEE
  支払総額 = 報酬合計 + 追加送付料合計 + 実費合計

特約と契約署名:
  全員: 特約 = 欄そのまま（空なら「特になし」）・署名者 = 全員のメールアドレス
  代表者のみ／空（初期値扱い・空は通知に明記）: 申述人 2 名以上なら
  REPRESENTATIVE_CLAUSE（代表者氏名を差し込み）+ 改行 + 欄（空なら定型文のみ）・
  1 名なら定型文なし（欄が空なら「特になし」）。署名者 = 代表者のみ。
  TOKUYAKU_MAX_CHARS の判定は欄（人が書いた部分）のみ（contract_webhook の
  tokuyaku_problem が担う・定型文は含めない）。

書き込み: 起点レコードのみ（委任契約書・契約書ステータス・cloudsign_document_id）。
他の申述人レコードは読むだけ。値（氏名・住所・メール）はログ・通知に載せない。
"""

import copy
import datetime
import hashlib
import io
import json
import re
import unicodedata
from dataclasses import dataclass
from zoneinfo import ZoneInfo

from docx import Document
from docx.text.paragraph import Paragraph

from hub import kintone
from hub.docx_builder import fill_template
from hub.houki_case_store import APP_HOUKI_CASE, CREDITOR_TABLE

# ── 雛形（人承認済み現物・HOUKI-CONTRACT-FINAL-DOCX の確定 v1） ──────────────
TEMPLATE_PATH = "docx_templates/houki/委任契約書.docx"
TEMPLATE_SHA256 = "519d4bb6765a266514f808e37c8d5f40533d752a694f5989276c3ac4e68d85c9"
# 雛形の全段落テキスト（差し込み記号を含む 62 段落・"\n" 連結）の sha256。
# 雛形の文言が 1 字でも変わればテストが落ちる（凍結条項の段落単位 pin）
TEMPLATE_PARAGRAPHS_SHA256 = (
    "9f15e22b915ed1f845880e76d844d6af0ea339fcfbbc7149c50c49826d25b2c3")
TEMPLATE_PARAGRAPH_COUNT = 62
OUTPUT_FILENAME = "委任契約書_相続放棄.docx"
OUTPUT_PDF_NAME = "委任契約書_相続放棄.pdf"

# 生成物の実行時凍結検証（第2条 見出し+1 項＝報酬条項・雛形逐語）
FROZEN_CLAUSE = (
    "第2条（費用および支払方法）",
    "1　本件の弁護士報酬は、申述人1名の場合は金88,000円とし、2名目以降は1名につき"
    "金33,000円を加算する。受理通知書の写しの送付は送付先3か所までを含み、4か所目以降は"
    "送付先1か所につき金1,100円を加算する。送付先の数は通知を送る相手先ごとに数え、"
    "申述人が複数であっても同じ相手先は1か所と数える。いずれも消費税込みとする。",
)

PLACEHOLDERS = (
    "{{被相続人氏名}}", "{{申述人数}}", "{{報酬合計}}", "{{追加送付件数}}",
    "{{追加送付料合計}}", "{{実費合計}}", "{{支払総額}}", "{{特約}}",
    "{{契約年}}", "{{契約月}}", "{{契約日}}", "{{申述人一覧}}",
)
APPLICANTS_KEY = "{{申述人一覧}}"
TOKUYAKU_KEY = "{{特約}}"

# ── 弁護士凍結事項（費用・定型文） ───────────────────────────────────────────
FEE_BASE = 88_000                 # 申述人 1 名目（税込）
FEE_ADDITIONAL = 33_000           # 2 名目以降 1 名につき（税込）
DEPOSIT_PER_APPLICANT = 5_000     # 実費預託金 1 名につき
EXTRA_SEND_FEE = 1_100            # 受理通知書の追加送付 1 か所につき（税込）
INCLUDED_DESTINATIONS = 3         # 報酬に含む送付先の数
REPRESENTATIVE_CLAUSE = (
    "甲らは、申述人{name}を代表者と定め、代表者が甲ら全員のために本契約に電子署名する。")
TOKUYAKU_NONE = "特になし"

# ── App 40 欄 ────────────────────────────────────────────────────────────────
FIELD_GROUP = "被相続人グループID"
FIELD_CONTRACT_STATUS = "契約書ステータス"
FIELD_CS_DOC_ID = "cloudsign_document_id"
# fix2 HCGF1-02: CloudSign 下書き作成の呼出し後に document ID を確定できなかった
# （結果不明）ことを cloudsign_document_id に永続的に残す印。非空の欄は規則 4
# （member_state_problems）と _regeneration_guard（cs_registered）の両方が遮断する
# ので、印がある限りグループID の付け替えでも本人でも再登録できない。解除は人の
# 操作のみ（MANUAL_RESOLUTION_TEXT の手順）。既に非空なら上書きしない
CLOUDSIGN_RESULT_UNKNOWN_MARK = "結果不明:要手動確認"
# fix5 HCGF4-01: 回収情報はレコード自身（契約書回収メモ・MULTI_LINE_TEXT・App 40 実測
# 2026-09-06）に 1 行ずつ追記して持つ（既存本文を保全）。LINE 通知には document ID を
# 載せない（redaction 規律: external_ref は全 sink 抑止）。指紋・docx 差し込みの対象外
FIELD_RECOVERY_MEMO = "契約書回収メモ"
MANUAL_RESOLUTION_TEXT = (
    "cloudsign_document_id に「結果不明:要手動確認」を記録しました。CloudSign 画面で"
    "下書きの有無を確認し、(i) 下書きがあれば本物の document ID を欄に入れて契約書"
    "ステータスを「クラウドサイン登録済」に、(ii) 無ければ欄を空にして再度"
    "「クラウドサイン登録」に設定してください。")
# 管理レコード判定の分類（閉集合・通知の理由と JSON の skip に使う）
NOT_MANAGER = "not_manager"
GROUP_INVALID = "group_invalid"
GROUP_MISSING = "group_missing"
OTHER_MEMBER_HAS_CONTRACT = "other_member_has_contract"
FIELD_SIGN_MODE = "契約署名"
SIGN_ALL = "全員"
SIGN_REPRESENTATIVE = "代表者のみ"
FIELD_NAME = "顧客名"
FIELD_ADDR = "住所"
FIELD_DECEDENT = "被相続人氏名"
FIELD_EMAIL = "メールアドレス"
FIELD_TOKUYAKU = "特約"
_EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
JST = ZoneInfo("Asia/Tokyo")


class HoukiContractIntegrityError(RuntimeError):
    """雛形/生成物の凍結検証に失敗（生成しない）。"""


def _fv(record: dict, code: str) -> str:
    return str((record.get(code) or {}).get("value") or "").strip()


def _rid(record: dict) -> str:
    return _fv(record, "$id")


# ── 費用 ─────────────────────────────────────────────────────────────────────
def normalize_creditor(name) -> str:
    """送付先名の正規化: 前後空白除去・NFKC（全角/半角の英数記号統一）・
    全角/半角スペース除去。空文字は「空行」= 数えない。"""
    s = unicodedata.normalize("NFKC", str(name or ""))
    return re.sub(r"[\s　]+", "", s).strip()


def destinations(records: list[dict]) -> set[str]:
    """申述人集合全員の 債権者一覧 の正規化後和集合。"""
    out: set[str] = set()
    for rec in records:
        rows = ((rec.get(CREDITOR_TABLE) or {}).get("value") or [])
        for row in rows:
            raw = ((row.get("value") or {}).get("債権者名") or {}).get("value")
            key = normalize_creditor(raw)
            if key:
                out.add(key)
    return out


def compute_fees(n: int, destination_count: int) -> dict:
    """税込・整数円。キーは差し込み名と同じ語。"""
    extra_count = max(0, destination_count - INCLUDED_DESTINATIONS)
    fee = FEE_BASE + FEE_ADDITIONAL * (n - 1)
    deposit = DEPOSIT_PER_APPLICANT * n
    extra_fee = extra_count * EXTRA_SEND_FEE
    return {
        "申述人数": n, "送付先数": destination_count,
        "報酬合計": fee, "実費合計": deposit,
        "追加送付件数": extra_count, "追加送付料合計": extra_fee,
        "支払総額": fee + extra_fee + deposit,
    }


def fmt_yen(n: int) -> str:
    return f"{int(n):,}"


# ── 申述人集合 ────────────────────────────────────────────────────────────────
def _escape(value: str) -> str:
    return value.replace("\\", "\\\\").replace('"', '\\"')


def manager_check(record: dict) -> tuple[str | None, str]:
    """起点レコードが管理レコードか。戻り値 (分類 or None=管理レコード, グループID)。
    空=自分自身が管理レコード。数値でない=group_invalid。自番号と不一致=not_manager
    （存在確認は gather_applicants が行い group_missing に振り分ける）。"""
    group = _fv(record, FIELD_GROUP)
    if not group:
        return None, ""
    if not group.isdigit():
        return GROUP_INVALID, group
    if group != _rid(record):
        return NOT_MANAGER, group
    return None, group


async def gather_applicants(record: dict) -> tuple[list[dict], str, list[str]]:
    """管理レコードを代表者として申述人集合を返す（管理レコード → レコード番号昇順）。
    戻り値 (集合, グループID, 問題分類の一覧)。問題があれば集合は [起点] のまま。"""
    kind, group = manager_check(record)
    if kind == GROUP_INVALID:
        return [record], group, [GROUP_INVALID]
    if kind == NOT_MANAGER:
        try:
            await kintone.get_record(APP_HOUKI_CASE, group)
        except kintone.KintoneError:
            return [record], group, [GROUP_MISSING]
        return [record], group, [NOT_MANAGER]
    if not group:
        return [record], "", []
    rows = await kintone.search_records(
        APP_HOUKI_CASE, f'{FIELD_GROUP} = "{_escape(group)}" order by $id asc limit 500')
    rep_id = _rid(record)
    others = sorted((r for r in rows if _rid(r) != rep_id), key=lambda r: int(_rid(r) or 0))
    return [record] + others, group, []


# 規則 4 の対象状態（票の閉集合から 要確認 を除く。理由: 非管理レコードは規則 1 で
# 要確認 になるため、含めると「B→A」の順で管理レコードが永久に生成できない。
# トリガ値〔契約書作成/クラウドサイン登録〕は人が誤って押した B の webhook 待ち状態
# なので残存状態ではない）
MEMBER_BLOCKING_STATUSES = frozenset({"契約書作成中", "契約書作成済",
                                      "クラウドサイン登録中", "クラウドサイン登録済"})


def member_state_problems(records: list[dict]) -> list[str]:
    """規則 4（二重防御）: 管理レコード以外に 契約書ステータス ∈ MEMBER_BLOCKING_STATUSES
    または cloudsign_document_id 非空があれば other_member_has_contract（レコード番号のみ）。
    グループID の付け替えで残った契約状態を検知する。"""
    bad = [_rid(r) for r in records[1:]
           if _fv(r, FIELD_CONTRACT_STATUS) in MEMBER_BLOCKING_STATUSES
           or _fv(r, FIELD_CS_DOC_ID)]
    if bad:
        return [f"{OTHER_MEMBER_HAS_CONTRACT}（レコード番号 {', '.join(bad)}）"]
    return []


def manager_problem_text(kind: str, group: str) -> str:
    if kind == NOT_MANAGER:
        return (f"{NOT_MANAGER}: 契約書の操作は代表者のレコード（No.{group}）で"
                "行ってください")
    if kind == GROUP_MISSING:
        return (f"{GROUP_MISSING}: 被相続人グループID（No.{group}）のレコードが"
                "App 40 に存在しません")
    return f"{GROUP_INVALID}: 被相続人グループID が数値ではありません"


def applicant_problems(records: list[dict]) -> list[str]:
    """要確認理由（固定語彙+レコード番号のみ・値は載せない）。"""
    problems: list[str] = []
    decedents = {_fv(r, FIELD_DECEDENT) for r in records}
    if len(decedents) != 1 or "" in decedents:
        problems.append("被相続人氏名が申述人集合で一致しない（または空）")
    for r in records:
        if not _fv(r, FIELD_NAME):
            problems.append(f"顧客名 未入力（レコード番号 {_rid(r)}）")
        if not _fv(r, FIELD_ADDR):
            problems.append(f"住所 未入力（レコード番号 {_rid(r)}）")
    return problems


def sign_mode(record: dict) -> tuple[str, bool]:
    """(署名方式, 欄が空だったか)。空は 代表者のみ として扱う。"""
    raw = _fv(record, FIELD_SIGN_MODE)
    if raw == SIGN_ALL:
        return SIGN_ALL, False
    return SIGN_REPRESENTATIVE, raw == ""


def signers(records: list[dict], mode: str) -> list[dict]:
    return list(records) if mode == SIGN_ALL else [records[0]]


def participants(records: list[dict], mode: str) -> tuple[list[tuple[str, str]], list[str]]:
    """CloudSign 宛先 [(email, 顧客名)] と要確認理由。"""
    out: list[tuple[str, str]] = []
    problems: list[str] = []
    for r in signers(records, mode):
        email = _fv(r, FIELD_EMAIL)
        if not email:
            problems.append(f"メールアドレス 未入力（レコード番号 {_rid(r)}）")
        elif not _EMAIL_RE.fullmatch(email):
            problems.append(f"メールアドレス 形式不正（レコード番号 {_rid(r)}）")
        else:
            out.append((email, _fv(r, FIELD_NAME)))
    return out, problems


def compose_tokuyaku(records: list[dict], mode: str) -> str:
    """特約本文（雛形の {{特約}} へ apply_tokuyaku で差し込む文字列・改行=段落）。"""
    rep = records[0]
    user = str((rep.get(FIELD_TOKUYAKU) or {}).get("value") or "").strip()
    if mode == SIGN_ALL or len(records) < 2:
        return user or TOKUYAKU_NONE
    clause = REPRESENTATIVE_CLAUSE.format(name=_fv(rep, FIELD_NAME))
    return clause + ("\n" + user if user else "")


# ── 差し込み ──────────────────────────────────────────────────────────────────
def build_fill_data(records: list[dict], fees: dict,
                    now: datetime.datetime | None = None) -> dict:
    """fill_template 用 10 キー（{{特約}} と {{申述人一覧}} は専用処理のため含まない）。"""
    now = now or datetime.datetime.now(JST)
    rep = records[0]
    return {
        "{{被相続人氏名}}": _fv(rep, FIELD_DECEDENT),
        "{{申述人数}}": str(fees["申述人数"]),
        "{{報酬合計}}": fmt_yen(fees["報酬合計"]),
        "{{追加送付件数}}": str(fees["追加送付件数"]),
        "{{追加送付料合計}}": fmt_yen(fees["追加送付料合計"]),
        "{{実費合計}}": fmt_yen(fees["実費合計"]),
        "{{支払総額}}": fmt_yen(fees["支払総額"]),
        "{{契約年}}": str(now.year), "{{契約月}}": str(now.month), "{{契約日}}": str(now.day),
    }


def _set_text(para, text: str) -> None:
    para.runs[0].text = text
    for r in para.runs[1:]:
        r.text = ""


def expand_applicants(docx_bytes: bytes, records: list[dict]) -> bytes:
    """{{申述人一覧}} の段落を申述人ごとに「住所　…」「氏名　…」の 2 段落へ展開
    （段落書式は元段落を複製・2 名以上は人と人の間に空段落 1 つ・表は使わない）。"""
    doc = Document(io.BytesIO(docx_bytes))
    anchor = next((p for p in doc.paragraphs if p.text == APPLICANTS_KEY), None)
    if anchor is None:
        raise HoukiContractIntegrityError("applicants placeholder missing")
    template_p = copy.deepcopy(anchor._p)
    lines: list[str] = []
    for i, rec in enumerate(records):
        if i:
            lines.append("")
        lines.append(f"住所　{_fv(rec, FIELD_ADDR)}")
        lines.append(f"氏名　{_fv(rec, FIELD_NAME)}")
    _set_text(anchor, lines[0])
    prev = anchor._p
    for line in lines[1:]:
        new_p = copy.deepcopy(template_p)
        prev.addnext(new_p)
        _set_text(Paragraph(new_p, anchor._parent), line)
        prev = new_p
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def recovery_memo_line(existing_id: str, doc_id: str | None,
                       now: datetime.datetime | None = None) -> str:
    """回収メモの 1 行。ok（D あり）=二重下書きの疑い、unknown=結果不明。"""
    stamp = (now or datetime.datetime.now(JST)).strftime("%Y-%m-%d %H:%M")
    if doc_id:
        return f"{stamp} 二重下書きの疑い: 既存 {existing_id} / 今回 {doc_id}"
    return f"{stamp} 結果不明: 既存 {existing_id} のまま"


def append_memo(existing_text: str, line: str) -> str:
    """既存本文を保全して末尾に 1 行追加。"""
    base = str(existing_text or "").rstrip("\n")
    return (base + "\n" + line) if base else line


def verify_template_integrity() -> None:
    data = open(TEMPLATE_PATH, "rb").read()
    if hashlib.sha256(data).hexdigest() != TEMPLATE_SHA256:
        raise HoukiContractIntegrityError("template hash mismatch")


def template_paragraphs_sha256(path: str = TEMPLATE_PATH) -> str:
    texts = [p.text for p in Document(path).paragraphs]
    return hashlib.sha256("\n".join(texts).encode("utf-8")).hexdigest()


def render(records: list[dict], mode: str, fees: dict,
           now: datetime.datetime | None = None) -> bytes:
    """雛形の整合検証 → 10 キー差し込み → 特約（定型文込み）→ 申述人一覧の展開 →
    凍結条項検証（第2条）。docx 添付と CloudSign 用 PDF の両方がこの単一経路。"""
    from contract_webhook import apply_tokuyaku, verify_frozen_clause
    verify_template_integrity()
    docx_bytes = fill_template(TEMPLATE_PATH, build_fill_data(records, fees, now))
    docx_bytes = apply_tokuyaku(docx_bytes, compose_tokuyaku(records, mode))
    docx_bytes = expand_applicants(docx_bytes, records)
    verify_frozen_clause(docx_bytes, FROZEN_CLAUSE)
    return docx_bytes


# ── 準備（contract_webhook の共通経路が呼ぶ） ───────────────────────────────
@dataclass
class HoukiPlan:
    records: list[dict]
    group: str
    mode: str
    mode_blank: bool
    fees: dict
    applicant_problems: list[str]
    email_problems: list[str]
    participants: list[tuple[str, str]]
    skip: str = ""          # 管理レコード規則の分類（空=該当なし）

    def fingerprint_source(self) -> dict:
        """fix1 HCG-02/03: 指紋の対象（構造のまま）。申述人集合の全レコード
        （番号・顧客名・住所・メール・被相続人氏名・債権者一覧の正規化配列）、
        管理レコードの 契約署名・特約、費用 5 値、署名者メール配列。"""
        rep = self.records[0]
        return {
            "group": self.group,
            "mode": self.mode,
            "members": [{
                "id": _rid(r), "name": _fv(r, FIELD_NAME), "addr": _fv(r, FIELD_ADDR),
                "email": _fv(r, FIELD_EMAIL), "decedent": _fv(r, FIELD_DECEDENT),
                "creditors": [normalize_creditor(
                    ((row.get("value") or {}).get("債権者名") or {}).get("value"))
                    for row in ((r.get(CREDITOR_TABLE) or {}).get("value") or [])],
            } for r in self.records],
            "manager": {"sign": _fv(rep, FIELD_SIGN_MODE),
                        "tokuyaku": str((rep.get(FIELD_TOKUYAKU) or {}).get("value") or "")},
            "fees": {k: self.fees[k] for k in ("報酬合計", "実費合計", "追加送付件数",
                                                "追加送付料合計", "支払総額")},
            "signers": [email for email, _n in self.participants],
        }

    @property
    def fingerprint(self) -> str:
        """JSON 直列化（sort_keys・区切り固定）→ sha256。配列境界・項目境界を保持する
        （文字列連結の衝突=HCG-03 を排除）。"""
        raw = json.dumps(self.fingerprint_source(), sort_keys=True, ensure_ascii=False,
                         separators=(",", ":"))
        return hashlib.sha256(raw.encode("utf-8")).hexdigest()

    def summary(self) -> str:
        f = self.fees
        mode = "全員" if self.mode == SIGN_ALL else "代表者のみ"
        if self.mode_blank:
            mode += "（契約署名 が空のため代表者のみとして作成）"
        lines = [f"・申述人数: {f['申述人数']}名",
                 f"・送付先数: {f['送付先数']}か所（追加送付 {f['追加送付件数']}か所）",
                 f"・支払総額: {fmt_yen(f['支払総額'])}円",
                 f"・署名方式: {mode}"]
        if not self.group:
            lines.append("・被相続人グループID 空・1 名として作成")
        return "\n".join(lines)


async def plan(record: dict) -> HoukiPlan:
    records, group, manager_problems = await gather_applicants(record)
    mode, blank = sign_mode(record)
    fees = compute_fees(len(records), len(destinations(records)))
    if manager_problems:
        problems = [manager_problem_text(k, group) for k in manager_problems]
    else:
        problems = member_state_problems(records) + applicant_problems(records)
    parts, email_problems = participants(records, mode)
    return HoukiPlan(records, group, mode, blank, fees, problems, email_problems, parts,
                     skip=(manager_problems[0] if manager_problems
                           else (OTHER_MEMBER_HAS_CONTRACT
                                 if member_state_problems(records) else "")))
