"""時効援用通知書テンプレートの収載スクリプト（JIKOU-NOTICE-1）

大野修正版の実物（Desktop/claude/時効援用通知書　黒田達人.docx・誤字・
信用情報文言修正済み=本文凍結の正）から、個人情報をプレースホルダ化した
テンプレートを docx_templates/jikou/時効援用通知書.docx へ生成する。
生成物の SHA-256 を notice_webhook.py の TEMPLATE_SHA256 に pin する
（契約書と同型）。

置換規則（値は**実物の記載ブロックから読み取る**＝本スクリプトに個人情報を
持たない）:
  - p00 日付行（全体）              → {{通知日付}}
  - 通知人氏名（p02/p11/p20 の 3 箇所）→ {{通知人氏名}}
  - p19 ふりがな                    → {{ふりがな}}
  - p21 生年月日                    → {{生年月日}}
  - p22 住所（記載ブロック）        → {{通知人住所}}
  - p22 の直後に「旧　住　所」行を同書式で追加 → {{旧住所}}
    （生成時: 旧住所が空なら段落ごと削除・非空なら差し込み）

凍結（大野裁定・改変しない）:
  - 宛先「債権者各位」
  - 冒頭の事務所住所ブロック（〒・住所・建物名・通知代理人 弁護士名・
    TEL/FAX＝p03〜p08）
  - 表題・本文全段落（信用情報機関への削除依頼・債務承認非該当の付言含む）

run 書式の保持（fix2・実機で発見）:
  元票の「ふりがな」「生年月日」ラベルは 5 文字幅の均等割り付け
  （w:fitText val=1050）が run 単位で付いている。初版は段落全体を先頭 run
  へ潰していたため、この 2 行だけ行全体が 1050 twips に圧縮され小さな崩れた
  字になった。本版は**値の run だけ**をプレースホルダへ置換し、ラベル run
  （fitText 含む）と各 run の rPr をそのまま残す。プレースホルダは必ず単一
  run 内に収まる（notice_webhook 側も run 単位で差し込む）。凍結段落の run は
  触らない。

実行: python make_notice_template.py
"""

import copy
import hashlib
import re
from pathlib import Path

from docx import Document

SOURCE = Path.home() / "Desktop" / "claude" / "時効援用通知書　黒田達人.docx"
OUT = Path("docx_templates") / "jikou" / "時効援用通知書.docx"

# 記載ブロックのラベル（値の抽出と、置換後の体裁維持に使う）
_LABELED = {
    19: ("ふりがな", "{{ふりがな}}"),
    20: ("債務者氏名", "{{通知人氏名}}"),
    21: ("生年月日", "{{生年月日}}"),
    22: ("住　　　所", "{{通知人住所}}"),
}


def _replace_span(p, old: str, new: str) -> None:
    """段落内の部分文字列 old（唯一）を new へ置換し、run 書式を保持する。

    old にかかる最初の run へ「前置き+new」を置き（その run の rPr を継承）、
    末尾 run には後置きだけを残し、中間 run は削除する。old が 1 run に
    収まる場合はその run 内で置換。結果として new は必ず単一 run 内。
    """
    runs = p.runs
    full = "".join(r.text for r in runs)
    if full.count(old) != 1:
        raise SystemExit(f"置換対象が一意でありません: {old!r} in {full!r}")
    start = full.index(old)
    end = start + len(old)
    pos = 0
    first = last = None
    for i, r in enumerate(runs):
        r_start, r_end = pos, pos + len(r.text)
        pos = r_end
        if first is None and r_end > start:
            first = i
            prefix = r.text[:start - r_start]
        if first is not None and r_end >= end:
            last = i
            suffix = r.text[end - r_start:]
            break
    if first == last:
        runs[first].text = prefix + new + suffix
        return
    runs[first].text = prefix + new
    runs[last].text = suffix
    for r in runs[first + 1:last + 1]:        # 中間 run と空になった末尾 run
        if r is not runs[last] or not suffix:
            r._r.getparent().remove(r._r)


def _label_split(text: str, label: str) -> tuple[str, str]:
    """『ラベル+空白+値』を（ラベル+空白, 値）に分ける。"""
    m = re.match(rf"^({re.escape(label)}[\s　]*)(.*)$", text)
    if not m or not m.group(2):
        raise SystemExit(f"記載ブロックの形が想定と異なります: {label}")
    return m.group(1), m.group(2)


def main() -> None:
    if not SOURCE.exists():
        raise SystemExit(f"大野修正版が見つかりません: {SOURCE}")
    doc = Document(SOURCE)
    paras = doc.paragraphs
    if len(paras) != 24 or paras[1].text != "債権者各位":
        raise SystemExit("段落構成が想定（24 段落・宛先=債権者各位）と異なります")

    # 実物から値を読む（氏名は記載ブロックの債務者氏名を正とする）
    values = {}
    for idx, (label, key) in _LABELED.items():
        _prefix, value = _label_split(paras[idx].text, label)
        values[key] = value
    name = values["{{通知人氏名}}"]

    # 記載ブロック 4 行: 値の run だけをプレースホルダ化（ラベル run と
    # 元の空白 run は書式ごとそのまま）
    for idx, (label, key) in _LABELED.items():
        _replace_span(paras[idx], values[key], key)

    # 旧住所行を住所行の直後に同書式で追加（{{旧住所}}・空なら生成時に削除）
    addr_p = paras[22]
    new_el = copy.deepcopy(addr_p._element)
    addr_p._element.addnext(new_el)
    doc2_paras = doc.paragraphs  # 挿入後に取り直し
    old_addr_p = doc2_paras[23]
    _replace_span(old_addr_p, "住　　　所", "旧　住　所")
    _replace_span(old_addr_p, "{{通知人住所}}", "{{旧住所}}")

    # 日付行（全体）と通知人氏名（冒頭・本文）
    _replace_span(doc2_paras[0], doc2_paras[0].text, "{{通知日付}}")
    for idx in (2, 11):
        if name not in doc2_paras[idx].text:
            raise SystemExit(f"p{idx:02d} に氏名が見つかりません")
        _replace_span(doc2_paras[idx], name, "{{通知人氏名}}")

    # プレースホルダは単一 run 内・記載ブロックの値 run に fitText がない
    for p in doc.paragraphs:
        if "{{" in p.text:
            holders = [r for r in p.runs if "{{" in r.text and "}}" in r.text]
            if len(holders) != 1 or holders[0].text.count("{{") != 1:
                raise SystemExit(f"プレースホルダが単一 run に収まっていません: {p.text}")
            rpr = holders[0]._r.rPr
            if rpr is not None and any(c.tag.endswith("}fitText") for c in rpr):
                raise SystemExit(f"値 run に fitText が混入しています: {p.text}")

    # 残存個人情報の機械検査（氏名・ふりがな・生年月日・記載住所の値）
    all_text = "\n".join(p.text for p in doc.paragraphs)
    for key, value in values.items():
        for token in re.split(r"[\s　]+", value):
            if token and token in all_text:
                raise SystemExit(f"個人情報が残存しています: {key}")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    sha = hashlib.sha256(OUT.read_bytes()).hexdigest()
    print(f"saved: {OUT} paras={len(doc.paragraphs)}")
    print(f"TEMPLATE_SHA256 = {sha}")


if __name__ == "__main__":
    main()
