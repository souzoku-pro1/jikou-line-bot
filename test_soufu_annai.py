"""channels/soufu_annai.py（M4 送付案内・T2-1）のテスト

T2-1 完了条件: ソート・フィルタ（無効除外/ユニット）・未定義キー・
同期検査異常系・複数行差込。
"""

import asyncio
import io
import os
import re
import unittest
import zipfile
from unittest.mock import AsyncMock, patch

os.environ.setdefault("KINTONE_SUBDOMAIN", "testsub")

from docx import Document

from channels import soufu_annai
from channels.soufu_annai import (
    SoufuAnnaiError,
    build_enclosure_text,
    build_soufu_annai_docx,
    check_block_sync,
    fetch_blocks,
)
from hub import kintone
from hub.docx_builder import fill_template_multiline

_OFFICE_ENV = {
    "OFFICE_NAME": "大野法律事務所",
    "OFFICE_ZIP": "332-0000",
    "OFFICE_ADDRESS": "埼玉県川口市テスト町1-2-3",
    "OFFICE_TEL": "048-000-0000",
    "OFFICE_FAX": "048-000-0001",
    "OFFICE_ATTORNEY": "大野　太郎",
}


def _all_text(docx_bytes: bytes) -> str:
    """段落＋表セルの全文（新書式は書類表を含むため表も走査）"""
    doc = Document(io.BytesIO(docx_bytes))
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                parts.append(c.text)
    return "\n".join(parts)


def block(key, order="1", units=("時効援用",), note="", name=None):
    return {
        "$id": {"value": "1"},
        "ブロックキー": {"value": key},
        "表示名": {"value": name or key},
        "案内文": {"value": note},
        "対象ユニット": {"value": list(units)},
        "返送要否": {"value": "不要"},
        "表示順": {"value": order},
    }


def run(coro):
    return asyncio.run(coro)


class TestFetchBlocks(unittest.TestCase):
    def test_sorted_by_display_order(self):
        records = [block("B", order="20"), block("A", order="10"), block("C", order="5")]
        with patch("hub.kintone.search_records", new=AsyncMock(return_value=records)) as s:
            got = run(fetch_blocks("時効援用", ["A", "B", "C"]))
        keys = [b["ブロックキー"]["value"] for b in got]
        self.assertEqual(keys, ["C", "A", "B"])
        query = s.await_args.args[1]
        self.assertIn('有効 in ("yes")', query, "無効ブロックはクエリで除外")
        self.assertIn("order by 表示順", query)

    def test_only_selected_keys_are_returned(self):
        records = [block("A"), block("B"), block("C")]
        with patch("hub.kintone.search_records", new=AsyncMock(return_value=records)):
            got = run(fetch_blocks("時効援用", ["B"]))
        self.assertEqual([b["ブロックキー"]["value"] for b in got], ["B"])

    def test_undefined_key_raises(self):
        """未定義キー（マスタに無い/無効化済み）の選択はエラー（→エラー遷移+警報）"""
        with patch("hub.kintone.search_records", new=AsyncMock(return_value=[block("A")])):
            with self.assertRaises(SoufuAnnaiError) as ctx:
                run(fetch_blocks("時効援用", ["A", "存在しないキー"]))
        self.assertIn("存在しないキー", str(ctx.exception))

    def test_unit_mismatch_raises(self):
        records = [block("A", units=("相続放棄",))]
        with patch("hub.kintone.search_records", new=AsyncMock(return_value=records)):
            with self.assertRaises(SoufuAnnaiError) as ctx:
                run(fetch_blocks("時効援用", ["A"]))
        self.assertIn("時効援用", str(ctx.exception))

    def test_empty_selection_raises(self):
        with self.assertRaises(SoufuAnnaiError):
            run(fetch_blocks("時効援用", []))


class TestEnclosureText(unittest.TestCase):
    def test_lines_with_and_without_note(self):
        blocks = [block("委任契約書2通", note="ご署名のうえ1通ご返送ください"),
                  block("返信用封筒")]
        text = build_enclosure_text(blocks)
        self.assertEqual(text.split("\n"), [
            "■ 委任契約書2通",
            "　ご署名のうえ1通ご返送ください",
            "■ 返信用封筒",
        ])


class TestMultilineFill(unittest.TestCase):
    """fill_template_multiline: 改行入り値が Word の改行(<w:br/>)になる（複数行差込）"""

    def _docx_xml(self, docx_bytes):
        with zipfile.ZipFile(io.BytesIO(docx_bytes)) as z:
            return z.read("word/document.xml").decode("utf-8")

    def test_multiline_value_becomes_breaks(self):
        import tempfile
        from pathlib import Path
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "t.docx"
            d = Document()
            d.add_paragraph("{{一覧}}")
            d.save(str(path))
            out = fill_template_multiline(str(path), {"{{一覧}}": "1行目\n2行目\n3行目"})
        xml = self._docx_xml(out)
        self.assertEqual(len(re.findall(r"<w:br\s*/>", xml)), 2, "3行→改行2つ")
        for line in ("1行目", "2行目", "3行目"):
            self.assertIn(line, xml)

    def test_single_line_behaves_like_fill_template(self):
        import tempfile
        from pathlib import Path
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "t.docx"
            d = Document()
            d.add_paragraph("{{名前}}様")
            d.save(str(path))
            out = fill_template_multiline(str(path), {"{{名前}}": "山田"})
        doc = Document(io.BytesIO(out))
        self.assertIn("山田様", doc.paragraphs[0].text)


def make_shipping_record(**over):
    rec = {
        "$id": {"value": "9"},
        "ユニット種別": {"value": "時効援用"},
        "宛先名": {"value": "山田太郎"},
        "宛先郵便番号": {"value": "332-0001"},
        "宛先住所": {"value": "埼玉県川口市朝日1-2-3"},
        "顧客名表示用": {"value": "山田太郎"},
        "件名": {"value": "委任契約書の送付"},
        "本文_特記事項": {"value": "同封の返信用封筒をご利用ください。"},
        "同封物選択": {"value": ["委任契約書2通", "返信用封筒"]},
    }
    rec.update(over)
    return rec


class TestBuildDocx(unittest.TestCase):
    def _run(self, record):
        records = [block("委任契約書2通", order="1", note="ご返送ください"),
                   block("返信用封筒", order="2")]
        with patch.dict("os.environ", _OFFICE_ENV), \
             patch("hub.kintone.search_records", new=AsyncMock(return_value=records)):
            return run(build_soufu_annai_docx(record))

    def test_generates_docx_with_all_content(self):
        out = self._run(make_shipping_record())
        self.assertTrue(out.startswith(b"PK"))
        text = _all_text(out)
        for expected in ("山田太郎",                      # 宛先
                         "書　類　送　付",     # 事務所書式の表題
                         "拝啓　時下ますますご清祥",   # 事務所書式の挨拶文
                         "大野法律事務所", "弁護士　大野　太郎",
                         "ＴＥＬ：048-000-0000", "ＦＡＸ：048-000-0001",
                         "令和",
                         "委任契約書2通",                  # 書類表の書類名
                         "ご返送ください",                  # 書類表の備考（案内文）
                         "同封の返信用封筒"):               # 特記事項（人の記入値）
            self.assertIn(expected, text)
        self.assertNotIn("{{", text, "未置換プレースホルダが残っていない")

    def test_table_rows_match_block_count(self):
        """レイアウト崩れがないこと: 行数が同封物の件数に追従（ヘッダ+件数）"""
        for n in (1, 2, 5):
            records = [block(f"書類{i}", order=str(i)) for i in range(n)]
            rec = make_shipping_record(
                同封物選択={"value": [f"書類{i}" for i in range(n)]},
                本文_特記事項={"value": "x"})
            with patch.dict("os.environ", _OFFICE_ENV), \
                 patch("hub.kintone.search_records", new=AsyncMock(return_value=records)):
                out = run(build_soufu_annai_docx(rec))
            doc = Document(io.BytesIO(out))
            table = doc.tables[0]
            self.assertEqual(len(table.rows), 1 + n, f"n={n}: ヘッダ+{n}行")
            self.assertEqual(len(table.columns), 4, "4列（No./書類名/部数/備考）を維持")
            # No. 列の連番と部数既定値
            self.assertEqual(table.rows[1].cells[0].text, "1")
            self.assertEqual(table.rows[1].cells[2].text, "1", "部数列が無い場合の既定は1")

    def test_remarks_column_reflects_return_flag(self):
        """備考列: 案内文なし+返送要否=要 → 定型の返送依頼文言"""
        b1 = block("要返送書類", note="")
        b1["返送要否"] = {"value": "要"}
        rec = make_shipping_record(同封物選択={"value": ["要返送書類"]},
                                   本文_特記事項={"value": "x"})
        with patch.dict("os.environ", _OFFICE_ENV), \
             patch("hub.kintone.search_records", new=AsyncMock(return_value=[b1])):
            out = run(build_soufu_annai_docx(rec))
        doc = Document(io.BytesIO(out))
        self.assertIn("ご返送をお願いいたします", doc.tables[0].rows[1].cells[3].text)

    def test_customer_line_omitted_when_empty_or_same(self):
        """顧客名表示用が空・宛先と同一のときは「（ご依頼者：…）」を印字しない"""
        for customer in ("", "山田太郎"):  # 空 / 宛先と同一
            rec = make_shipping_record(顧客名表示用={"value": customer})
            out = self._run(rec)
            self.assertNotIn("ご依頼者", _all_text(out), f"customer={customer!r}")

    def test_customer_line_printed_when_different(self):
        rec = make_shipping_record(宛先名={"value": "アコム株式会社"},
                                   顧客名表示用={"value": "山田太郎"})
        out = self._run(rec)
        self.assertIn("（ご依頼者：山田太郎　様）", _all_text(out))

    def test_recipient_block_is_three_lines(self):
        """宛先ブロック: 〒（1行目）／住所（2行目）／宛名+様（3行目）の縦積み"""
        out = self._run(make_shipping_record())
        doc = Document(io.BytesIO(out))
        addr_para = next(p for p in doc.paragraphs if "〒332-0001" in p.text)
        self.assertIn("<w:br", addr_para._p.xml,
                      "〒と住所の間に Word の改行がある（横並びでない）")
        self.assertIn("埼玉県川口市朝日1-2-3", addr_para.text)
        self.assertNotIn("〒332-0001　埼玉県", addr_para.text.replace("\n", ""),
                         "旧レイアウト（全角スペースの横並び）になっていない")
        name_para = next(p for p in doc.paragraphs if "山田太郎" in p.text and "様" in p.text)
        self.assertIsNot(addr_para, name_para, "宛名は別段落（3行目）")

    def test_recipient_block_without_zip(self):
        """郵便番号が空なら 〒 の行を出さない（住所から始まる）"""
        out = self._run(make_shipping_record(宛先郵便番号={"value": ""}))
        doc = Document(io.BytesIO(out))
        text = "\n".join(p.text for p in doc.paragraphs)
        self.assertNotIn("〒\n", text)
        addr_para = next(p for p in doc.paragraphs if "埼玉県川口市朝日1-2-3" in p.text)
        self.assertNotIn("〒", addr_para.text)

    def test_missing_office_info_raises(self):
        empty = {"OFFICE_NAME": "", "OFFICE_ADDRESS": ""}
        records = [block("委任契約書2通"), block("返信用封筒")]
        with patch.dict("os.environ", {**_OFFICE_ENV, **empty}), \
             patch("hub.kintone.search_records", new=AsyncMock(return_value=records)):
            with self.assertRaises(SoufuAnnaiError) as ctx:
                run(build_soufu_annai_docx(make_shipping_record()))
        self.assertIn("OFFICE_NAME", str(ctx.exception))

    def test_missing_unit_raises(self):
        with self.assertRaises(SoufuAnnaiError):
            run(build_soufu_annai_docx(make_shipping_record(ユニット種別={"value": ""})))

    def test_adapter_prepare_returns_artifacts(self):
        records = [block("委任契約書2通"), block("返信用封筒")]
        with patch.dict("os.environ", _OFFICE_ENV), \
             patch("hub.kintone.search_records", new=AsyncMock(return_value=records)):
            result = run(soufu_annai.SoufuAnnaiAdapter().prepare(make_shipping_record()))
        self.assertEqual(len(result.artifacts), 2, "T2-2: docx + 宛名ラベルPDF")
        self.assertEqual(result.artifacts[0].filename, "送付案内.docx")
        self.assertTrue(result.artifacts[0].content.startswith(b"PK"))
        self.assertEqual(result.artifacts[1].filename, "宛名ラベル.pdf")
        self.assertTrue(result.artifacts[1].content.startswith(b"%PDF"))

    def test_adapter_is_registered(self):
        """T2-2: CHANNEL_REGISTRY に登録済み（ディスパッチャから呼ばれる）"""
        import channels
        adapter = channels.get_adapter("送付案内")
        self.assertIsNotNone(adapter)
        self.assertIsInstance(adapter, soufu_annai.SoufuAnnaiAdapter)


class _FakeToolResponse:
    """create_message_with_fallback の戻り（tool_use ブロック）の代役"""

    def __init__(self, note):
        class B:
            type = "tool_use"
            input = {"note": note}
        self.content = [B()]


class TestTokkiNote(unittest.TestCase):
    """AI 特記事項（T2-2）: 加飾であり必須依存にしない"""

    def _blocks(self):
        return [block("委任契約書2通", note="ご返送ください"), block("返信用封筒")]

    def test_generated_note_is_used_and_written_back(self):
        with patch.dict("os.environ", _OFFICE_ENV), \
             patch("hub.kintone.search_records",
                   new=AsyncMock(return_value=self._blocks())), \
             patch("channels.soufu_annai.create_message_with_fallback",
                   new=AsyncMock(return_value=_FakeToolResponse("返信用封筒にてご返送ください。"))):
            result = run(soufu_annai.SoufuAnnaiAdapter().prepare(
                make_shipping_record(本文_特記事項={"value": ""})))
        self.assertEqual(result.fields["本文_特記事項"], "返信用封筒にてご返送ください。")
        doc = Document(io.BytesIO(result.artifacts[0].content))
        text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("返信用封筒にてご返送ください。", text)

    def test_ai_failure_continues_with_blank(self):
        """██ AI 生成失敗 → 空欄で prepare 続行（警報なし・成果物は生成される）██"""
        with patch.dict("os.environ", _OFFICE_ENV), \
             patch("hub.kintone.search_records",
                   new=AsyncMock(return_value=self._blocks())), \
             patch("channels.soufu_annai.create_message_with_fallback",
                   new=AsyncMock(side_effect=RuntimeError("api down"))), \
             patch("hub.notify.notify_admin_line", new=AsyncMock()) as alert:
            result = run(soufu_annai.SoufuAnnaiAdapter().prepare(
                make_shipping_record(本文_特記事項={"value": ""})))
        self.assertEqual(len(result.artifacts), 2, "docx+ラベルは生成される")
        self.assertNotIn("本文_特記事項", result.fields, "空欄のまま（書き戻さない）")
        alert.assert_not_awaited()

    def test_human_note_is_not_overwritten(self):
        """人が書いた特記事項があるときは AI を呼ばない"""
        ai = AsyncMock(return_value=_FakeToolResponse("AIの文"))
        with patch.dict("os.environ", _OFFICE_ENV), \
             patch("hub.kintone.search_records",
                   new=AsyncMock(return_value=self._blocks())), \
             patch("channels.soufu_annai.create_message_with_fallback", new=ai):
            result = run(soufu_annai.SoufuAnnaiAdapter().prepare(
                make_shipping_record(本文_特記事項={"value": "人間が書いた注意書き"})))
        ai.assert_not_awaited()
        self.assertNotIn("本文_特記事項", result.fields)
        doc = Document(io.BytesIO(result.artifacts[0].content))
        self.assertIn("人間が書いた注意書き", "\n".join(p.text for p in doc.paragraphs))


class TestPrepareArtifacts(unittest.TestCase):
    def test_prepare_outputs_docx_and_label_pdf_and_metadata(self):
        records = [block("委任契約書2通"), block("返信用封筒", units=("時効援用",))]
        records[0]["返送要否"] = {"value": "要"}
        with patch.dict("os.environ", _OFFICE_ENV), \
             patch("hub.kintone.search_records", new=AsyncMock(return_value=records)):
            result = run(soufu_annai.SoufuAnnaiAdapter().prepare(make_shipping_record()))
        names = [a.filename for a in result.artifacts]
        self.assertEqual(names, ["送付案内.docx", "宛名ラベル.pdf"])
        self.assertTrue(result.artifacts[1].content.startswith(b"%PDF"))
        import json as _json
        meta = _json.loads(result.fields["チャネル固有データ"])
        self.assertTrue(meta["needs_return"], "返送要否=要 のブロックを含むためフラグON")
        self.assertEqual(meta["blocks"], ["委任契約書2通", "返信用封筒"])

    def test_dispatch_is_manual_mailing(self):
        result = run(soufu_annai.SoufuAnnaiAdapter().dispatch(make_shipping_record()))
        self.assertTrue(result.manual_mailing)


class TestEndToEnd(unittest.IsolatedAsyncioTestCase):
    """██ 起票→承認→出力の一巡（T2-2 完了条件）██
    下書き Webhook → prepare（docx+ラベル添付・承認待ち・弁護士通知）
    → 承認済 Webhook → claim → 発送処理中 → 印刷指示（manual_mailing で停止）"""

    async def test_full_cycle(self):
        import copy

        from hub import dispatch

        store_rec = {
            "$id": {"value": "9"}, "$revision": {"value": "1"},
            "発送ステータス": {"value": "下書き"},
            "チャネル": {"value": "送付案内"},
            "ユニット種別": {"value": "時効援用"},
            "件名": {"value": "委任契約書の送付"},
            "顧客名表示用": {"value": "山田太郎"},
            "宛先名": {"value": "山田太郎"},
            "宛先郵便番号": {"value": "332-0001"},
            "宛先住所": {"value": "埼玉県川口市1-1"},
            "本文_特記事項": {"value": ""},
            "同封物選択": {"value": ["委任契約書2通", "返信用封筒"]},
            "実行済み": {"value": "no"},
        }
        records = {"9": store_rec}
        uploaded, updates = [], []

        async def fake_get(app, rid):
            return copy.deepcopy(records[rid])

        async def fake_update(app, rid, fields, revision=None):
            rec = records[rid]
            cur = int(rec["$revision"]["value"])
            if revision is not None and int(revision) != cur:
                raise kintone.KintoneConflict(409, "GAIA_CO02", "conflict")
            for k, v in fields.items():
                rec[k] = {"value": v}
            rec["$revision"] = {"value": str(cur + 1)}
            updates.append(dict(fields))

        async def fake_upload(app, filename, content, mime):
            uploaded.append(filename)
            return f"fk_{len(uploaded)}"

        blocks = [block("委任契約書2通", order="1"), block("返信用封筒", order="2")]
        notify_admin, notify_attorney = AsyncMock(), AsyncMock()

        with patch.dict("os.environ", _OFFICE_ENV), \
             patch("hub.kintone.get_record", new=fake_get), \
             patch("hub.kintone.update_record", new=fake_update), \
             patch("hub.kintone.upload_file", new=fake_upload), \
             patch("hub.kintone.search_records", new=AsyncMock(return_value=blocks)), \
             patch("channels.soufu_annai.create_message_with_fallback",
                   new=AsyncMock(return_value=_FakeToolResponse("ご返送ください。"))), \
             patch("hub.notify.notify_admin_line", new=notify_admin), \
             patch("hub.notify.notify_attorney_approval", new=notify_attorney):

            # ① 起票（下書き保存）→ Webhook
            await dispatch.process_dispatch("9")
            self.assertEqual(records["9"]["発送ステータス"]["value"], "承認待ち")
            self.assertEqual(uploaded, ["送付案内.docx", "宛名ラベル.pdf"])
            self.assertEqual(records["9"]["成果物"]["value"],
                             [{"fileKey": "fk_1"}, {"fileKey": "fk_2"}])
            self.assertEqual(records["9"]["本文_特記事項"]["value"], "ご返送ください。")
            notify_attorney.assert_awaited_once()

            # ② 弁護士が kintone 上で承認（人の操作をストアで再現）
            records["9"]["発送ステータス"] = {"value": "承認済"}

            # ③ 承認済 Webhook → claim → 発送処理中 → 印刷指示で停止
            await dispatch.process_dispatch("9")
            self.assertEqual(records["9"]["実行済み"]["value"], "yes")
            self.assertEqual(records["9"]["発送ステータス"]["value"], "発送処理中")
            texts = [c.args[0] for c in notify_admin.await_args_list]
            self.assertTrue(any("印刷・投函" in t for t in texts))

            # ④ 二重 Webhook でも再実行されない（冪等）
            await dispatch.process_dispatch("9")
            self.assertEqual(records["9"]["発送ステータス"]["value"], "発送処理中")


class TestBlockSync(unittest.TestCase):
    """App30/32 同期検査（監視項目D）の異常系"""

    _ENV = {"APP_ENCLOSURE": "32", "TOKEN_ENCLOSURE": "d",
            "APP_SHIPPING": "30", "TOKEN_SHIPPING": "d"}

    def _run(self, options, keys, fields_error=None):
        form = {"同封物選択": {"type": "CHECK_BOX",
                               "options": {o: {"label": o, "index": str(i)}
                                           for i, o in enumerate(options)}}}
        records = [{"ブロックキー": {"value": k}} for k in keys]
        get_form = AsyncMock(return_value=form)
        if fields_error:
            get_form = AsyncMock(side_effect=fields_error)
        with patch.dict("os.environ", self._ENV), \
             patch("hub.kintone.get_form_fields", new=get_form), \
             patch("hub.kintone.search_records", new=AsyncMock(return_value=records)):
            return run(check_block_sync())

    def test_key_missing_from_app30_options_is_reported(self):
        problems = self._run(options=["（未設定）"], keys=["委任契約書2通"])
        self.assertEqual(len(problems), 1)
        self.assertIn("委任契約書2通", problems[0])
        self.assertIn("App 30", problems[0])

    def test_synced_keys_pass(self):
        problems = self._run(options=["（未設定）", "委任契約書2通", "返信用封筒"],
                             keys=["委任契約書2通", "返信用封筒"])
        self.assertEqual(problems, [])

    def test_fetch_error_is_reported_as_problem(self):
        problems = self._run(options=[], keys=[],
                             fields_error=kintone.KintoneError(500, "X", "down"))
        self.assertEqual(len(problems), 1)
        self.assertIn("同期検査の実行に失敗", problems[0])

    def test_env_unset_skips(self):
        with patch.dict("os.environ", {"APP_ENCLOSURE": "", "APP_SHIPPING": ""}):
            problems = run(check_block_sync())
        self.assertEqual(problems, [])


if __name__ == "__main__":
    unittest.main()
