"""units/souzoku/zaisan_mokuroku.py（S3 財産目録）のテスト

検証（設計 02 §7 の S3 該当観点）:
種別セクション振り分け・合計（3桁カンマ+円・コード側計算）・債務の部・
空セクションの「該当なし」行・前提条件ガード（評価確定=no 混在で生成拒否）・
特定情報の列分解（推奨書式・非準拠フォールバック）・注記・
App 財産のモック取得（env 未設定は安全側で拒否）・ゴールデンファイル比較。
App 財産は未作成のため kintone 接続はすべてモック。
"""

import asyncio
import io
import os
import unittest
from datetime import date
from unittest.mock import AsyncMock, patch

os.environ.setdefault("KINTONE_SUBDOMAIN", "testsub")

from docx import Document

from units.souzoku.guards import ValuationNotConfirmed, ensure_valuations_confirmed
from units.souzoku.zaisan_mokuroku import (
    ZaisanMokurokuError,
    build_zaisan_mokuroku_docx,
    fetch_zaisan_records,
    generate_zaisan_mokuroku,
)

_OFFICE_ENV = {
    "OFFICE_NAME": "大野法律事務所",
    "OFFICE_ATTORNEY": "大野　太郎",
}


def run(coro):
    return asyncio.run(coro)


def _all_text(docx_bytes: bytes) -> str:
    doc = Document(io.BytesIO(docx_bytes))
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                parts.append(c.text)
    return "\n".join(parts)


def zrec(rec_id, shubetsu, tokutei, hyouka, *, kakutei="yes", kijunbi="",
         source="", name="山田太郎"):
    """App 財産のレコード（cu-app35-zaisan.md の19フィールドのうち本機能が読む分）"""
    return {
        "$id": {"value": str(rec_id)},
        "財産種別": {"value": shubetsu},
        "特定情報": {"value": tokutei},
        "評価額": {"value": "" if hyouka == "" else str(hyouka)},
        "評価確定": {"value": kakutei},
        "評価基準日": {"value": kijunbi},
        "データ源": {"value": source},
        "被相続人名表示用": {"value": name},
        "有効": {"value": "yes"},
    }


def _full_records():
    """全セクションを埋める固定データ（合計: 積極 34,704,551 / 消極 6,000,000）"""
    return [
        zrec(1, "不動産_土地",
             "所在 川口市朝日一丁目 / 地番 12番3 / 地目 宅地 / 地積 120.45㎡ / 持分 2分の1",
             15000000, kijunbi="2026-04-01", source="OCR_課税明細"),
        zrec(2, "不動産_建物",
             "所在 川口市朝日一丁目12番地3 / 家屋番号 12番3 / 種類 居宅 / 床面積 95.20㎡",
             8000000, kijunbi="2026-04-01", source="OCR_課税明細"),
        zrec(3, "預貯金", "みずほ銀行 川口支店 普通預金 口座番号1234567",
             3204551, kijunbi="2026-03-15", source="OCR_残高証明"),
        zrec(4, "預貯金", "ゆうちょ銀行 － 通常貯金 口座番号98765432",
             1000000, kijunbi="2026-03-15", source="手入力"),
        zrec(5, "有価証券", "○○証券　△△投資信託", 2500000, source="手入力"),
        zrec(6, "生命保険", "○○生命 終身保険 証券番号12-345678",
             5000000, source="ヒアリング"),
        zrec(7, "債務", "○○銀行 住宅ローン（残債）", 4200000, source="手入力"),
        zrec(8, "葬儀費用", "葬儀費用一式", 1800000, source="ヒアリング"),
    ]


def _build(records, **kw):
    with patch.dict("os.environ", _OFFICE_ENV):
        return build_zaisan_mokuroku_docx(records, created=date(2026, 7, 5), **kw)


class TestGuard(unittest.TestCase):
    """前提条件ガード（02 §6・協議書/遺言と共有する単一実装）"""

    def test_unconfirmed_rows_raise_with_record_ids(self):
        records = [zrec(1, "預貯金", "x", 1, kakutei="yes"),
                   zrec(2, "預貯金", "y", 1, kakutei="no"),
                   zrec(3, "預貯金", "z", 1, kakutei="")]
        with self.assertRaises(ValuationNotConfirmed) as cm:
            ensure_valuations_confirmed(records)
        self.assertIn("2 件", str(cm.exception))
        self.assertIn("2, 3", str(cm.exception))

    def test_all_confirmed_passes(self):
        ensure_valuations_confirmed([zrec(1, "預貯金", "x", 1)])

    def test_empty_list_passes(self):
        ensure_valuations_confirmed([])


class TestFetch(unittest.TestCase):
    """App 財産（未作成）の取得はモック。env 未設定は API を呼ばず安全側で拒否"""

    def test_env_unset_raises_without_api_call(self):
        mock = AsyncMock()
        with patch.dict("os.environ", {"APP_ZAISAN": ""}), \
             patch("hub.kintone.search_records", new=mock):
            with self.assertRaises(ZaisanMokurokuError) as cm:
                run(fetch_zaisan_records("26", "100"))
        self.assertIn("APP_ZAISAN", str(cm.exception))
        mock.assert_not_awaited()

    def test_query_filters_case_and_active(self):
        mock = AsyncMock(return_value=[])
        with patch.dict("os.environ", {"APP_ZAISAN": "35", "TOKEN_ZAISAN": "t"}), \
             patch("hub.kintone.search_records", new=mock):
            run(fetch_zaisan_records("26", "100"))
        app, query = mock.await_args.args
        self.assertEqual(app.app_id_env, "APP_ZAISAN")
        self.assertIn('案件アプリID = "26"', query)
        self.assertIn('案件レコードID = "100"', query)
        self.assertIn('有効 in ("yes")', query)


class TestSectionRouting(unittest.TestCase):
    """種別セクション振り分けと特定情報の列分解（02 §2.3 推奨書式）"""

    def setUp(self):
        self.doc = Document(io.BytesIO(_build(_full_records())))
        # 表の並びはテンプレート規約: 0=不動産 1=預貯金 2=有価証券その他 3=債務
        self.fudousan, self.yokin, self.sonota, self.saimu = self.doc.tables

    def test_row_counts_per_section(self):
        self.assertEqual([len(t.rows) for t in
                          (self.fudousan, self.yokin, self.sonota, self.saimu)],
                         [1 + 2, 1 + 2, 1 + 2, 1 + 2], "ヘッダ+件数")

    def test_fudousan_columns_parsed(self):
        row = self.fudousan.rows[1]
        self.assertEqual([c.text for c in row.cells],
                         ["川口市朝日一丁目", "12番3", "宅地", "120.45㎡",
                          "2分の1", "15,000,000円"])
        row2 = self.fudousan.rows[2]  # 建物: 家屋番号・種類・床面積の別名キー
        self.assertEqual([c.text for c in row2.cells],
                         ["川口市朝日一丁目12番地3", "12番3", "居宅", "95.20㎡",
                          "", "8,000,000円"])

    def test_yokin_columns_parsed(self):
        self.assertEqual([c.text for c in self.yokin.rows[1].cells],
                         ["みずほ銀行", "川口支店", "普通預金", "1234567",
                          "3,204,551円"])

    def test_sonota_section_collects_non_realestate_non_deposit(self):
        contents = [r.cells[0].text for r in self.sonota.rows[1:]]
        self.assertEqual(contents, ["○○証券　△△投資信託",
                                    "○○生命 終身保険 証券番号12-345678"])

    def test_saimu_section(self):
        self.assertEqual([c.text for c in self.saimu.rows[1].cells],
                         ["○○銀行 住宅ローン（残債）", "4,200,000円"])
        self.assertEqual(self.saimu.rows[2].cells[1].text, "1,800,000円")


class TestTotalsAndNotes(unittest.TestCase):
    def setUp(self):
        self.text = _all_text(_build(_full_records()))

    def test_totals_computed_in_code_with_comma_yen(self):
        self.assertIn("積極財産合計　34,704,551円", self.text)
        self.assertIn("消極財産合計　6,000,000円", self.text)
        self.assertIn("純資産額　28,704,551円", self.text)

    def test_cover_fields(self):
        self.assertIn("被相続人　山田太郎", self.text)
        self.assertIn("作成日　令和8年7月5日", self.text)
        self.assertIn("作成者　大野法律事務所　弁護士　大野　太郎", self.text)

    def test_notes_kijunbi_and_sources(self):
        self.assertIn("評価基準日: 令和8年3月15日、令和8年4月1日", self.text)
        self.assertIn("出典資料: OCR_課税明細、OCR_残高証明、手入力、ヒアリング",
                      self.text)

    def test_no_placeholder_left(self):
        self.assertNotIn("{{", self.text)


class TestEmptySections(unittest.TestCase):
    def test_empty_sections_get_nashi_row(self):
        """預貯金のみ → 他3セクションは「該当なし」1行（S2 empty_text）"""
        records = [zrec(1, "預貯金", "みずほ銀行 川口支店 普通預金 口座番号1234567",
                        1000000)]
        doc = Document(io.BytesIO(_build(records)))
        fudousan, yokin, sonota, saimu = doc.tables
        for table in (fudousan, sonota, saimu):
            self.assertEqual(len(table.rows), 2, "ヘッダ+該当なし行")
            self.assertEqual(table.rows[1].cells[0].text, "該当なし")
        self.assertEqual(len(yokin.rows), 2)
        self.assertEqual(yokin.rows[1].cells[0].text, "みずほ銀行")

    def test_totals_with_empty_negative_section(self):
        records = [zrec(1, "預貯金", "みずほ銀行 川口支店 普通預金 口座番号1234567",
                        1000000)]
        text = _all_text(_build(records))
        self.assertIn("消極財産合計　0円", text)
        self.assertIn("純資産額　1,000,000円", text)


class TestGenerationRefusal(unittest.TestCase):
    def test_unconfirmed_valuation_refuses_generation(self):
        """評価確定=no が1件でもあれば生成しない（02 §6）"""
        records = _full_records()
        records[3]["評価確定"]["value"] = "no"
        with self.assertRaises(ValuationNotConfirmed):
            _build(records)

    def test_empty_records_refuse_generation(self):
        with self.assertRaises(ZaisanMokurokuError):
            _build([])


class TestTokuteiJohoFallback(unittest.TestCase):
    """推奨書式（02 §2.3）に合わない特定情報は全文を先頭列に（情報を落とさない）"""

    def test_fudousan_nonstandard_goes_to_first_column(self):
        records = [zrec(1, "不動産_土地", "川口市の実家の土地", 1000000)]
        doc = Document(io.BytesIO(_build(records)))
        row = doc.tables[0].rows[1]
        self.assertEqual(row.cells[0].text, "川口市の実家の土地")
        self.assertEqual([c.text for c in row.cells[1:5]], ["", "", "", ""])
        self.assertEqual(row.cells[5].text, "1,000,000円")

    def test_yokin_nonstandard_goes_to_first_column(self):
        records = [zrec(1, "預貯金", "タンス預金", 300000)]
        doc = Document(io.BytesIO(_build(records)))
        row = doc.tables[1].rows[1]
        self.assertEqual(row.cells[0].text, "タンス預金")
        self.assertEqual(row.cells[4].text, "300,000円")

    def test_blank_hyoukagaku_prints_blank_and_sums_as_zero(self):
        records = [zrec(1, "預貯金", "みずほ銀行 川口支店 普通預金 口座番号1234567",
                        ""),
                   zrec(2, "有価証券", "○○株式", 500000)]
        text = _all_text(_build(records))
        self.assertIn("積極財産合計　500,000円", text)


class TestGenerateEndToEnd(unittest.TestCase):
    def test_generate_with_mocked_kintone(self):
        mock = AsyncMock(return_value=_full_records())
        with patch.dict("os.environ",
                        {**_OFFICE_ENV, "APP_ZAISAN": "35", "TOKEN_ZAISAN": "t"}), \
             patch("hub.kintone.search_records", new=mock):
            out = run(generate_zaisan_mokuroku("26", "100"))
        self.assertTrue(out.startswith(b"PK"))
        text = _all_text(out)
        self.assertIn("被相続人　山田太郎", text, "被相続人名表示用から補完")
        self.assertIn("34,704,551円", text)


GOLDEN = """\
財　産　目　録

被相続人　山田太郎
作成日　令和8年7月5日
作成者　大野法律事務所　弁護士　大野　太郎

第１　積極財産の部
Ⅰ　不動産
Ⅱ　預貯金
Ⅲ　有価証券・その他

第２　消極財産の部（債務・葬儀費用）

積極財産合計　34,704,551円
消極財産合計　6,000,000円
純資産額　28,704,551円

（注記）
評価基準日: 令和8年3月15日、令和8年4月1日
出典資料: OCR_課税明細、OCR_残高証明、手入力、ヒアリング
所在
地番・家屋番号
地目・種類
地積・床面積
持分
評価額
川口市朝日一丁目
12番3
宅地
120.45㎡
2分の1
15,000,000円
川口市朝日一丁目12番地3
12番3
居宅
95.20㎡

8,000,000円
金融機関
支店
種別
口座番号
死亡日残高
みずほ銀行
川口支店
普通預金
1234567
3,204,551円
ゆうちょ銀行
－
通常貯金
98765432
1,000,000円
銘柄・内容
数量
評価額
○○証券　△△投資信託

2,500,000円
○○生命 終身保険 証券番号12-345678

5,000,000円
内容
金額
○○銀行 住宅ローン（残債）
4,200,000円
葬儀費用一式
1,800,000円"""


class TestGoldenFile(unittest.TestCase):
    def test_full_text_matches_golden(self):
        """固定入力→抽出全文の完全一致（意図しない出力変化の検知・02 §7）"""
        text = _all_text(_build(_full_records()))
        self.assertEqual(text, GOLDEN)


if __name__ == "__main__":
    unittest.main()
