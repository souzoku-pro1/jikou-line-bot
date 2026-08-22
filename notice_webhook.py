"""時効援用通知書 自動生成 Webhook — JIKOU-NOTICE-1

CONTRACT-GEN で確立した構造の同型（App21・path=/notice/・secret 共用）:

状態機械（CAS＝$revision 楽観ロック・契約書ステータス DROP_DOWN を共用し
本フローの閉集合 3 値+要確認を追加）:
  時効援用通知作成 --（CAS 勝者: 作成→作成中)--> 時効援用通知作成中
    --（N 通生成+凍結検証+upload×N+添付 PUT〔revision=claim+1〕）-->
    時効援用通知作成済
  時効援用通知作成(必須欠落・債権者 0 件) → 変更なし（不足名のみ通知）
  時効援用通知作成中 + 添付なし → 回収: CAS 再claim → 再生成
  時効援用通知作成中 + 添付あり → 自動上書きせず「要確認」+管理者通知
  時効援用通知作成済 → already_done skip／要確認・他値 → stale_status skip
  CAS 敗者（409）→ 作用 0（fix2 流儀: 409 のみ cas_lost・障害系は 500 →
  kintone 再配送）

大野裁定（改変禁止）:
  - 宛先表記はテンプレのまま「債権者各位」（変更しない）
  - 冒頭の事務所住所ブロック（〒・住所・建物・通知代理人・TEL/FAX）は
    固定＝凍結本文の一部
  - 旧住所行: old_address が空なら該当段落ごと削除・非空なら差し込み
  - 対象債権者 1〜3（委任契約書と同じ欄構成）の非空ごとに 1 通生成し、
    全 docx を notice_file へ複数添付（単一レコード・1 PUT）。宛先が
    「債権者各位」固定のため各通の本文は同一（枠番ファイル名で区別）

本文凍結検証（実行時）:
  - テンプレは収載現物の SHA-256 と完全一致（TEMPLATE_SHA256）
  - 生成物の全段落テキストが「テンプレへ同一の差し込み・旧住所行削除を
    適用した期待列」と完全一致（＝プレースホルダ差し込み箇所以外の本文は
    テンプレと逐語一致）。プレースホルダ残存も拒否。不一致は添付せず 500
"""

import hashlib
import io
import logging
from datetime import date

from docx import Document
from fastapi import APIRouter, Request
from fastapi.responses import JSONResponse

from hub import kintone as hub_kintone
from hub.docx_builder import to_wareki
from hub.redact import emit
from hub.webhook_auth import extract_record_id, verify_token

logger = logging.getLogger("notice")

_APP = hub_kintone.KintoneApp(
    "App 21 (案件)", "KINTONE_APP_ID", "KINTONE_API_TOKEN")

FIELD_STATUS     = "契約書ステータス"      # contract と同一 DROP_DOWN を共用
STATUS_TRIGGER   = "時効援用通知作成"
STATUS_WORKING   = "時効援用通知作成中"
STATUS_DONE      = "時効援用通知作成済"
STATUS_REVIEW    = "要確認"
FIELD_ATTACHMENT = "notice_file"           # ラベル「時効援用通知書」（CU 新設）
FIELD_FURIGANA   = "furigana"              # ラベル「ふりがな」（CU 新設）
FIELD_OLD_ADDR   = "old_address"           # ラベル「旧住所」（CU 新設）
TEMPLATE_PATH    = "docx_templates/jikou/時効援用通知書.docx"
# 収載現物（make_notice_template.py 生成・コミット済み artifact）の pin
TEMPLATE_SHA256 = (
    "828fcee68bbee193e3b5f6946f294fce6a10a481d02bd834016dfc7f243fa725")

_REQUIRED_NAME   = "顧客名"
_REQUIRED_ADDR   = "住所"
_REQUIRED_BIRTH  = "生年月日"
_CREDITOR_FIELDS = ("問い合わせ業者名", "対象債権者2", "対象債権者3")
_OLD_ADDR_KEY    = "{{旧住所}}"

router = APIRouter()


class NoticeIntegrityError(RuntimeError):
    """テンプレ pin 不一致・生成物の凍結検証失敗（添付しない）。"""


def _fv(record: dict, code: str) -> str:
    return str((record.get(code) or {}).get("value") or "").strip()


def _creditors(record: dict) -> list[tuple[int, str]]:
    """fix1[01]: （元の枠番号, 値）の組で返す——枠1=問い合わせ業者名/
    枠2=対象債権者2/枠3=対象債権者3。空欄は除くが**枠番号は保持**し、
    ファイル命名「時効援用通知書_対象債権者N.docx」の N は常に元の枠番号
    （App21 の入力枠との監査対応）。枠番号昇順=生成・添付順も決定的。"""
    out = []
    for slot, code in enumerate(_CREDITOR_FIELDS, start=1):
        v = _fv(record, code)
        if v:
            out.append((slot, v))
    return out


def _missing_fields(record: dict) -> list[str]:
    missing = [label for label, code in (
        ("顧客名", _REQUIRED_NAME), ("住所", _REQUIRED_ADDR),
        ("生年月日", _REQUIRED_BIRTH), ("ふりがな", FIELD_FURIGANA),
    ) if not _fv(record, code)]
    if not _creditors(record):
        missing.append("債権者（問い合わせ業者名/対象債権者2/対象債権者3 の"
                       "いずれか1つ以上）")
    return missing


def build_fill_data(record: dict) -> dict:
    return {
        "{{通知日付}}":   to_wareki(date.today()),
        "{{通知人氏名}}": _fv(record, _REQUIRED_NAME),
        "{{ふりがな}}":   _fv(record, FIELD_FURIGANA),
        "{{生年月日}}":   _fv(record, _REQUIRED_BIRTH),
        "{{通知人住所}}": _fv(record, _REQUIRED_ADDR),
    }


def verify_template_integrity() -> None:
    data = open(TEMPLATE_PATH, "rb").read()
    if hashlib.sha256(data).hexdigest() != TEMPLATE_SHA256:
        raise NoticeIntegrityError("template hash mismatch")


def _set_paragraph_text(p, text: str) -> None:
    if not p.runs:
        p.add_run(text)
        return
    p.runs[0].text = text
    for r in p.runs[1:]:
        r.text = ""


def build_notice_docx(fill: dict, old_address: str) -> bytes:
    """テンプレへ差し込み（旧住所行は空なら段落ごと削除）して bytes を返す。"""
    doc = Document(TEMPLATE_PATH)
    for p in list(doc.paragraphs):
        text = p.text
        if _OLD_ADDR_KEY in text:
            if old_address:
                _set_paragraph_text(p, text.replace(_OLD_ADDR_KEY,
                                                    old_address))
            else:
                p._element.getparent().remove(p._element)
            continue
        if "{{" in text:
            for k, v in fill.items():
                text = text.replace(k, v)
            _set_paragraph_text(p, text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def expected_paragraph_texts(fill: dict, old_address: str) -> list[str]:
    """凍結検証の期待列: テンプレ全段落へ同一変換を適用したテキスト列。
    プレースホルダ以外の本文はテンプレ逐語のまま残る＝全段落完全一致検査で
    「差し込み箇所以外の逐語一致」を保証する。"""
    tdoc = Document(TEMPLATE_PATH)
    out = []
    for p in tdoc.paragraphs:
        text = p.text
        if _OLD_ADDR_KEY in text:
            if not old_address:
                continue
            text = text.replace(_OLD_ADDR_KEY, old_address)
        else:
            for k, v in fill.items():
                text = text.replace(k, v)
        out.append(text)
    return out


def verify_notice_docx(docx_bytes: bytes, fill: dict,
                       old_address: str) -> None:
    """生成物の全段落がテンプレ由来の期待列と完全一致（凍結本文の逐語
    一致+差し込み値の一致+旧住所行の有無）・プレースホルダ残存なし。"""
    got = [p.text for p in Document(io.BytesIO(docx_bytes)).paragraphs]
    if got != expected_paragraph_texts(fill, old_address):
        raise NoticeIntegrityError("body mismatch against template")
    joined = "\n".join(got)
    if "{{" in joined or "}}" in joined:
        raise NoticeIntegrityError("unfilled placeholder")


async def _notify(text: str) -> None:
    """管理者 LINE 通知（best-effort・固定文言+レコード番号/件数のみ）。"""
    try:
        from hub.notify import notify_admin_line
        await notify_admin_line(text)
    except Exception:
        logger.error("[NOTICE] admin notify failed (fixed text)")


async def _generate_and_attach(record_id: str, record: dict,
                               final_revision: str) -> None:
    """N 通生成 → 凍結検証 → upload×N → 添付+作成済（revision CAS PUT）。"""
    verify_template_integrity()
    fill = build_fill_data(record)
    old_address = _fv(record, FIELD_OLD_ADDR)
    docx_bytes = build_notice_docx(fill, old_address)
    verify_notice_docx(docx_bytes, fill, old_address)
    creditors = _creditors(record)
    logger.info("[NOTICE] generated record_id=%s bytes=%d copies=%d",
                emit(record_id, "record_id", "log", "operator"),
                len(docx_bytes), len(creditors))
    mime = ("application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.document")
    # fix1[02]（採用方式 (a)）: upload を全通成功させてから単一 PUT へ進む。
    # 途中 upload/最終 PUT の失敗で未添付の fileKey が残っても、kintone
    # 公式仕様により一時保管領域のファイルは「レコードやスペースなどに
    # 添付されない場合、3 日間で削除される」ため孤立 fileKey は自動回収
    # される（cybozu developer network「ファイルをアップロードする」
    # https://cybozu.dev/ja/kintone/docs/rest-api/files/upload-file/
    # 制限事項・2026-08-23 確認）。再実行（reconcile 回収）は最初から
    # 再 upload してよい
    file_keys = []
    for slot, _name in creditors:
        file_keys.append(await hub_kintone.upload_file(
            _APP, f"時効援用通知書_対象債権者{slot}.docx", docx_bytes, mime))
    await hub_kintone.update_record(_APP, record_id, {
        FIELD_ATTACHMENT: [{"fileKey": k} for k in file_keys],
        FIELD_STATUS: STATUS_DONE,
    }, revision=final_revision)
    logger.info("[NOTICE] attached record_id=%s files=%d",
                emit(record_id, "record_id", "log", "operator"),
                len(file_keys))


async def _claim(record_id: str, revision: str, to_status: str) -> str | None:
    """CAS（fix2 流儀）: 409 競合のみ cas_lost（None）。障害系は再送出し
    外側で 500 → kintone 再配送へ。"""
    try:
        await hub_kintone.update_record(
            _APP, record_id, {FIELD_STATUS: to_status}, revision=revision)
    except hub_kintone.KintoneError as e:
        if getattr(e, "status", None) == 409:
            return None
        raise
    return str(int(revision) + 1)


async def _reconcile_working(record_id: str, record: dict, revision: str):
    """「作成中」で停止した行の回収（contract_webhook と同規則）。"""
    attachment = (record.get(FIELD_ATTACHMENT) or {}).get("value") or []
    if attachment:
        next_rev = await _claim(record_id, revision, STATUS_REVIEW)
        if next_rev is None:
            return JSONResponse(status_code=200, content={
                "ok": True, "skip": "cas_lost"})
        await _notify(
            f"【時効援用通知書】案件 No.{record_id} は生成が中断した状態で"
            "既に添付ファイルが存在するため、自動では上書きせず"
            f"「{STATUS_REVIEW}」にしました。添付内容を確認のうえ、再生成する"
            f"場合は添付を削除してステータスを「{STATUS_TRIGGER}」に設定して"
            "ください")
        return JSONResponse(status_code=200, content={
            "ok": True, "skip": "needs_review"})
    next_rev = await _claim(record_id, revision, STATUS_WORKING)
    if next_rev is None:
        return JSONResponse(status_code=200,
                            content={"ok": True, "skip": "cas_lost"})
    await _generate_and_attach(record_id, record, next_rev)
    return JSONResponse(status_code=200, content={
        "ok": True, "record_id": record_id, "recovered": True})


@router.post("/notice/{secret}")
async def notice_webhook(secret: str, request: Request):
    if not verify_token(secret or "", "DOCUMENT_WEBHOOK_SECRET"):
        return JSONResponse(status_code=403, content={"error": "forbidden"})

    try:
        body = await request.json()
    except Exception:
        return JSONResponse(status_code=400, content={"error": "invalid json"})

    app_in_body = str(((body.get("app") or {}) if isinstance(body, dict)
                       else {}).get("id") or "")
    if not app_in_body.isdigit() or app_in_body != str(_APP.app_id()):
        logger.warning("[NOTICE] app mismatch in webhook body")
        return JSONResponse(status_code=200,
                            content={"ok": True, "skip": "app_mismatch"})

    record_id = extract_record_id(body)
    if not record_id:
        logger.warning("[NOTICE] record id missing in webhook body")
        return JSONResponse(status_code=200,
                            content={"ok": True, "skip": "no_record_id"})

    # 本文ステータス gate（contract/クラウドサイン系の値・自 update の echo
    # はここで落ちる=共用 DROP_DOWN でも相互不干渉）
    try:
        status_in_webhook = body["record"][FIELD_STATUS]["value"]
    except (KeyError, TypeError):
        status_in_webhook = None
    if status_in_webhook != STATUS_TRIGGER:
        logger.info("[NOTICE] not triggered record_id=%s",
                    emit(record_id, "record_id", "log", "operator"))
        return JSONResponse(status_code=200,
                            content={"ok": True, "skip": "not_triggered"})

    try:
        record = await hub_kintone.get_record(_APP, record_id)
        current = _fv(record, FIELD_STATUS)
        revision = _fv(record, "$revision")

        if current == STATUS_DONE:
            logger.info("[NOTICE] already done record_id=%s",
                        emit(record_id, "record_id", "log", "operator"))
            return JSONResponse(status_code=200,
                                content={"ok": True, "skip": "already_done"})
        if current == STATUS_WORKING:
            return await _reconcile_working(record_id, record, revision)
        if current != STATUS_TRIGGER or not revision.isdigit():
            logger.info("[NOTICE] stale status record_id=%s",
                        emit(record_id, "record_id", "log", "operator"))
            return JSONResponse(status_code=200,
                                content={"ok": True, "skip": "stale_status"})

        # fail-closed: 必須欠落・債権者 0 件は生成しない（状態も動かさない・
        # 通知はフィールド名のみ=値・PII 非搭載）
        missing = _missing_fields(record)
        if missing:
            logger.info("[NOTICE] missing required fields record_id=%s "
                        "count=%d",
                        emit(record_id, "record_id", "log", "operator"),
                        len(missing))
            await _notify(
                f"【時効援用通知書】案件 No.{record_id} は必須項目が未入力の"
                f"ため生成しませんでした。不足: {'・'.join(missing)}。"
                "kintone で入力後、契約書ステータスを"
                f"「{STATUS_TRIGGER}」に設定し直してください")
            return JSONResponse(status_code=200, content={
                "ok": True, "skip": "missing_fields", "missing": missing})

        next_rev = await _claim(record_id, revision, STATUS_WORKING)
        if next_rev is None:
            logger.info("[NOTICE] cas lost record_id=%s",
                        emit(record_id, "record_id", "log", "operator"))
            return JSONResponse(status_code=200,
                                content={"ok": True, "skip": "cas_lost"})
        await _generate_and_attach(record_id, record, next_rev)
    except Exception as e:
        logger.error("[NOTICE] error record_id=%s cls=%s: %s",
                     emit(record_id, "record_id", "log", "operator"),
                     type(e).__name__,
                     emit(str(e), "vendor_raw", "log", "operator"))
        return JSONResponse(status_code=500,
                            content={"error": "internal_error"})

    return JSONResponse(status_code=200,
                        content={"ok": True, "record_id": record_id})
