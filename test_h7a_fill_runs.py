"""SOUZOKU-HOUKI-H7A: _fill_runs の共通ヘルパ昇格の固定。

- hub/docx_builder.fill_runs = notice_webhook._fill_runs（JIKOU-NOTICE-1-fix2）
  の逐語昇格。notice 側は同一関数への委譲（identity pin）＝挙動変更ゼロ。
- fill_runs の契約: プレースホルダを含む run の中だけで置換し、他 run の
  text・rPr（w:fitText 等）に触れない。run 跨ぎのプレースホルダは置換されず
  残る（呼び出し側の残存検査で fail-closed 拒否する設計）。
- fill_template（段落潰し・既存利用者互換のため実装不変）には既知問題と
  「新規テンプレは fill_runs を使う」の注記が明文化されていること。
"""

import io
import unittest

from docx import Document
from docx.oxml.ns import qn

from hub import docx_builder


def _make_para(runs: list[str]):
    """指定テキストの run 列を持つ段落を生成（先頭 run に w:fitText を付与）。"""
    doc = Document()
    p = doc.add_paragraph()
    for i, text in enumerate(runs):
        r = p.add_run(text)
        if i == 0:
            rpr = r._r.get_or_add_rPr()
            fit = rpr.makeelement(qn("w:fitText"), {qn("w:val"): "1050"})
            rpr.append(fit)
    return doc, p


def _fittext_vals(p):
    out = []
    for r in p.runs:
        rpr = r._r.rPr
        fits = [] if rpr is None else rpr.findall(qn("w:fitText"))
        out.append([f.get(qn("w:val")) for f in fits])
    return out


class TestFillRunsContract(unittest.TestCase):
    def test_replaces_only_within_placeholder_run(self):
        doc, p = _make_para(["ふりがな", "{{ふりがな}}", "様"])
        docx_builder.fill_runs(p, {"{{ふりがな}}": "やまだたろう"})
        self.assertEqual([r.text for r in p.runs],
                         ["ふりがな", "やまだたろう", "様"])
        # 先頭ラベル run の fitText は不変・他 run へ波及しない（run 数も不変）
        self.assertEqual(_fittext_vals(p), [["1050"], [], []])

    def test_multiple_keys_and_untouched_runs(self):
        _doc, p = _make_para(["ラベル", "{{a}}と{{b}}", "末尾{{c}}"])
        docx_builder.fill_runs(p, {"{{a}}": "A", "{{b}}": "B", "{{c}}": "C"})
        self.assertEqual([r.text for r in p.runs], ["ラベル", "AとB", "末尾C"])

    def test_cross_run_placeholder_left_intact(self):
        # run 跨ぎ（"{{" が片側 run に無い形）は置換されず残る＝呼び出し側の
        # 残存検査（fail-closed）で拒否される設計の pin
        _doc, p = _make_para(["{{ふり", "がな}}"])
        docx_builder.fill_runs(p, {"{{ふりがな}}": "x"})
        self.assertEqual([r.text for r in p.runs], ["{{ふり", "がな}}"])

    def test_no_placeholder_paragraph_untouched(self):
        _doc, p = _make_para(["本文のみ", "そのまま"])
        docx_builder.fill_runs(p, {"{{k}}": "v"})
        self.assertEqual([r.text for r in p.runs], ["本文のみ", "そのまま"])


class TestPromotionWiring(unittest.TestCase):
    def test_notice_delegates_to_hub_helper(self):
        # notice 側は昇格ヘルパそのもの（identity）＝逐語移設・挙動変更ゼロ
        import notice_webhook
        self.assertIs(notice_webhook._fill_runs, docx_builder.fill_runs)

    def test_fill_template_docstring_notes_known_issue(self):
        doc = docx_builder.fill_template.__doc__ or ""
        self.assertIn("既知問題", doc)
        self.assertIn("fitText", doc)
        self.assertIn("fill_runs", doc)
        self.assertIn("新規テンプレ", doc)
        helper_doc = docx_builder.fill_runs.__doc__ or ""
        self.assertIn("単一 run", helper_doc)
        self.assertIn("rPr", helper_doc)

    def test_notice_build_still_works_via_helper(self):
        # 実テンプレでの生成が委譲後も成立（凍結検証は既存 test_jikou_notice1
        # が担う——ここでは生成物が docx として開け、差し込みが反映される
        # ことだけ確認する軽量スモーク）
        import notice_webhook as nw
        fill = {
            "{{通知日付}}": "令和8年8月25日",
            "{{通知人氏名}}": "試験太郎",
            "{{ふりがな}}": "しけんたろう",
            "{{生年月日}}": "平成2年1月1日",
            "{{通知人住所}}": "埼玉県川口市1-1",
        }
        data = nw.build_notice_docx(fill, "")
        doc = Document(io.BytesIO(data))
        text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("試験太郎", text)
        self.assertIn("しけんたろう", text)
        self.assertNotIn("{{", text)   # 全キー充足で残存なし


if __name__ == "__main__":
    unittest.main()
