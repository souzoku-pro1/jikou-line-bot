"""財産目録テンプレート（標準形）の生成スクリプト（S3）

docx_templates/souzoku/財産目録.docx を生成する。

- このテンプレートは標準形で、後日オーナーの書式こだわりで差し替える前提。
  差し替え時は config.EXPECTED_DOCX_TEMPLATES に登録したプレースホルダを
  すべて残すこと（daily_healthcheck のテンプレート検査が欠落を検知する）
- プレースホルダ規約は送付案内と同一: スカラーは {{名前}}、可変行は {{行:列名}}。
  各表の行マーカー（{{行:不動産}} 等）は S2 fill_table_rows の row_marker で、
  「該当なし」行の印字位置を兼ねるため各表テンプレート行の先頭セルに置く

実行: python make_zaisan_mokuroku_template.py
"""

import logging
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from hub.redact import emit

logger = logging.getLogger("make_zaisan_mokuroku_template")

OUT = Path("docx_templates") / "souzoku" / "財産目録.docx"


def _heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)


def _table(doc: Document, header: list[str], template_cells: list[str]) -> None:
    table = doc.add_table(rows=2, cols=len(header))
    table.style = "Table Grid"
    for i, text in enumerate(header):
        run = table.rows[0].cells[i].paragraphs[0].add_run(text)
        run.bold = True
    for i, text in enumerate(template_cells):
        table.rows[1].cells[i].text = text


def build() -> Document:
    doc = Document()

    # ── 表紙情報 ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("財　産　目　録")
    run.bold = True
    run.font.size = Pt(18)
    doc.add_paragraph("")
    doc.add_paragraph("被相続人　{{被相続人名}}")
    doc.add_paragraph("作成日　{{作成日}}")
    doc.add_paragraph("作成者　{{作成者}}")
    doc.add_paragraph("")

    # ── 積極財産の部 ──
    _heading(doc, "第１　積極財産の部")
    _heading(doc, "Ⅰ　不動産")
    _table(doc,
           ["所在", "地番・家屋番号", "地目・種類", "地積・床面積", "持分", "評価額"],
           ["{{行:不動産}}{{行:所在}}", "{{行:地番家屋番号}}", "{{行:地目種別}}",
            "{{行:地積床面積}}", "{{行:持分}}", "{{行:評価額}}"])
    _heading(doc, "Ⅱ　預貯金")
    _table(doc,
           ["金融機関", "支店", "種別", "口座番号", "死亡日残高"],
           ["{{行:預貯金}}{{行:金融機関}}", "{{行:支店}}", "{{行:種別}}",
            "{{行:口座番号}}", "{{行:死亡日残高}}"])
    _heading(doc, "Ⅲ　有価証券・その他")
    _table(doc,
           ["銘柄・内容", "数量", "評価額"],
           ["{{行:有価証券}}{{行:銘柄内容}}", "{{行:数量}}", "{{行:評価額}}"])
    doc.add_paragraph("")

    # ── 消極財産の部 ──
    _heading(doc, "第２　消極財産の部（債務・葬儀費用）")
    _table(doc,
           ["内容", "金額"],
           ["{{行:債務}}{{行:内容}}", "{{行:金額}}"])
    doc.add_paragraph("")

    # ── 合計・注記 ──
    doc.add_paragraph("積極財産合計　{{積極財産合計}}")
    doc.add_paragraph("消極財産合計　{{消極財産合計}}")
    doc.add_paragraph("純資産額　{{純資産額}}")
    doc.add_paragraph("")
    doc.add_paragraph("（注記）")
    doc.add_paragraph("評価基準日: {{評価基準日}}")
    doc.add_paragraph("出典資料: {{出典資料}}")
    return doc


from hub.logging_setup import configure_app_logging  # PR-4b: CLI 起動時の logging 配線


if __name__ == "__main__":
    configure_app_logging()   # __main__ の最初に配線（PR-5: 順序固定）
    OUT.parent.mkdir(parents=True, exist_ok=True)
    build().save(str(OUT))
    logger.info("generated: %s", emit(str(OUT), "freetext", "log", "operator"))
