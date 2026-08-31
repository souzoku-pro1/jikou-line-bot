"""SOUZOKU-HOUKI-H7C: 相続放棄申述書の自動生成。

固定する仕様:
- テンプレの正 = docx_templates/houki/相続放棄申述書.docx（公式様式を
  scripts/make_shinjutsu_template.py でプレースホルダ化・SHA-256 pin・
  プレースホルダ 33 個・全て単一 run・値 run に fitText なし）
- 弁護士決定（改変禁止）: 番号系は丸数字置換（閉じたマッピング・
  1 グループ 1 個のみ・空値グループは無印）・語句系（元号/都道府県）不触・
  「相続財産の概略」欄なしの様式のまま
- 凍結検証: 生成物の全段落（本文+全セル）がテンプレ由来の期待列と完全一致
  （sanctioned substitution=許可した丸数字置換のみ許容）・プレースホルダ
  残存なし
- fail-closed: HOUKI_WEBHOOK_TOKEN 未設定/時効側と同値=404・必須欠落
  （申述人氏名/被相続人氏名/死亡日）・知った日導出不能・マッピング不能=
  生成拒否+管理者通知（閉集合語彙+レコード番号のみ）
- 冪等: 申述書欄 非空=skip（作用 0）・添付 PUT は $revision CAS
  （並行二重 POST の敗者=cas_lost・作用 0）・アップロード失敗=要確認通知+500
"""

import datetime
import hashlib
import os
import unittest
from unittest.mock import AsyncMock, patch

_ENV = {
    "ANTHROPIC_API_KEY": "dummy", "LINE_CHANNEL_SECRET": "dummy_secret",
    "LINE_CHANNEL_ACCESS_TOKEN": "dummy_token", "KINTONE_SUBDOMAIN": "testsub",
    "KINTONE_APP_ID": "21", "KINTONE_API_TOKEN": "dummy",
    "SOUZOKU_KINTONE_APP_ID": "26", "SOUZOKU_KINTONE_API_TOKEN": "dummy",
    "CLOUDSIGN_CLIENT_ID": "c", "CLOUDSIGN_WEBHOOK_SECRET": "cs",
    "KINTONE_WEBHOOK_TOKEN": "kintone-token",
    "DOCUMENT_WEBHOOK_SECRET": "doc-secret",
    "APP_APPROVAL": "29", "TOKEN_APPROVAL": "d", "HEALTHCHECK_DISABLED": "1",
    "STRIPE_WEBHOOK_SECRET": "w", "GOOGLE_VISION_API_KEY": "dummy_vision",
    "APP_CHATLOG": "28", "TOKEN_CHATLOG": "d",
    "APP_HOUKI": "40", "TOKEN_HOUKI": "d",
    "HOUKI_LINE_CHANNEL_SECRET": "houki_secret",
    "HOUKI_LINE_CHANNEL_ACCESS_TOKEN": "houki_token",
    "HOUKI_WEBHOOK_TOKEN": "houki-hook",
}
for _k, _v in _ENV.items():
    os.environ.setdefault(_k, _v)

from docx import Document  # noqa: E402
from docx.table import _Cell  # noqa: E402
from fastapi.testclient import TestClient  # noqa: E402

import main  # noqa: E402
import shinjutsu_webhook as sw  # noqa: E402
from hub import houki_shinjutsu as hs  # noqa: E402
from hub.kintone import KintoneConflict, KintoneError  # noqa: E402

_client = TestClient(main.app)
_URL = "/souzoku-houki/shinjutsu/houki-hook"

TODAY = datetime.date(2026, 8, 31)
_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _rec(**fields):
    return {k: {"value": v} for k, v in fields.items()}


def _check_rows(*pairs):
    return [{"value": {"書類名": {"value": n}, "書類状態": {"value": s}}}
            for n, s in pairs]


def _full_record(**over):
    base = {
        "$revision": "5", "申述書": [],
        "顧客名": "山田花子", "furigana": "ヤマダハナコ",
        "生年月日": "1976-03-13",
        "住所": "〒340-0048 埼玉県草加市原町1-19-47",
        "電話番号": "090-1234-5678",
        "被相続人氏名": "山田太郎", "被相続人ふりがな": "ヤマダタロウ",
        "被相続人本籍": "長野県木曽郡上松町本町通り4-43",
        "被相続人最後の住所": "長野県木曽郡上松町本町通り4-43",
        "死亡日_申告": "2023-11-21",
        "相続の開始を知った日": "2026-06-01",
        "続柄": "子", "続柄その他": "",
        "知った日の区分": "死亡の通知をうけた日", "知った日の区分その他": "",
        "放棄の理由": "債務超過のため。",
        "管轄家庭裁判所": "長野家庭裁判所木曽福島出張所",
        "書類チェック": _check_rows(
            ("被相続人死亡記載戸籍", "受領"), ("申述人戸籍", "受領"),
            ("被相続人住民票除票", "受領")),
    }
    base.update(over)
    return _rec(**base)


def _all_texts(docx_bytes: bytes) -> str:
    import io
    doc = Document(io.BytesIO(docx_bytes))
    return "\n".join(p.text for p in hs._walk_paragraphs(doc))


# ── テンプレ pin ──────────────────────────────────────────────────────────────
class TestTemplatePin(unittest.TestCase):
    def test_sha256_pinned(self):
        data = open(hs.TEMPLATE_PATH, "rb").read()
        self.assertEqual(hashlib.sha256(data).hexdigest(), hs.TEMPLATE_SHA256)

    def test_placeholders_single_run_count_and_no_fittext(self):
        doc = Document(hs.TEMPLATE_PATH)
        n = 0
        for p in hs._walk_paragraphs(doc):
            for r in p.runs:
                if "{{" not in r.text:
                    continue
                self.assertEqual(r.text.count("{{"), r.text.count("}}"),
                                 r.text)
                n += r.text.count("{{")
                rpr = r._r.find(_W + "rPr")
                if rpr is not None:
                    self.assertIsNone(rpr.find(_W + "fitText"), r.text)
        self.assertEqual(n, 33)
        self.assertEqual(set(hs.PLACEHOLDER_DEFAULTS), {
            ph for p in hs._walk_paragraphs(doc) for ph in
            __import__("re").findall(r"\{\{[^{}]+\}\}", p.text)})

    def test_fittext_labels_preserved(self):
        # 収載が fitText ラベル run に触れていない（時効通知書の崩れ前歴の pin）
        import zipfile
        xml = zipfile.ZipFile(hs.TEMPLATE_PATH).read(
            "word/document.xml").decode("utf-8")
        self.assertEqual(xml.count("<w:fitText"), 10)

    def test_empty_cell_insertions_present(self):
        doc = Document(hs.TEMPLATE_PATH)
        t4 = doc.tables[4]
        last = _Cell(t4.rows[7]._tr.tc_lst[2], t4)
        self.assertIn("{{被相続人最後の住所}}",
                      "".join(p.text for p in last.paragraphs))
        dname = _Cell(t4.rows[8]._tr.tc_lst[2], t4)
        joined = "".join(p.text for p in dname.paragraphs)
        self.assertIn("{{被相続人フリガナ}}", joined)
        self.assertIn("{{被相続人氏名}}", joined)

    def test_court_runs_inserted_around_fittext_label(self):
        doc = Document(hs.TEMPLATE_PATH)
        t2 = doc.tables[2]
        p0 = _Cell(t2.rows[0]._tr.tc_lst[0], t2).paragraphs[0]
        self.assertEqual([r.text for r in p0.runs],
                         ["{{裁判所前}}", "家庭裁判", "所", "{{裁判所後}}"])

    def test_circle_digits_unique_in_template_cells(self):
        doc = Document(hs.TEMPLATE_PATH)
        for group, (ti, ri, ci) in hs._CIRCLE_CELLS.items():
            cell = _Cell(doc.tables[ti].rows[ri]._tr.tc_lst[ci],
                         doc.tables[ti])
            text = "".join(p.text for p in cell.paragraphs)
            digits = {"kankei": "１２３４５６７", "shitta": "１２３４",
                      "riyu": "１２３４５６"}[group]
            for d in digits:
                self.assertEqual(text.count(d), 1, (group, d))


# ── fix1[H7C-01]: canonical pin（自己参照でない正本性の防壁） ─────────────────────
class TestCanonicalPin(unittest.TestCase):
    """バイナリ SHA を追随更新する攻撃（テンプレ改変+TEMPLATE_SHA256 更新）を
    canonical 定数（収載スクリプト由来・コード固定）が検出することの固定。"""

    def _tampered_template(self, mutate):
        """コミット済みテンプレへ mutate を適用し、(一時ファイルパス,
        改変後バイナリ SHA) を返す＝攻撃シナリオの「①改変 ②SHA 追随」。"""
        import io as _io
        import tempfile
        doc = Document(hs.TEMPLATE_PATH)
        mutate(doc)
        buf = _io.BytesIO()
        doc.save(buf)
        data = buf.getvalue()
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        tmp.write(data)
        tmp.close()
        self.addCleanup(os.unlink, tmp.name)
        return tmp.name, hashlib.sha256(data).hexdigest()

    def test_canary_committed_template_matches_canonical(self):
        self.assertEqual(hs.canonical_sha256(hs.TEMPLATE_PATH),
                         hs.TEMPLATE_CANONICAL_SHA256)

    def test_canonical_is_zip_timestamp_independent(self):
        # 同一内容を保存し直す（zip タイムスタンプが変わりバイナリ SHA は
        # 変動し得る）→ canonical は不変
        import io as _io
        doc = Document(hs.TEMPLATE_PATH)
        buf = _io.BytesIO()
        doc.save(buf)
        self.assertEqual(hs.canonical_sha256(buf.getvalue()),
                         hs.TEMPLATE_CANONICAL_SHA256)

    def test_fixed_text_one_char_change_detected_despite_sha_update(self):
        # Codex 指摘の攻撃そのもの: 不動文字 1 文字改変+バイナリ SHA 追随でも
        # canonical 照合で拒否される
        def mutate(doc):
            t5 = doc.tables[5]
            cell = _Cell(t5.rows[1]._tr.tc_lst[0], t5)
            run = cell.paragraphs[1].runs[0]
            assert "相　続　の　放　棄" in run.text
            run.text = run.text.replace("放", "抛", 1)
        path, sha = self._tampered_template(mutate)
        with patch.object(hs, "TEMPLATE_PATH", path), \
             patch.object(hs, "TEMPLATE_SHA256", sha):
            with self.assertRaises(hs.ShinjutsuIntegrityError):
                hs.verify_template_integrity()
            with self.assertRaises(hs.ShinjutsuIntegrityError):
                hs.generate(_full_record(), today=TODAY)

    def test_placeholder_count_change_detected_despite_sha_update(self):
        def mutate(doc):
            t4 = doc.tables[4]
            cell = _Cell(t4.rows[2]._tr.tc_lst[2], t4)
            run = cell.paragraphs[0].runs[0]
            assert run.text == "{{申述人フリガナ}}"
            run.text = "　"      # プレースホルダを 1 個削る（個数改変）
        path, sha = self._tampered_template(mutate)
        with patch.object(hs, "TEMPLATE_PATH", path), \
             patch.object(hs, "TEMPLATE_SHA256", sha):
            with self.assertRaises(hs.ShinjutsuIntegrityError):
                hs.verify_template_integrity()

    def test_placeholder_position_change_detected_despite_sha_update(self):
        def mutate(doc):
            t2 = doc.tables[2]
            p0 = _Cell(t2.rows[0]._tr.tc_lst[0], t2).paragraphs[0]
            runs = p0.runs
            assert runs[0].text == "{{裁判所前}}"
            runs[0].text, runs[3].text = runs[3].text, runs[0].text  # 位置交換
        path, sha = self._tampered_template(mutate)
        with patch.object(hs, "TEMPLATE_PATH", path), \
             patch.object(hs, "TEMPLATE_SHA256", sha):
            with self.assertRaises(hs.ShinjutsuIntegrityError):
                hs.verify_template_integrity()

    def test_canonical_constant_anchors_expected_texts(self):
        # 期待列の基準は canonical 定数に錨づく（定数を偽値にすると凍結検証
        # 一式が拒否される=自己生成に退行していないことの pin）
        with patch.object(hs, "TEMPLATE_CANONICAL_SHA256", "0" * 64):
            with self.assertRaises(hs.ShinjutsuIntegrityError):
                hs.generate(_full_record(), today=TODAY)
            fill = dict(hs.PLACEHOLDER_DEFAULTS)
            with self.assertRaises(hs.ShinjutsuIntegrityError):
                hs.expected_paragraph_texts(fill, {})


# ── fix2[H7C-fix1-01]: 表示書式・全パートを含む canonical（見た目改変の遮断） ──────
class TestCanonicalDisplayAttrs(unittest.TestCase):
    """テキスト同一のまま表示を壊す改変（白文字・vanish・フォント・罫線/
    セル書式・styles パート）が、バイナリ SHA 追随でも canonical で拒否
    されることの固定。"""

    def _reject(self, path, sha):
        with patch.object(hs, "TEMPLATE_PATH", path), \
             patch.object(hs, "TEMPLATE_SHA256", sha):
            with self.assertRaises(hs.ShinjutsuIntegrityError):
                hs.verify_template_integrity()

    def _tampered_doc(self, mutate):
        import io as _io
        import tempfile
        doc = Document(hs.TEMPLATE_PATH)
        mutate(doc)
        buf = _io.BytesIO()
        doc.save(buf)
        data = buf.getvalue()
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        tmp.write(data)
        tmp.close()
        self.addCleanup(os.unlink, tmp.name)
        return tmp.name, hashlib.sha256(data).hexdigest()

    def _tampered_part(self, part_name, mutate_text):
        import io as _io
        import tempfile
        import zipfile
        src = zipfile.ZipFile(hs.TEMPLATE_PATH)
        buf = _io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as out:
            for info in src.infolist():
                data = src.read(info.filename)
                if info.filename == part_name:
                    data = mutate_text(data.decode("utf-8")).encode("utf-8")
                out.writestr(info, data)
        data = buf.getvalue()
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        tmp.write(data)
        tmp.close()
        self.addCleanup(os.unlink, tmp.name)
        return tmp.name, hashlib.sha256(data).hexdigest()

    def _shushi_run(self, doc):
        t5 = doc.tables[5]
        run = _Cell(t5.rows[1]._tr.tc_lst[0], t5).paragraphs[1].runs[0]
        assert "相　続　の　放　棄" in run.text
        return run

    def test_white_color_attack_rejected(self):
        from docx.shared import RGBColor

        def mutate(doc):
            self._shushi_run(doc).font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        self._reject(*self._tampered_doc(mutate))

    def test_vanish_attack_rejected(self):
        def mutate(doc):
            self._shushi_run(doc).font.hidden = True
        self._reject(*self._tampered_doc(mutate))

    def test_rfonts_change_rejected(self):
        def mutate(doc):
            self._shushi_run(doc).font.name = "Arial"
        self._reject(*self._tampered_doc(mutate))

    def test_cell_format_change_rejected(self):
        from docx.oxml.ns import qn

        def mutate(doc):
            t5 = doc.tables[5]
            tc = t5.rows[1]._tr.tc_lst[0]
            tc_pr = tc.get_or_add_tcPr()
            shd = tc_pr.makeelement(qn("w:shd"), {})
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:fill"), "000000")
            tc_pr.append(shd)
        self._reject(*self._tampered_doc(mutate))

    def test_styles_part_tamper_rejected(self):
        def mutate(text):
            assert "明朝" in text
            return text.replace("明朝", "ゴシック", 1)
        self._reject(*self._tampered_part("word/styles.xml", mutate))

    def test_theme_part_tamper_rejected(self):
        def mutate(text):
            return text.replace("<a:theme", "<a:theme foo=\"1\"", 1)
        self._reject(*self._tampered_part("word/theme/theme1.xml", mutate))


# ── 生成・凍結・丸数字 ─────────────────────────────────────────────────────────
class TestBuildAndFreeze(unittest.TestCase):
    def test_happy_path_contents(self):
        data = hs.generate(_full_record(), today=TODAY)
        text = _all_texts(data)
        self.assertIn("山田花子", text)
        self.assertIn("ヤマダハナコ", text)
        self.assertIn("山田太郎", text)
        self.assertIn("長野県木曽郡上松町本町通り4-43", text)
        self.assertIn("ヤマダタロウ", text)
        # 丸数字: 子=①・死亡の通知をうけた日=②・債務超過のため。=⑤ の 3 個のみ
        self.assertIn("①", text)
        self.assertIn("②　死亡の通知をうけた日", text)
        self.assertIn("⑤　債務超過のため。", text)
        self.assertEqual(sum(text.count(c) for c in hs._CIRCLED.values()), 3)
        # 非該当番号は原文のまま（例: 続柄の２ 孫・理由の３ 遺産が少ない。）
        self.assertIn("孫　　３", text)
        self.assertIn("３　遺産が少ない。", text)
        # 提出日（2026-08-31=令和８年８月３１日）・死亡日（令和５年）・
        # 知った日（令和８年）・年齢（50）
        self.assertIn("８", text)
        self.assertIn("５", text)
        self.assertIn("（５０歳）", text)
        # 裁判所（前=長野・後=木曽福島出張所）
        self.assertIn("長野家庭裁判所", text.replace("{{", ""))
        self.assertIn("木曽福島出張所", text)
        # 添付書類: 戸籍 2 通 ■・除票 ■・予備行の □ は残る
        self.assertIn("■", text)
        self.assertIn("合計　２　通", text)
        self.assertEqual(text.count("■"), 2)
        self.assertEqual(text.count("□"), 1)
        # プレースホルダ残存なし
        self.assertNotIn("{{", text)

    def test_empty_groups_leave_no_circles(self):
        rec = _full_record(続柄="", 知った日の区分="", 放棄の理由="",
                           管轄家庭裁判所="", 書類チェック=[])
        data = hs.generate(rec, today=TODAY)
        text = _all_texts(data)
        self.assertEqual(sum(text.count(c) for c in hs._CIRCLED.values()), 0)
        self.assertEqual(text.count("□"), 3)      # 添付書類 3 行とも □ のまま
        self.assertNotIn("■", text)

    def test_optional_blanks_keep_form_look(self):
        # 任意項目が空でも生成は止めない（空欄=様式の既定空白のまま）
        rec = _full_record(furigana="", 生年月日="", 電話番号="",
                           被相続人ふりがな="", 被相続人本籍="",
                           被相続人最後の住所="", 住所="")
        data = hs.generate(rec, today=TODAY)
        self.assertNotIn("{{", _all_texts(data))

    def test_sonota_values_only_when_selected(self):
        rec = _full_record(続柄="その他", 続柄その他="伯父",
                           知った日の区分="その他",
                           知った日の区分その他="債権者からの通知")
        text = _all_texts(hs.generate(rec, today=TODAY))
        self.assertIn("⑦", text)
        self.assertIn("（伯父）", text)
        self.assertIn("④", text)
        self.assertIn("債権者からの通知", text)

    def test_tampered_body_rejected(self):
        fill = hs.build_fill_data(_full_record(), today=TODAY)
        circles, _ = hs._circles(_full_record())
        data = hs.build_shinjutsu_docx(fill, circles)
        import io
        doc = Document(io.BytesIO(data))
        doc.paragraphs[4].runs[0].text = "改変"   # (注) 行を改変
        buf = io.BytesIO()
        doc.save(buf)
        with self.assertRaises(hs.ShinjutsuIntegrityError):
            hs.verify_shinjutsu_docx(buf.getvalue(), fill, circles)

    def test_unfilled_placeholder_rejected(self):
        fill = hs.build_fill_data(_full_record(), today=TODAY)
        circles, _ = hs._circles(_full_record())
        broken = dict(fill)
        broken.pop("{{申述人氏名}}")
        data = hs.build_shinjutsu_docx(broken, circles)
        with self.assertRaises(hs.ShinjutsuIntegrityError):
            hs.verify_shinjutsu_docx(data, broken, circles)

    def test_template_hash_mismatch_rejected(self):
        with patch.object(hs, "TEMPLATE_SHA256", "0" * 64):
            with self.assertRaises(hs.ShinjutsuIntegrityError):
                hs.generate(_full_record(), today=TODAY)


class TestRejections(unittest.TestCase):
    def _reasons(self, **over):
        with self.assertRaises(hs.ShinjutsuRejection) as ctx:
            hs.build_fill_data(_full_record(**over), today=TODAY)
        return ctx.exception.reasons

    def test_missing_required(self):
        r = self._reasons(顧客名="", 被相続人氏名="", 死亡日_申告="")
        self.assertIn(hs.REJECT_MISSING_NAME, r)
        self.assertIn(hs.REJECT_MISSING_DECEASED, r)
        self.assertIn(hs.REJECT_MISSING_DEATH, r)

    def test_shitta_missing_and_era(self):
        self.assertIn(hs.REJECT_SHITTA_MISSING,
                      self._reasons(相続の開始を知った日=""))
        self.assertIn(hs.REJECT_SHITTA_ERA,
                      self._reasons(相続の開始を知った日="2019-04-30"))

    def test_death_era_pre_heisei(self):
        self.assertIn(hs.REJECT_DEATH_ERA, self._reasons(死亡日_申告="1988-12-01"))

    def test_unmappable_closed_sets(self):
        self.assertIn(hs.REJECT_KANKEI_UNMAPPED, self._reasons(続柄="いとこ"))
        self.assertIn(hs.REJECT_SHITTA_KUBUN_UNMAPPED,
                      self._reasons(知った日の区分="想定外"))
        self.assertIn(hs.REJECT_RIYU_UNMAPPED,
                      self._reasons(放棄の理由="なんとなく"))

    def test_court_without_marker_rejected(self):
        self.assertIn(hs.REJECT_COURT_UNMAPPED,
                      self._reasons(管轄家庭裁判所="さいたま家裁"))


# ── 受け口 ───────────────────────────────────────────────────────────────────
class _WebhookBase(unittest.TestCase):
    def setUp(self):
        self.get = AsyncMock(return_value=_full_record())
        self.upload = AsyncMock(return_value="fk-1")
        self.update = AsyncMock()
        self.notify = AsyncMock(return_value=True)
        patches = [
            patch.object(sw.hub_kintone, "get_record", self.get),
            patch.object(sw.hub_kintone, "upload_file", self.upload),
            patch.object(sw.hub_kintone, "update_record", self.update),
            patch("hub.notify.notify_admin_line", self.notify),
        ]
        for p in patches:
            p.start()
            self.addCleanup(p.stop)

    def post(self, url=_URL, record_id="12", body=None):
        if body is None:
            body = {"record_id": record_id}
        return _client.post(url, json=body)


class TestWebhookAuth(_WebhookBase):
    def test_token_unset_404(self):
        with patch.dict(os.environ, {"HOUKI_WEBHOOK_TOKEN": ""}):
            resp = self.post()
        self.assertEqual(resp.status_code, 404)
        self.get.assert_not_awaited()

    def test_token_same_as_jikou_404(self):
        for env in ("DOCUMENT_WEBHOOK_SECRET", "KINTONE_WEBHOOK_TOKEN"):
            with self.subTest(env=env):
                with patch.dict(os.environ,
                                {"HOUKI_WEBHOOK_TOKEN": os.environ[env]}):
                    resp = self.post(
                        f"/souzoku-houki/shinjutsu/{os.environ[env]}")
                self.assertEqual(resp.status_code, 404)
        self.get.assert_not_awaited()

    def test_wrong_secret_403(self):
        resp = self.post("/souzoku-houki/shinjutsu/wrong")
        self.assertEqual(resp.status_code, 403)
        self.get.assert_not_awaited()

    def test_app_mismatch_zero_effects(self):
        resp = self.post(body={"app": {"id": "21"},
                               "record": {"$id": {"value": "12"}}})
        self.assertEqual(resp.json().get("skip"), "app_mismatch")
        self.get.assert_not_awaited()

    def test_kintone_webhook_body_shape_accepted(self):
        resp = self.post(body={"app": {"id": "40"},
                               "record": {"$id": {"value": "12"}}})
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(resp.json().get("record_id"), "12")


class TestWebhookFlow(_WebhookBase):
    def test_happy_path_attaches_once(self):
        resp = self.post()
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(resp.json().get("record_id"), "12")
        self.upload.assert_awaited_once()
        fname = self.upload.await_args.args[1]
        self.assertEqual(fname, "相続放棄申述書_12_山田花子.docx")
        self.update.assert_awaited_once()
        args = self.update.await_args
        self.assertEqual(args.args[2],
                         {"申述書": [{"fileKey": "fk-1"}]})
        self.assertEqual(args.kwargs.get("revision"), "5")
        self.notify.assert_not_awaited()

    def test_already_attached_skip_idempotent(self):
        self.get.return_value = _full_record(
            申述書=[{"fileKey": "existing"}])
        resp = self.post()
        self.assertEqual(resp.json().get("skip"), "already_attached")
        self.upload.assert_not_awaited()
        self.update.assert_not_awaited()

    def test_concurrent_loser_cas_lost(self):
        self.update.side_effect = KintoneConflict(409, "GAIA_CO02", "conflict")
        resp = self.post()
        self.assertEqual(resp.json().get("skip"), "cas_lost")
        self.assertEqual(resp.status_code, 200)

    def test_attach_failure_notifies_500(self):
        self.update.side_effect = KintoneError(500, "GAIA_XX", "down")
        resp = self.post()
        self.assertEqual(resp.status_code, 500)
        self.notify.assert_awaited_once()
        text = self.notify.await_args.args[0]
        self.assertIn("要確認", text)
        self.assertIn("No.12", text)
        self.assertNotIn("山田", text)     # PII 非搭載

    def test_rejection_notifies_closed_vocab_only(self):
        self.get.return_value = _full_record(相続の開始を知った日="")
        resp = self.post()
        self.assertEqual(resp.json().get("skip"), "rejected")
        self.upload.assert_not_awaited()
        self.update.assert_not_awaited()
        self.notify.assert_awaited_once()
        text = self.notify.await_args.args[0]
        self.assertIn(hs.REJECT_SHITTA_MISSING, text)
        self.assertIn("No.12", text)
        self.assertNotIn("山田", text)     # PII 非搭載

    def test_no_record_id_skip(self):
        resp = self.post(body={"foo": 1})
        self.assertEqual(resp.json().get("skip"), "no_record_id")
        self.get.assert_not_awaited()


if __name__ == "__main__":
    unittest.main()
