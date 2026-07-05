"""hub/docx_builder.py の単体テスト（T0-3）

検証: 差込（run 分割・表セル含む）・和暦変換・テンプレ規約解決・
プレースホルダ検査・re-export 互換・実テンプレートと検査レジストリの一致。
"""

import io
import os
import tempfile
import unittest
from datetime import date
from pathlib import Path
from unittest.mock import patch

from docx import Document

os.environ.setdefault("KINTONE_SUBDOMAIN", "testsub")

from hub import docx_builder
from hub.docx_builder import (
    TemplateNotFound,
    fill_table_rows,
    fill_template,
    fill_template_with_table,
    list_placeholders,
    resolve_template,
    to_wareki,
    validate_template,
)


def _make_docx(path: Path, paragraphs=(), table_cells=(), split_run=None):
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    if split_run:
        p = doc.add_paragraph()
        for part in split_run:
            p.add_run(part)
    if table_cells:
        table = doc.add_table(rows=1, cols=len(table_cells))
        for i, text in enumerate(table_cells):
            table.rows[0].cells[i].text = text
    doc.save(str(path))


def _all_text(docx_bytes: bytes) -> str:
    doc = Document(io.BytesIO(docx_bytes))
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                parts.extend(p.text for p in c.paragraphs)
    return "\n".join(parts)


class TestFillTemplate(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory()
        self.addCleanup(self.tmp.cleanup)
        self.path = Path(self.tmp.name) / "t.docx"

    def test_replaces_in_paragraphs_and_tables(self):
        _make_docx(self.path,
                   paragraphs=["{{氏名}}様", "日付: {{日付}}"],
                   table_cells=["{{氏名}}", "固定文言"])
        out = fill_template(str(self.path), {"{{氏名}}": "山田太郎", "{{日付}}": "令和8年7月3日"})
        text = _all_text(out)
        self.assertIn("山田太郎様", text)
        self.assertIn("日付: 令和8年7月3日", text)
        self.assertIn("固定文言", text)
        self.assertNotIn("{{", text, "未置換のプレースホルダが残っていない")

    def test_replaces_placeholder_split_across_runs(self):
        """Word 編集で run 分割されたプレースホルダも置換できる（既存実装の要点）"""
        _make_docx(self.path, split_run=["{{氏", "名}}", "様"])
        out = fill_template(str(self.path), {"{{氏名}}": "山田"})
        self.assertIn("山田様", _all_text(out))

    def test_no_match_leaves_document_unchanged(self):
        _make_docx(self.path, paragraphs=["プレースホルダなし"])
        out = fill_template(str(self.path), {"{{氏名}}": "山田"})
        self.assertIn("プレースホルダなし", _all_text(out))

    def test_returns_valid_docx_bytes(self):
        _make_docx(self.path, paragraphs=["{{a}}"])
        out = fill_template(str(self.path), {"{{a}}": "b"})
        self.assertTrue(out.startswith(b"PK"))


class TestToWareki(unittest.TestCase):
    def test_reiwa_first_day(self):
        self.assertEqual(to_wareki(date(2019, 5, 1)), "令和元年5月1日")

    def test_reiwa(self):
        self.assertEqual(to_wareki(date(2026, 7, 3)), "令和8年7月3日")

    def test_heisei_last_day(self):
        self.assertEqual(to_wareki(date(2019, 4, 30)), "平成31年4月30日")

    def test_heisei_first_day(self):
        self.assertEqual(to_wareki(date(1989, 1, 8)), "平成元年1月8日")

    def test_before_heisei_falls_back_to_seireki(self):
        self.assertEqual(to_wareki(date(1989, 1, 7)), "1989年01月07日")


class TestResolveTemplate(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory()
        self.addCleanup(self.tmp.cleanup)
        self.base = Path(self.tmp.name)
        (self.base / "jikou").mkdir()
        _make_docx(self.base / "jikou" / "送付案内.docx", paragraphs=["x"])

    def test_resolves_by_convention(self):
        path = resolve_template("時効援用", "送付案内", base_dir=str(self.base))
        self.assertEqual(path, self.base / "jikou" / "送付案内.docx")

    def test_unknown_unit_raises(self):
        with self.assertRaises(TemplateNotFound):
            resolve_template("存在しないユニット", "送付案内", base_dir=str(self.base))

    def test_missing_file_raises(self):
        with self.assertRaises(TemplateNotFound):
            resolve_template("時効援用", "存在しない種別", base_dir=str(self.base))

    def test_new_unit_needs_only_config_entry(self):
        """拡張ポイント検証: UNIT_CONFIG のエントリ追加だけで新ユニットが解決できる"""
        (self.base / "houki").mkdir()
        _make_docx(self.base / "houki" / "申述書.docx", paragraphs=["y"])
        new_unit = {"相続放棄": {"template_dir": "houki"}}
        with patch.dict(docx_builder.UNIT_CONFIG, new_unit):
            path = resolve_template("相続放棄", "申述書", base_dir=str(self.base))
        self.assertEqual(path, self.base / "houki" / "申述書.docx")


class TestValidateTemplate(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory()
        self.addCleanup(self.tmp.cleanup)
        self.path = Path(self.tmp.name) / "t.docx"

    def test_all_keys_present_returns_empty(self):
        _make_docx(self.path, paragraphs=["{{日付}}"], table_cells=["{{氏名}}"])
        self.assertEqual(validate_template(self.path, ["{{日付}}", "{{氏名}}"]), [])

    def test_missing_keys_are_reported(self):
        _make_docx(self.path, paragraphs=["{{日付}}"])
        missing = validate_template(self.path, ["{{日付}}", "{{氏名}}", "{{住所}}"])
        self.assertEqual(missing, ["{{氏名}}", "{{住所}}"])

    def test_missing_file_raises(self):
        with self.assertRaises(TemplateNotFound):
            validate_template(self.path / "nai.docx", ["{{x}}"])

    def test_list_placeholders(self):
        _make_docx(self.path, paragraphs=["{{b}} と {{a}}"], table_cells=["{{a}}"])
        self.assertEqual(list_placeholders(self.path), ["{{a}}", "{{b}}"])


class TestRealTemplateRegistry(unittest.TestCase):
    """config.EXPECTED_DOCX_TEMPLATES が実テンプレートと一致していること
    （日次死活監視が誤警報しないことの担保）"""

    def test_registered_templates_pass_validation(self):
        from config import EXPECTED_DOCX_TEMPLATES
        for path, keys in EXPECTED_DOCX_TEMPLATES.items():
            with self.subTest(template=path):
                self.assertEqual(validate_template(path, keys), [],
                                 f"{path} に登録済みプレースホルダが欠けている")

    def test_healthcheck_check_templates_is_clean(self):
        import daily_healthcheck
        self.assertEqual(daily_healthcheck.check_templates(), [])


def _add_row_table(doc: Document, header: list[str], template_cells: list[str]):
    """ヘッダ1行＋テンプレート行1行の表を追加する（fill_table_rows テスト用）"""
    table = doc.add_table(rows=2, cols=len(template_cells))
    for i, text in enumerate(header):
        table.rows[0].cells[i].text = text
    for i, text in enumerate(template_cells):
        table.rows[1].cells[i].text = text
    return table


class TestFillTableRows(unittest.TestCase):
    """souzoku-shorui 02 §2（S2）: 行複製・書式継承・空リスト・row_marker"""

    def setUp(self):
        self.doc = Document()

    def test_duplicates_rows_preserving_order(self):
        table = _add_row_table(self.doc, ["No.", "名称"], ["{{行:No}}", "{{行:名称}}"])
        fill_table_rows(self.doc, [{"No": 1, "名称": "土地"},
                                   {"No": 2, "名称": "建物"},
                                   {"No": 3, "名称": "預金"}])
        self.assertEqual(len(table.rows), 1 + 3, "ヘッダ+件数分の行")
        self.assertEqual([r.cells[1].text for r in table.rows[1:]],
                         ["土地", "建物", "預金"], "rows の順序どおり")
        self.assertEqual(table.rows[1].cells[0].text, "1", "非文字列値は str() される")
        all_text = "\n".join(c.text for r in table.rows for c in r.cells)
        self.assertNotIn("{{", all_text, "テンプレート行が残っていない")

    def test_copied_rows_inherit_run_formatting(self):
        """書式継承: テンプレート行の run 書式（太字）が複製行に引き継がれる"""
        table = _add_row_table(self.doc, ["名称"], ["{{行:名称}}"])
        table.rows[1].cells[0].paragraphs[0].runs[0].bold = True
        fill_table_rows(self.doc, [{"名称": "土地"}, {"名称": "建物"}])
        for row in table.rows[1:]:
            run = row.cells[0].paragraphs[0].runs[0]
            self.assertTrue(run.bold, f"複製行 {row.cells[0].text!r} の太字が継承される")

    def test_empty_rows_removes_template_row(self):
        """空リスト既定: テンプレート行を削除するだけ（従来互換・T2-2 送付案内）"""
        table = _add_row_table(self.doc, ["名称"], ["{{行:名称}}"])
        fill_table_rows(self.doc, [])
        self.assertEqual(len(table.rows), 1, "ヘッダのみ残る")
        self.assertEqual(table.rows[0].cells[0].text, "名称")

    def test_empty_rows_with_empty_text_inserts_nashi_row(self):
        """空リスト+empty_text: 書式継承の1行を残し「該当なし」を差し込む（02 §7）"""
        table = _add_row_table(self.doc, ["No.", "名称", "評価額"],
                               ["{{行:No}}", "{{行:名称}}", "{{行:評価額}}"])
        table.rows[1].cells[0].paragraphs[0].runs[0].bold = True
        fill_table_rows(self.doc, [], empty_text="該当なし")
        self.assertEqual(len(table.rows), 2, "ヘッダ+該当なし行")
        self.assertEqual(table.rows[1].cells[0].text, "該当なし",
                         "行内の最初のプレースホルダ位置に empty_text")
        self.assertEqual(table.rows[1].cells[1].text, "")
        self.assertEqual(table.rows[1].cells[2].text, "")
        self.assertTrue(table.rows[1].cells[0].paragraphs[0].runs[0].bold,
                        "該当なし行もテンプレート行の書式を継承")

    def test_row_marker_targets_each_section_table(self):
        """row_marker: 可変表が複数あるテンプレート（財産目録の種別セクション）を
        セクションごとに呼び分けて差し込める"""
        t1 = _add_row_table(self.doc, ["Ⅰ 不動産"],
                            ["{{行:不動産}}{{行:特定情報}}"])
        t2 = _add_row_table(self.doc, ["Ⅱ 預貯金"],
                            ["{{行:預貯金}}{{行:特定情報}}"])
        fill_table_rows(self.doc, [{"特定情報": "所在 川口市"}],
                        row_marker="{{行:不動産}}")
        fill_table_rows(self.doc, [{"特定情報": "○○銀行"}, {"特定情報": "△△銀行"}],
                        row_marker="{{行:預貯金}}")
        self.assertEqual([r.cells[0].text for r in t1.rows],
                         ["Ⅰ 不動産", "所在 川口市"], "マーカーは空文字になる")
        self.assertEqual([r.cells[0].text for r in t2.rows],
                         ["Ⅱ 預貯金", "○○銀行", "△△銀行"])

    def test_row_marker_with_empty_rows_puts_empty_text_at_marker(self):
        _add_row_table(self.doc, ["Ⅳ 動産"], ["{{行:動産}}", "{{行:評価額}}"])
        fill_table_rows(self.doc, [], row_marker="{{行:動産}}", empty_text="該当なし")
        table = self.doc.tables[0]
        self.assertEqual(len(table.rows), 2)
        self.assertEqual(table.rows[1].cells[0].text, "該当なし",
                         "row_marker の位置に empty_text")
        self.assertEqual(table.rows[1].cells[1].text, "")

    def test_row_marker_not_found_leaves_document_unchanged(self):
        """マーカー不在は何もしない（validate_template で検知する規約）"""
        table = _add_row_table(self.doc, ["名称"], ["{{行:名称}}"])
        fill_table_rows(self.doc, [{"名称": "x"}], row_marker="{{行:存在しない}}")
        self.assertEqual(len(table.rows), 2, "テンプレート行が残っている")
        self.assertIn("{{行:名称}}", table.rows[1].cells[0].text)

    def test_marker_split_across_runs_is_replaced(self):
        """Word 編集で run 分割されたプレースホルダも差し込める（fill_template と同じ要点）"""
        table = self.doc.add_table(rows=1, cols=1)
        para = table.rows[0].cells[0].paragraphs[0]
        para.add_run("{{行:名")
        para.add_run("称}}")
        fill_table_rows(self.doc, [{"名称": "土地"}])
        self.assertEqual(table.rows[0].cells[0].text, "土地")

    def test_multiline_value_becomes_word_break(self):
        """値の改行は Word の改行（<w:br/>）になる（_replace_multiline 系の共通挙動）"""
        table = _add_row_table(self.doc, ["特定情報"], ["{{行:特定情報}}"])
        fill_table_rows(self.doc, [{"特定情報": "所在 川口市\n地番 12番3"}])
        para = table.rows[1].cells[0].paragraphs[0]
        self.assertIn("<w:br", para._p.xml)
        self.assertIn("所在 川口市", para.text)
        self.assertIn("地番 12番3", para.text)


class TestFillTemplateWithTableRegression(unittest.TestCase):
    """T2-2 送付案内が使う従来シグネチャ（位置引数・row_marker なし）の回帰"""

    def test_scalar_and_table_fill_together(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "t.docx"
            doc = Document()
            doc.add_paragraph("{{宛先}}様")
            _add_row_table(doc, ["No.", "書類名"], ["{{行:No}}", "{{行:書類名}}"])
            doc.save(str(path))
            out = fill_template_with_table(
                str(path), {"{{宛先}}": "山田太郎"},
                [{"No": "1", "書類名": "委任契約書"}])
        text = _all_text(out)
        self.assertIn("山田太郎様", text)
        self.assertIn("委任契約書", text)
        self.assertNotIn("{{", text)


class TestReExport(unittest.TestCase):
    def test_document_webhook_reexports(self):
        """既存の import 経路（document_webhook.fill_template / to_wareki）が生きていること"""
        import document_webhook
        self.assertIs(document_webhook.fill_template, fill_template)
        self.assertIs(document_webhook.to_wareki, to_wareki)


if __name__ == "__main__":
    unittest.main()
