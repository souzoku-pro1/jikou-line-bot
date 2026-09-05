# -*- coding: utf-8 -*-
"""相続放棄 相談カード（来所用・A4 紙）の雛形を hub/houki_card.CARD_ITEMS から生成する
— HOUKI-CARD-TEMPLATE

成果物:
  - docx（docx_templates/houki/相談カード.docx）: python-docx で生成。zip の時刻を
    固定した決定的バイト列（同一定義→同一 SHA-256・test_houki_card が pin）
  - PDF（印刷用）: reportlab platypus で同じ項目定義から描画（A4 縦・2 ページ以内・
    各ページ下部に「相続放棄 相談カード v1」と ページ番号=読み取り時の版判定）

設計（読み取り精度のための規約）:
  - 各項目に固定の項目番号（1〜N）と見出し。番号→App 40 欄コードは hub/houki_card
    が単一の正
  - 記入欄は枠線付きの空欄。日付は 西暦［　］年［　］月［　］日 の 3 枠。ふりがなは
    氏名欄の上段に小枠
  - 選択肢は横並び・各選択肢の左に □（該当に ✓ または ○ の注記）。文言は
    HEARING_CHOICE_FIELDS の逐語。未成年後見関与は載せない
  - 債権者は 3 行の表（債権者名／住所または連絡先／裁判所からの書類 □あり □なし）
  - 右上に 記入日・受付番号（事務所記入）。末尾に個人情報の注記

使い方:
  python scripts/make_houki_card.py <docx出力先> [<pdf出力先>]
"""

import copy
import hashlib
import io
import os
import sys
import zipfile

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (KeepTogether, Paragraph, SimpleDocTemplate,
                                Spacer, Table, TableStyle)

sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), ".."))
from hub import houki_card as hc  # noqa: E402

_FIXED_ZIP_TIME = (1980, 1, 1, 0, 0, 0)
_FONT_JA = "MS Mincho"
_PDF_FONT = "HeiseiMin-W3"
_BOX = "□"


# ══════════════════════════════════════════════════════════════════════════════
# 共通: 記入欄の文字列
# ══════════════════════════════════════════════════════════════════════════════
def _choice_line(item: hc.CardItem) -> str:
    parts = [f"{_BOX} {c}" for c in item.choices]
    line = "　".join(parts)
    if "続柄その他" in item.fields:
        line += "（　　　　　　　　）"
    return line


def _kana_lines() -> tuple[str, str]:
    return ("ふりがな：", "氏　　名：")


# ══════════════════════════════════════════════════════════════════════════════
# docx
# ══════════════════════════════════════════════════════════════════════════════
def _set_cell_borders(cell, sz: int = 6) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tc_pr.append(borders)


def _run(para, text: str, size: float = 9.5, bold: bool = False):
    r = para.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.name = _FONT_JA
    r._element.rPr.rFonts.set(qn("w:eastAsia"), _FONT_JA)
    return r


def _para(container, text: str = "", size: float = 9.5, bold: bool = False,
          align=None, space_after: float = 2.0):
    p = container.add_paragraph()
    if text:
        _run(p, text, size, bold)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if align is not None:
        p.alignment = align
    return p


def _cell_text(cell, text: str, size: float = 9.5, bold: bool = False) -> None:
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    if text:
        _run(p, text, size, bold)


def _row_height(row, height_mm: float) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    h = OxmlElement("w:trHeight")
    h.set(qn("w:val"), str(int(height_mm * 56.7)))   # twips（1mm≈56.7）
    h.set(qn("w:hRule"), "atLeast")
    tr_pr.append(h)


def _add_page_field(paragraph) -> None:
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)


def _item_rows_docx(table, item: hc.CardItem) -> None:
    """1 項目=1 行（左: 番号+見出し／右: 記入欄）。"""
    row = table.add_row()
    left, right = row.cells
    _set_cell_borders(left)
    _set_cell_borders(right)
    _cell_text(left, f"{item.number}. {item.label}", 8.5, bold=True)
    if item.note and item.kind != "creditors":
        _para(left, f"※{item.note}", 7.5, space_after=0)
    if item.kind == "text":
        _cell_text(right, "")
        _row_height(row, 9)
    elif item.kind == "kana_text":
        kana, name = _kana_lines()
        _cell_text(right, kana, 7.5)
        _para(right, name, 9.5, space_after=0)
        _row_height(row, 13)
    elif item.kind == "date":
        _cell_text(right, hc.DATE_BOX, 10)
        _row_height(row, 9)
    elif item.kind == "choice" or item.kind == "check_only":
        _cell_text(right, _choice_line(item), 9.5)
        _row_height(row, 9)
    elif item.kind == "free":
        _cell_text(right, "")
        for _ in range(item.lines - 1):
            _para(right, "", 9.5, space_after=0)
        _row_height(row, 6 * item.lines + 3)


def _creditor_table_docx(doc, item: hc.CardItem) -> None:
    _para(doc, f"{item.number}. {item.label}　※{item.note}", 8.5, bold=True,
          space_after=1)
    cols = ("No.",) + hc.CREDITOR_COLUMNS
    t = doc.add_table(rows=1, cols=len(cols))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = (Mm(10), Mm(60), Mm(70), Mm(40))
    for i, c in enumerate(cols):
        cell = t.rows[0].cells[i]
        cell.width = widths[i]
        _set_cell_borders(cell)
        _cell_text(cell, c, 8.5, bold=True)
    for n in range(1, hc.CREDITOR_ROWS + 1):
        row = t.add_row()
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]
            _set_cell_borders(cell)
        _cell_text(row.cells[0], str(n), 9)
        _cell_text(row.cells[3], "　".join(f"{_BOX} {c}" for c in hc.CREDITOR_DOC_CHOICES), 9)
        _row_height(row, 8)
    _para(doc, "", 4, space_after=2)


def build_docx() -> bytes:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Mm(210), Mm(297)
    sec.left_margin = sec.right_margin = Mm(14)
    sec.top_margin, sec.bottom_margin = Mm(12), Mm(12)
    normal = doc.styles["Normal"]
    normal.font.name = _FONT_JA
    normal.font.size = Pt(9.5)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), _FONT_JA)

    # ヘッダ: 表題（左）+ 記入日/受付番号（右）
    head = doc.add_table(rows=1, cols=2)
    head.alignment = WD_TABLE_ALIGNMENT.CENTER
    l, r = head.rows[0].cells
    l.width, r.width = Mm(95), Mm(85)
    _cell_text(l, hc.CARD_TITLE, 16, bold=True)
    _para(l, "大野法律事務所（相続放棄専門窓口）", 8.5, space_after=0)
    _set_cell_borders(r)
    _cell_text(r, f"記入日　{hc.DATE_BOX}", 9)
    _para(r, "受付番号（事務所記入）［　　　　　　　］", 9, space_after=0)
    _para(doc, hc.CHECK_NOTE, 8.5, space_after=3)

    for gnum, gtitle in hc.GROUPS:
        _para(doc, f"第{gnum}群　{gtitle}", 10.5, bold=True, space_after=1)
        items = [it for it in hc.CARD_ITEMS if it.group == gnum]
        table = None
        for item in items:
            if item.kind == "creditors":
                table = None
                _creditor_table_docx(doc, item)
                continue
            if table is None:
                table = doc.add_table(rows=0, cols=2)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
            _item_rows_docx(table, item)
            table.rows[-1].cells[0].width = Mm(70)
            table.rows[-1].cells[1].width = Mm(110)
        _para(doc, "", 4, space_after=2)

    _para(doc, hc.PRIVACY_NOTE, 8, space_after=0)

    # フッタ: 版 + ページ番号（読み取り時の版判定）
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(footer, f"{hc.CARD_FOOTER}　—　ページ ", 8)
    _add_page_field(footer)

    out = io.BytesIO()
    doc.save(out)
    return deterministic_zip(out.getvalue())


def deterministic_zip(docx_bytes: bytes) -> bytes:
    src = zipfile.ZipFile(io.BytesIO(docx_bytes))
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as dst:
        for info in src.infolist():
            zi = zipfile.ZipInfo(info.filename, date_time=_FIXED_ZIP_TIME)
            zi.compress_type = zipfile.ZIP_DEFLATED
            zi.external_attr = info.external_attr
            dst.writestr(zi, src.read(info.filename))
    return buf.getvalue()


# ══════════════════════════════════════════════════════════════════════════════
# PDF（印刷用・同じ項目定義から描画）
# ══════════════════════════════════════════════════════════════════════════════
_TTF_CANDIDATES = (
    ("MSMincho", r"C:\Windows\Fonts\msmincho.ttc", 0),
    ("YuMincho", r"C:\Windows\Fonts\yumin.ttf", 0),
)


def _register_pdf_font() -> str:
    """印刷用 PDF のフォント。Windows の TTF（MS 明朝・游明朝）があれば ✓ や
    ダッシュも描ける TTF を使い、無ければ CID フォント（HeiseiMin-W3）に縮退。"""
    for name, path, idx in _TTF_CANDIDATES:
        if name in pdfmetrics.getRegisteredFontNames():
            return name
        if os.path.exists(path):
            try:
                pdfmetrics.registerFont(TTFont(name, path, subfontIndex=idx))
                return name
            except Exception:
                continue
    if _PDF_FONT not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(UnicodeCIDFont(_PDF_FONT))
    return _PDF_FONT


def _pdf_styles():
    font = _register_pdf_font()
    base = ParagraphStyle("base", fontName=font, fontSize=9, leading=12)
    return {
        "title": ParagraphStyle("title", parent=base, fontSize=15, leading=19),
        "small": ParagraphStyle("small", parent=base, fontSize=7.5, leading=10),
        "group": ParagraphStyle("group", parent=base, fontSize=10.5, leading=14),
        "label": ParagraphStyle("label", parent=base, fontSize=8.5, leading=11),
        "field": ParagraphStyle("field", parent=base, fontSize=9.5, leading=12.5),
        "body": base,
    }


def _footer(canvas_, doc_):
    canvas_.saveState()
    canvas_.setFont(_register_pdf_font(), 8)
    canvas_.drawCentredString(A4[0] / 2, 8 * mm,
                              f"{hc.CARD_FOOTER}　—　ページ {doc_.page}")
    canvas_.restoreState()


def build_pdf() -> bytes:
    st = _pdf_styles()
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=14 * mm, rightMargin=14 * mm,
                            topMargin=12 * mm, bottomMargin=14 * mm,
                            title=hc.CARD_FOOTER)
    story = []
    head = Table([[Paragraph(f"<b>{hc.CARD_TITLE}</b>", st["title"]),
                   Paragraph(f"記入日　{hc.DATE_BOX}<br/>受付番号（事務所記入）［　　　　　　　］",
                             st["body"])]],
                 colWidths=[95 * mm, 85 * mm])
    head.setStyle(TableStyle([("BOX", (1, 0), (1, 0), 0.6, colors.black),
                              ("VALIGN", (0, 0), (-1, -1), "TOP")]))
    story += [head, Paragraph("大野法律事務所（相続放棄専門窓口）", st["small"]),
              Paragraph(hc.CHECK_NOTE, st["small"]), Spacer(1, 3 * mm)]
    grid = TableStyle([("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                       ("VALIGN", (0, 0), (-1, -1), "TOP"),
                       ("LEFTPADDING", (0, 0), (-1, -1), 3),
                       ("RIGHTPADDING", (0, 0), (-1, -1), 3)])
    for gnum, gtitle in hc.GROUPS:
        block = [Paragraph(f"<b>第{gnum}群　{gtitle}</b>", st["group"])]
        rows, heights = [], []
        for item in [it for it in hc.CARD_ITEMS if it.group == gnum]:
            if item.kind == "creditors":
                if rows:
                    t = Table(rows, colWidths=[70 * mm, 110 * mm], rowHeights=heights)
                    t.setStyle(grid)
                    block.append(t)
                    rows, heights = [], []
                block.append(Paragraph(f"<b>{item.number}. {item.label}</b>　※{item.note}",
                                       st["label"]))
                cdata = [[Paragraph(c, st["label"]) for c in ("No.",) + hc.CREDITOR_COLUMNS]]
                for n in range(1, hc.CREDITOR_ROWS + 1):
                    cdata.append([str(n), "", "",
                                  "　".join(f"{_BOX} {c}" for c in hc.CREDITOR_DOC_CHOICES)])
                ct = Table(cdata, colWidths=[10 * mm, 60 * mm, 70 * mm, 40 * mm],
                           rowHeights=[6 * mm] + [8 * mm] * hc.CREDITOR_ROWS)
                ct.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                                        ("FONTNAME", (0, 0), (-1, -1), st["body"].fontName),
                                        ("FONTSIZE", (0, 0), (-1, -1), 9),
                                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE")]))
                block.append(ct)
                continue
            label = f"<b>{item.number}. {item.label}</b>"
            if item.note:
                label += f"<br/><font size=7.5>※{item.note}</font>"
            if item.kind == "text":
                field, h = "", 9 * mm
            elif item.kind == "kana_text":
                kana, name = _kana_lines()
                field, h = f"<font size=7.5>{kana}</font><br/>{name}", 13 * mm
            elif item.kind == "date":
                field, h = hc.DATE_BOX, 9 * mm
            elif item.kind in ("choice", "check_only"):
                field, h = _choice_line(item), 9 * mm
            else:
                field, h = "", (6 * item.lines + 3) * mm
            lp, fp = Paragraph(label, st["label"]), Paragraph(field, st["field"])
            rows.append([lp, fp])
            # 見出し文の折返し分は実測して行高を確保（固定高だけだと溢れる）
            _w, lh = lp.wrap(70 * mm - 6, 1000)
            _w, fh = fp.wrap(110 * mm - 6, 1000)
            heights.append(max(h, lh + 4, fh + 4))
        if rows:
            t = Table(rows, colWidths=[70 * mm, 110 * mm], rowHeights=heights)
            t.setStyle(grid)
            block.append(t)
        story += [KeepTogether(block), Spacer(1, 2.5 * mm)]
    story.append(Paragraph(hc.PRIVACY_NOTE, st["small"]))
    doc.build(story, onFirstPage=_footer, onLaterPages=_footer)
    return buf.getvalue()


def main(argv: list[str]) -> int:
    if len(argv) < 2:
        print(__doc__)
        return 2
    docx_bytes = build_docx()
    open(argv[1], "wb").write(docx_bytes)
    print("docx written:", argv[1])
    print("docx sha256:", hashlib.sha256(docx_bytes).hexdigest())
    if len(argv) >= 3:
        open(argv[2], "wb").write(build_pdf())
        print("pdf written:", argv[2])
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
