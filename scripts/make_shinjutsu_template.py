"""相続放棄申述書テンプレートの収載スクリプト（SOUZOKU-HOUKI-H7C）

公式様式（Desktop/claude/相続放棄　書式.docx・SHINJUTSU-TEMPLATE-SURVEY-2 で
実測した未記入のテキストベース様式・SOURCE_SHA256 で同一性を検証）から、
記入欄の空白 run をプレースホルダ化したテンプレートを
docx_templates/houki/相続放棄申述書.docx へ生成する。生成物の SHA-256 を
hub/houki_shinjutsu.py の TEMPLATE_SHA256 に pin する（JIKOU-NOTICE-1 と同型）。

原本はリポジトリに含めない（make_notice_template.py と同じ設計: SOURCE は
Desktop 参照・SOURCE_SHA256 pin で同一性を保証。原本は未記入公式様式で
PII なしだが、収載スクリプト+生成物 SHA pin で再現性は足りる）。

弁護士決定（凍結・改変禁止）:
- マル囲み選択（番号系）は生成時に丸数字置換（１→①等）——テンプレ側は
  不動文字のまま触らない
- 語句系の選択（元号 昭和/平成/令和・本籍の都道府県併記）は触らない
  （印刷後の手書き運用）
- 「相続財産の概略」欄なしのこの様式のまま（欄の追加はしない）

操作の分類（すべて SURVEY-2 の実測 run 構成を前提に位置で特定し、想定と
異なれば SystemExit で停止する。不動文字 run には一切触れない）:
(a) 空白 run の置換（U+3000/半角スペース列 → {{...}}）
(b) run 内の空白部分文字列の置換（『合計　　通』等、不動文字と同居
    する run の空白部だけ置換）
(c) 空セル 3 箇所への run 仕込み（被相続人の最後の住所・被相続人フリガナ/
    氏名。rPr は同表の既存 run から複製）
(d) 宛先裁判所名 run の挿入（fitText ラベル run の前後に新 run。ラベル run
    自体には触れない）
(e) 添付書類の □ run のプレースホルダ化（生成時に □/■ を値として差し込む。
    値なしの既定は □ ＝様式のまま）

canonical pin との照合（fix1[H7C-01]）:
  保存後、生成物の canonical manifest SHA（hub/houki_shinjutsu.
  canonical_sha256——全段落テキスト+run 分割+rPr 要点の決定的直列化・zip
  タイムスタンプ非依存）を hub/houki_shinjutsu.TEMPLATE_CANONICAL_SHA256 と
  照合し、不一致なら exit 1（テンプレの不動文字・プレースホルダ構成が
  pin 済み正本から変わっている）。
  **意図的更新の手順（原本改訂・プレースホルダ設計変更時のみ・票由来）**:
  1. 本 docstring 冒頭の SOURCE_SHA256 を新しい原本の実測値へ更新
  2. 本スクリプトを実行（canonical 不一致で exit 1・新 canonical が出力される）
  3. 出力された canonical を hub/houki_shinjutsu.TEMPLATE_CANONICAL_SHA256 へ、
     コミット現物の SHA-256 を TEMPLATE_SHA256 へ pin し直す
  4. 再実行して一致（exit 0）を確認・全 suite green → commit（Codex レビュー
     対象。canonical 定数の変更は官製様式の不動文言の変更を意味する）

実行: python scripts/make_shinjutsu_template.py
"""

import copy
import hashlib
import sys
from pathlib import Path

from docx import Document
from docx.table import _Cell

REPO = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(REPO))

from hub.houki_shinjutsu import (TEMPLATE_CANONICAL_SHA256,  # noqa: E402
                                 canonical_sha256)
from hub.redact import emit  # noqa: E402

SOURCE = Path.home() / "Desktop" / "claude" / "相続放棄　書式.docx"
SOURCE_SHA256 = (
    "e875e376d6152d89edd3bbd8550637fa8d7784e915959445ee257ac8e0464213")
OUT = REPO / "docx_templates" / "houki" / "相続放棄申述書.docx"

FW = "　"          # 全角スペース
W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def die(msg: str) -> None:
    raise SystemExit(f"収載中止: {msg}")


def cells(table, row_idx: int):
    """gridSpan/vMerge の展開を避け、XML の tc 順でセルを返す
    （SURVEY-2 のセル番地と一致させる）。"""
    tr = table.rows[row_idx]._tr
    return [_Cell(tc, table) for tc in tr.tc_lst]


def replace_run(run, expect: str, new_text: str) -> None:
    """空白 run 全体をプレースホルダへ置換（不動文字 run の誤置換は expect
    完全一致で防ぐ）。"""
    if run.text != expect:
        die(f"run 現物が想定と異なります: {run.text!r} != {expect!r}")
    run.text = new_text


def replace_in_run(run, old: str, new: str) -> None:
    """run 内の空白部分文字列だけを置換（同居する不動文字は保持）。"""
    if run.text.count(old) != 1:
        die(f"run 内の置換対象が一意でありません: {old!r} in {run.text!r}")
    run.text = run.text.replace(old, new)


def new_run_from(paragraph, rpr_source_run, text: str,
                 insert_before_run=None, with_break: bool = False):
    """rPr を複製した新 run を段落へ挿入する（fitText 付き run を複製元に
    しない——値 run へ均等割り付けを持ち込まない）。"""
    src_rpr = rpr_source_run._r.find(W_NS + "rPr")
    if src_rpr is not None and src_rpr.find(W_NS + "fitText") is not None:
        die("fitText 付き run を複製元にはできません")
    r = paragraph.add_run("")._r
    if src_rpr is not None:
        r.insert(0, copy.deepcopy(src_rpr))
    if with_break:
        br = r.makeelement(W_NS + "br", {})
        r.append(br)
    t = r.makeelement(W_NS + "t", {})
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    if insert_before_run is not None:
        insert_before_run._r.addprevious(r)
    return r


def main() -> None:
    if not SOURCE.exists():
        die(f"原本が見つかりません: {SOURCE}")
    data = SOURCE.read_bytes()
    if hashlib.sha256(data).hexdigest() != SOURCE_SHA256:
        die("原本の SHA-256 が SURVEY-2 実測値と一致しません")

    doc = Document(str(SOURCE))
    tables = doc.tables
    if len(tables) != 7:
        die(f"テーブル数が想定（7）と異なります: {len(tables)}")

    # ── TABLE2: 宛先裁判所+提出日 / 記名押印 ──────────────────────────────────
    t2 = tables[2]
    c = cells(t2, 0)[0]
    p0 = c.paragraphs[0]
    if [r.text for r in p0.runs] != ["家庭裁判", "所"]:
        die("宛先裁判所行の run 構成が想定と異なります")
    # (d) 裁判所名の前後 run 挿入（rPr は日付行『令和』run から複製=fitText なし）
    date_p = c.paragraphs[3]
    if date_p.runs[0].text != "令和":
        die("提出日行の run 構成が想定と異なります")
    new_run_from(p0, date_p.runs[0], "{{裁判所前}}",
                 insert_before_run=p0.runs[0])
    new_run_from(p0, date_p.runs[0], "{{裁判所後}}")
    # (a) 提出日: 年/月/日の '　　' 3 run
    blanks = [r for r in date_p.runs if r.text == FW * 2]
    if len(blanks) != 3:
        die("提出日行の空白 run が 3 個でありません")
    for r, ph in zip(blanks, ("{{提出年}}", "{{提出月}}", "{{提出日}}")):
        replace_run(r, FW * 2, ph)
    # (a) 記名押印: 14 全角空白 run
    c2 = cells(t2, 0)[2]
    p1 = c2.paragraphs[1]
    targets = [r for r in p1.runs if r.text == FW * 14]
    if len(targets) != 1:
        die("記名押印欄の空白 run が特定できません")
    replace_run(targets[0], FW * 14, "{{記名}}")

    # ── TABLE3: 添付書類（□・通数） ─────────────────────────────────────────
    t3 = tables[3]
    c = cells(t3, 0)[1]
    ps = c.paragraphs
    replace_run(ps[1].runs[0], "□", "{{チェック戸籍}}")           # (e)
    goukei = [r for r in ps[1].runs if "合計" in r.text]
    if len(goukei) != 1:
        die("戸籍通数の run が特定できません")
    replace_in_run(goukei[0], f"合計{FW * 2}通", "合計{{戸籍通数}}通")   # (b)
    replace_run(ps[2].runs[0], "□", "{{チェック除票}}")           # (e)
    # 3 つ目の □（ラベルなし予備行）は不動のまま

    # ── TABLE4: 当事者欄 ────────────────────────────────────────────────────
    t4 = tables[4]
    # 申述人 住所（row1 col2）
    addr = cells(t4, 1)[2]
    p0 = addr.paragraphs[0]
    runs = p0.runs
    dash = [i for i, r in enumerate(runs) if r.text == "－"]
    if len(dash) != 1:
        die("〒 行の－ run が特定できません")
    replace_run(runs[dash[0] - 2], FW * 2, "{{郵便前}}")
    replace_run(runs[dash[0] + 1], FW * 2, "{{郵便後}}")
    tel = [i for i, r in enumerate(runs) if r.text == "電話" + FW]
    if len(tel) != 1:
        die("電話ラベル run が特定できません")
    replace_run(runs[tel[0] + 1], FW, "{{電話1}}")
    paren = [i for i, r in enumerate(runs) if r.text == FW * 4]
    if len(paren) != 1:
        die("電話（）内の空白 run が特定できません")
    replace_run(runs[paren[0]], FW * 4, "{{電話2}}")
    if runs[-1].text != " " * 4:
        die("電話末尾の空白 run が想定と異なります")
    replace_run(runs[-1], " " * 4, "{{電話3}}")
    p1 = addr.paragraphs[1]
    replace_run(p1.runs[0], FW * 23, "{{申述人住所}}")

    # 申述人 フリガナ/氏名（row2 col2・1 段落に小/大 2 run）
    name = cells(t4, 2)[2]
    p0 = name.paragraphs[0]
    if [r.text for r in p0.runs] != [FW, FW * 2]:
        die("申述人氏名セルの run 構成が想定と異なります")
    replace_run(p0.runs[0], FW, "{{申述人フリガナ}}")
    replace_run(p0.runs[1], FW * 2, "{{申述人氏名}}")

    # 申述人 生年月日（row2 col3）: 数字部のみ。元号は語句系=不触（弁護士決定）
    birth = cells(t4, 2)[3]
    bp1 = birth.paragraphs[1]
    y_blanks = [r for r in bp1.runs if r.text == FW]
    if len(y_blanks) != 4:
        die("生年月日行の空白 run 構成が想定と異なります")
    replace_run(y_blanks[0], FW, "{{生年}}")     # 年の先頭空白
    replace_run(y_blanks[3], FW, "{{生月}}")     # 月の直前空白
    day = [r for r in bp1.runs if r.text == f" {FW} "]
    if len(day) != 1:
        die("生日 run が特定できません")
    replace_in_run(day[0], FW, "{{生日}}")
    bp2 = birth.paragraphs[2]
    replace_in_run(bp2.runs[0], f"（{FW * 4}歳）", "（{{年齢}}歳）")

    # 続柄 ７その他（　）（row3 col2）
    zoku = cells(t4, 3)[2]
    zp2 = zoku.paragraphs[2]
    replace_in_run(zp2.runs[1], f"（{FW * 9}）", "（{{続柄その他}}）")

    # 被相続人 本籍（row6 col2）: 県 run の後ろへ追記 run（(d) と同じ複製方式）
    honseki = cells(t4, 6)[2]
    hp1 = honseki.paragraphs[1]
    if hp1.runs[-1].text != "県":
        die("被相続人本籍行の run 構成が想定と異なります")
    new_run_from(hp1, hp1.runs[0], "{{被相続人本籍}}")

    # 被相続人 最後の住所（row7 col2・空セル）: (c) run 仕込み
    last_addr = cells(t4, 7)[2]
    if any(r.text for p in last_addr.paragraphs for r in p.runs):
        die("被相続人最後の住所セルが空でありません")
    new_run_from(last_addr.paragraphs[0], hp1.runs[0],
                 "{{被相続人最後の住所}}")

    # 被相続人 フリガナ/氏名（row8 col2・空セル）: (c) 小/大 2 run+改行
    dname = cells(t4, 8)[2]
    if any(r.text for p in dname.paragraphs for r in p.runs):
        die("被相続人氏名セルが空でありません")
    applicant_name = cells(t4, 2)[2].paragraphs[0]
    small_src, large_src = applicant_name.runs[0], applicant_name.runs[1]
    new_run_from(dname.paragraphs[0], small_src, "{{被相続人フリガナ}}")
    new_run_from(dname.paragraphs[0], large_src, "{{被相続人氏名}}",
                 with_break=True)

    # 被相続人 死亡日（row8 col3）: run 内の年月日空白のみ
    death = cells(t4, 8)[3]
    dp0 = death.paragraphs[0]
    tgt = [r for r in dp0.runs if "日死亡" in r.text]
    if len(tgt) != 1:
        die("死亡日 run が特定できません")
    replace_in_run(tgt[0], f"{FW * 3}年{FW * 3}月{FW * 3}日死亡",
                   "{{死亡年}}年{{死亡月}}月{{死亡日}}日死亡")

    # ── TABLE6: 申述の理由 ──────────────────────────────────────────────────
    t6 = tables[6]
    shitta = cells(t6, 1)[0]
    sp1 = shitta.paragraphs[1]
    replace_in_run(sp1.runs[0], f"令和{FW * 3}年{FW * 3}月{FW * 3}日",
                   "令和{{知年}}年{{知月}}月{{知日}}日")
    sp4 = shitta.paragraphs[4]
    others = [r for r in sp4.runs if "その他（" in r.text]
    if len(others) != 1:
        die("知った日その他 run が特定できません")
    replace_in_run(others[0], FW * 10, "{{知った日区分その他}}")
    # 放棄の理由（row3）の番号・６その他は不動のまま（丸数字置換は生成時）

    # ── 収載後の機械検査 ────────────────────────────────────────────────────
    def walk_paragraphs(d):
        yield from d.paragraphs
        for tb in d.tables:
            for row in tb.rows:
                for tc in row._tr.tc_lst:
                    yield from _Cell(tc, tb).paragraphs

    n_ph = 0
    for p in walk_paragraphs(doc):
        if "{{" not in p.text:
            continue
        for r in p.runs:
            if "{{" not in r.text:
                continue
            if r.text.count("{{") != r.text.count("}}"):
                die(f"プレースホルダが run を跨いでいます: {p.text!r}")
            n_ph += r.text.count("{{")
            rpr = r._r.find(W_NS + "rPr")
            if rpr is not None and rpr.find(W_NS + "fitText") is not None:
                die(f"値 run に fitText が混入しています: {r.text!r}")
    if n_ph != 33:
        die(f"プレースホルダ総数が想定（33）と異なります: {n_ph}")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    sha = hashlib.sha256(OUT.read_bytes()).hexdigest()
    # sink 規律（test_sink_ast_policy）: print 引数は定数+emit() のみ。
    # 保存先・個数は上の検査で定数確定済み。SHA-256 hex64 は record_id の
    # 値域（[A-Za-z0-9_-]{1,64}）を満たす素通し kind で出力する
    print("saved: docx_templates/houki/相続放棄申述書.docx")
    print("placeholders: 33")
    print("TEMPLATE_SHA256 =", emit(sha, "record_id", "log", "operator"))

    # fix1[H7C-01]: 生成物の canonical を pin 済み定数と照合（不一致=exit 1。
    # 意図的更新の手順は docstring 参照）
    canon = canonical_sha256(str(OUT))
    print("CANONICAL_SHA256 =", emit(canon, "record_id", "log", "operator"))
    if canon != TEMPLATE_CANONICAL_SHA256:
        die("生成物の canonical が pin（hub/houki_shinjutsu."
            "TEMPLATE_CANONICAL_SHA256）と一致しません。原本改訂等の意図的"
            "更新なら docstring の手順で pin を更新してください")
    print("canonical: OK (matches pinned TEMPLATE_CANONICAL_SHA256)")


if __name__ == "__main__":
    main()
