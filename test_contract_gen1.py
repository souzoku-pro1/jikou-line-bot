"""CONTRACT-GEN-1(+fix1): 時効援用委任契約書の自動生成。

固定する仕様:
- 入口ガード（fix1[01]）: webhook 本文 app.id が App21 実 ID と完全一致
  （欠落・非数字・別 App は get_record 含め作用 0）。正本再検証は完全一致
  （作成のみ続行・作成済=already_done・要確認/空/他値=stale_status で作用 0）。
- CAS 状態機械（fix1[02]）: 作成→作成中 を $revision CAS で確保した勝者のみ
  生成・upload（並行 2 本でも合計 1 回）。完了 PUT は revision=claim+1。
  作成中+添付なし=回収（再claim→再生成）／作成中+添付あり=自動上書きせず
  「要確認」+管理者通知／作成済への再配送=upload 増分 0。
- 凍結文言の構造保証（fix1[03]）: テンプレ SHA-256 固定 pin＋報酬条項
  （第2条全体）の正規化済み逐語一致 pin（テンプレ・生成物の両方）。
  改変生成物は添付せず 500。
- fail-closed: 顧客名/住所/債権者 1 社以上の欠落は生成せず不足フィールド名
  のみ通知（値は非搭載・状態も動かさない）。
"""

import hashlib
import io
import os
import unittest
from unittest.mock import AsyncMock, patch

_ENV = {
    "ANTHROPIC_API_KEY": "dummy", "LINE_CHANNEL_SECRET": "dummy_secret",
    "LINE_CHANNEL_ACCESS_TOKEN": "dummy_token", "KINTONE_SUBDOMAIN": "testsub",
    "KINTONE_APP_ID": "21", "KINTONE_API_TOKEN": "dummy",
    "SOUZOKU_KINTONE_APP_ID": "26", "SOUZOKU_KINTONE_API_TOKEN": "dummy",
    "CLOUDSIGN_CLIENT_ID": "c", "CLOUDSIGN_WEBHOOK_SECRET": "cs",
    "KINTONE_WEBHOOK_TOKEN": "kintone-token",
    "DOCUMENT_WEBHOOK_SECRET": "doc-secret",
    "APP_APPROVAL": "29", "TOKEN_APPROVAL": "d", "HEALTHCHECK_DISABLED": "1",
    "STRIPE_WEBHOOK_SECRET": "w", "GOOGLE_VISION_API_KEY": "dummy_vision",
}
for _k, _v in _ENV.items():
    os.environ.setdefault(_k, _v)

from docx import Document  # noqa: E402
from fastapi.testclient import TestClient  # noqa: E402

import contract_webhook as cw  # noqa: E402
import main  # noqa: E402
from config import EXPECTED_DOCX_TEMPLATES  # noqa: E402
from hub.kintone import KintoneError  # noqa: E402

_client = TestClient(main.app)
_URL = "/contract/doc-secret"


def _rec(**fields):
    return {k: {"value": v} for k, v in fields.items()}


def _full_record(**over):
    base = {"$revision": "5", "契約書ステータス": "契約書作成",
            "顧客名": "熊澤花子", "住所": "埼玉県川口市青木1-1-1",
            "問い合わせ業者名": "株式会社Aファイナンス",
            "対象債権者2": "", "対象債権者3": "",
            "委任契約書": []}
    base.update(over)
    return _rec(**base)


def _body(record_id="12", status="契約書作成", app_id="21"):
    body = {"record": {"$id": {"value": record_id},
                       "契約書ステータス": {"value": status}}}
    if app_id is not None:
        body["app"] = {"id": app_id}
    return body


def _docx_text(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    return "\n".join(p.text for p in doc.paragraphs)


def _tampered_docx() -> bytes:
    """報酬条項を改変した生成物（fix1[03] negative 用）。"""
    doc = Document(cw.TEMPLATE_PATH)
    for p in doc.paragraphs:
        if "44,000" in p.text and p.runs:
            p.runs[0].text = p.text.replace("44,000", "40,000")
            for r in p.runs[1:]:
                r.text = ""
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestTemplateIntegrity(unittest.TestCase):
    def test_template_sha256_pinned(self):
        # fix1[03]: 人承認済み現物の SHA-256 固定 pin
        data = open(cw.TEMPLATE_PATH, "rb").read()
        self.assertEqual(hashlib.sha256(data).hexdigest(),
                         cw.TEMPLATE_SHA256)

    def test_frozen_clause_verbatim_in_template(self):
        # fix1[03]: 報酬条項（第2条全体）の正規化済み逐語一致（44,000 の
        # 部分文字列確認を置き換える強化 pin）
        data = open(cw.TEMPLATE_PATH, "rb").read()
        self.assertEqual(cw._clause_of(data), cw.FROZEN_CLAUSE)

    def test_registered_keys_eight_unique(self):
        keys = EXPECTED_DOCX_TEMPLATES["docx_templates/jikou/委任契約書.docx"]
        self.assertEqual(len(keys), 8)          # 一意プレースホルダ 8 キー
        text = _docx_text(open(cw.TEMPLATE_PATH, "rb").read())
        for k in keys:
            self.assertIn(k, text)
        self.assertEqual(text.count("{{依頼者氏名}}"), 2)


class _WebhookBase(unittest.TestCase):
    def _post(self, *, record, body=None, url=_URL,
              upload=None, update=None):
        upload = upload or AsyncMock(return_value="fk-1")
        update = update or AsyncMock()
        notify = AsyncMock(return_value=True)
        get = AsyncMock(return_value=record)
        with patch.dict(os.environ, _ENV), \
             patch.object(cw.hub_kintone, "get_record", get), \
             patch.object(cw.hub_kintone, "upload_file", upload), \
             patch.object(cw.hub_kintone, "update_record", update), \
             patch("hub.notify.notify_admin_line", notify):
            r = _client.post(url, json=body or _body())
        return r, upload, update, notify, get


class TestEntryGuards(_WebhookBase):
    def test_wrong_secret_403(self):
        r, *_ = self._post(record=_full_record(), url="/contract/wrong")
        self.assertEqual(r.status_code, 403)

    def test_app_mismatch_zero_effects(self):
        # fix1[01]: 別 App・欠落・非数字は get_record 含め作用 0
        for app_id in ("26", None, "abc"):
            with self.subTest(app_id=app_id):
                r, upload, update, _n, get = self._post(
                    record=_full_record(), body=_body(app_id=app_id))
                self.assertEqual(r.json().get("skip"), "app_mismatch")
                get.assert_not_awaited()
                upload.assert_not_awaited()
                update.assert_not_awaited()

    def test_not_triggered_body_gate(self):
        # 自 update の echo（作成中/作成済）も本文 gate で作用 0
        for status in ("契約書作成中", "契約書作成済", "要確認", ""):
            with self.subTest(status=status):
                r, upload, update, _n, get = self._post(
                    record=_full_record(), body=_body(status=status))
                self.assertEqual(r.json().get("skip"), "not_triggered")
                get.assert_not_awaited()
                upload.assert_not_awaited()

    def test_stale_body_vs_authoritative_states(self):
        # fix1[01]: stale 本文（作成）×最新正本 3 態様 → 作用 0
        cases = {
            "契約書作成済": "already_done",
            "要確認": "stale_status",
            "": "stale_status",
        }
        for state, skip in cases.items():
            with self.subTest(state=state):
                r, upload, update, _n, _g = self._post(
                    record=_full_record(契約書ステータス=state))
                self.assertEqual(r.json().get("skip"), skip)
                upload.assert_not_awaited()
                update.assert_not_awaited()


class TestStateMachine(_WebhookBase):
    def test_happy_path_cas_then_attach(self):
        r, upload, update, notify, _g = self._post(record=_full_record())
        self.assertEqual(r.status_code, 200)
        self.assertEqual(r.json().get("record_id"), "12")
        # CONTRACT-GEN-1-fix1[02] 裁定由来: 1 PUT → CAS+完了 PUT の 2 段へ
        self.assertEqual(update.await_count, 2)
        cas_call, final_call = update.await_args_list
        self.assertEqual(cas_call.args[2], {"契約書ステータス": "契約書作成中"})
        self.assertEqual(cas_call.kwargs.get("revision"), "5")
        self.assertEqual(final_call.args[2]["契約書ステータス"], "契約書作成済")
        self.assertEqual(final_call.args[2]["委任契約書"],
                         [{"fileKey": "fk-1"}])
        self.assertEqual(final_call.kwargs.get("revision"), "6")  # claim+1
        # 生成物: 差し込み+凍結条項の逐語一致（fix1[03]）
        text = _docx_text(upload.await_args.args[2])
        self.assertEqual(text.count("熊澤花子"), 2)
        self.assertIn("株式会社Aファイナンス", text)
        self.assertNotIn("{{", text)
        self.assertEqual(cw._clause_of(upload.await_args.args[2]),
                         cw.FROZEN_CLAUSE)
        notify.assert_not_awaited()

    def test_concurrent_second_request_loses_cas(self):
        # fix1[02] negative: 並行 2 本 → 生成/upload 合計 1 回・敗者は作用 0
        r1, upload1, update1, _n, _g = self._post(record=_full_record())
        self.assertEqual(r1.json().get("record_id"), "12")
        self.assertEqual(upload1.await_count, 1)
        loser_update = AsyncMock(side_effect=KintoneError(409, "GAIA_CO02"))
        r2, upload2, update2, _n2, _g2 = self._post(
            record=_full_record(), update=loser_update)
        self.assertEqual(r2.json().get("skip"), "cas_lost")
        upload2.assert_not_awaited()             # 敗者は生成・upload 0 回
        self.assertEqual(update2.await_count, 1)  # 失敗した CAS のみ

    def test_recovery_after_crash_before_attach(self):
        # fix1[02] negative: upload 後 update 前に停止 → 作成中+添付なし →
        # 再配送で回収（再claim→再生成→作成済）
        record = _full_record(契約書ステータス="契約書作成中")
        record["$revision"] = {"value": "7"}
        r, upload, update, notify, _g = self._post(record=record)
        self.assertEqual(r.status_code, 200)
        self.assertTrue(r.json().get("recovered"))
        self.assertEqual(upload.await_count, 1)
        cas_call, final_call = update.await_args_list
        self.assertEqual(cas_call.args[2], {"契約書ステータス": "契約書作成中"})
        self.assertEqual(cas_call.kwargs.get("revision"), "7")
        self.assertEqual(final_call.kwargs.get("revision"), "8")
        self.assertEqual(final_call.args[2]["契約書ステータス"], "契約書作成済")
        notify.assert_not_awaited()

    def test_redelivery_after_done_uploads_nothing(self):
        # fix1[02] negative: update 後応答喪失 → 再配送は already_done で
        # upload 増分 0
        r, upload, update, _n, _g = self._post(
            record=_full_record(契約書ステータス="契約書作成済"))
        self.assertEqual(r.json().get("skip"), "already_done")
        upload.assert_not_awaited()
        update.assert_not_awaited()

    def test_working_with_attachment_goes_review_no_overwrite(self):
        # fix1[02] negative: 作成中+添付あり → 整合確認不能=自動上書きせず
        # 「要確認」+管理者通知・upload 0
        record = _full_record(契約書ステータス="契約書作成中")
        record["委任契約書"] = {"value": [{"fileKey": "old-file"}]}
        r, upload, update, notify, _g = self._post(record=record)
        self.assertEqual(r.json().get("skip"), "needs_review")
        upload.assert_not_awaited()
        self.assertEqual(update.await_count, 1)
        self.assertEqual(update.await_args.args[2],
                         {"契約書ステータス": "要確認"})
        notify.assert_awaited_once()
        self.assertIn("自動では上書きせず", notify.await_args.args[0])

    def test_claim_failures_not_silenced(self):
        # fix2（CONTRACT-GEN-04）: 409 競合のみ cas_lost（200/作用0）。
        # 5xx・transport_error・401/403 は HTTP 500 で kintone 再配送へ
        # （cas_lost に変換して沈黙させない）
        cases = {
            "conflict_409": (KintoneError(409, "GAIA_CO02", "競合"),
                             200, "cas_lost"),
            "server_500": (KintoneError(500, "GAIA_XX01", "x"), 500, None),
            "transport": (KintoneError(0, "transport_error", "x"), 500, None),
            "auth_401": (KintoneError(401, "CB_AU01", "x"), 500, None),
            "auth_403": (KintoneError(403, "GAIA_AP15", "x"), 500, None),
        }
        for label, (err, want_status, want_skip) in cases.items():
            with self.subTest(case=label):
                update = AsyncMock(side_effect=err)
                r, upload, _u, _n, _g = self._post(
                    record=_full_record(), update=update)
                self.assertEqual(r.status_code, want_status)
                if want_skip:
                    self.assertEqual(r.json().get("skip"), want_skip)
                upload.assert_not_awaited()      # いずれも生成/upload 0

    def test_review_transition_only_conflict_is_cas_lost(self):
        # fix2: 要確認への CAS でも 409 のみ cas_lost・障害系は 500
        record = _full_record(契約書ステータス="契約書作成中")
        record["委任契約書"] = {"value": [{"fileKey": "old-file"}]}
        for err, want in ((KintoneError(409, "GAIA_CO02", "x"), 200),
                          (KintoneError(500, "GAIA_XX01", "x"), 500),
                          (KintoneError(0, "transport_error", "x"), 500)):
            with self.subTest(status=err.status):
                update = AsyncMock(side_effect=err)
                r, upload, _u, notify, _g = self._post(
                    record=record, update=update)
                self.assertEqual(r.status_code, want)
                upload.assert_not_awaited()
                notify.assert_not_awaited()      # 遷移未成立時は通知しない

    def test_recovery_claim_failure_not_silenced(self):
        # fix2: 回収の再claim も同規則（500 系は HTTP 500）
        record = _full_record(契約書ステータス="契約書作成中")
        record["$revision"] = {"value": "7"}
        update = AsyncMock(side_effect=KintoneError(503, "GAIA_XX02", "x"))
        r, upload, _u, _n, _g = self._post(record=record, update=update)
        self.assertEqual(r.status_code, 500)
        upload.assert_not_awaited()

    def test_review_transition_cas_lost_no_notify(self):
        record = _full_record(契約書ステータス="契約書作成中")
        record["委任契約書"] = {"value": [{"fileKey": "old-file"}]}
        loser = AsyncMock(side_effect=KintoneError(409, "GAIA_CO02"))
        r, upload, _u, notify, _g = self._post(record=record, update=loser)
        self.assertEqual(r.json().get("skip"), "cas_lost")
        upload.assert_not_awaited()
        notify.assert_not_awaited()


class TestFailClosed(_WebhookBase):
    def test_missing_fields_rejected(self):
        cases = {
            "住所欠落": _full_record(住所=""),
            "氏名欠落": _full_record(顧客名="  "),
            "債権者全欠落": _full_record(問い合わせ業者名="",
                                         対象債権者2="", 対象債権者3=""),
        }
        for label, record in cases.items():
            with self.subTest(case=label):
                r, upload, update, notify, _g = self._post(record=record)
                self.assertEqual(r.json().get("skip"), "missing_fields")
                upload.assert_not_awaited()
                update.assert_not_awaited()      # 状態も動かさない
                notify.assert_awaited_once()
                sent = notify.await_args.args[0]
                self.assertIn("必須項目が未入力", sent)
                self.assertNotIn("熊澤", sent)   # 値は通知に載せない

    def test_upload_failure_500_stays_working(self):
        # fix1[02] 裁定由来: CAS 後の失敗は 作成中 のまま（回収経路で復旧）
        upload = AsyncMock(side_effect=RuntimeError("kintone down"))
        r, _u, update, _n, _g = self._post(record=_full_record(),
                                           upload=upload)
        self.assertEqual(r.status_code, 500)
        self.assertEqual(update.await_count, 1)  # CAS のみ・作成済へは進めない
        self.assertEqual(update.await_args.args[2],
                         {"契約書ステータス": "契約書作成中"})

    def test_tampered_output_rejected(self):
        # fix1[03]: 報酬条項が改変された生成物は添付せず 500
        with patch.object(cw, "fill_template",
                          lambda *_a, **_k: _tampered_docx()):
            r, upload, update, _n, _g = self._post(record=_full_record())
        self.assertEqual(r.status_code, 500)
        upload.assert_not_awaited()
        self.assertEqual(update.await_count, 1)  # CAS のみ

    def test_template_hash_mismatch_rejected(self):
        with patch.object(cw, "TEMPLATE_SHA256", "0" * 64):
            r, upload, _u, _n, _g = self._post(record=_full_record())
        self.assertEqual(r.status_code, 500)
        upload.assert_not_awaited()


class TestFillData(unittest.TestCase):
    def test_blank_slots_are_fullwidth_space(self):
        data = cw.build_fill_data(_full_record())
        for key in ("{{対象債権者2}}", "{{対象債権者3}}",
                    "{{契約年}}", "{{契約月}}", "{{契約日}}"):
            self.assertEqual(data[key], "　")

    def test_missing_detection_closed(self):
        self.assertEqual(cw._missing_fields(_full_record()), [])
        self.assertEqual(cw._missing_fields(_full_record(顧客名="")),
                         ["顧客名"])
        missing = cw._missing_fields(_rec(**{
            "顧客名": "", "住所": "", "問い合わせ業者名": "",
            "対象債権者2": "", "対象債権者3": ""}))
        self.assertEqual(len(missing), 3)


if __name__ == "__main__":
    unittest.main()
