"""contract_pdf — 生成済み docx の段落を PDF に描画（CONTRACT-GEN-2）

設計（票の条件「テンプレ docx を単一の正・本文をコードに二重管理しない」）:
  - 本文文言はこのモジュールに一切持たない。fill_template 済み docx の段落
    テキストをそのまま描画する（＝文言の正はテンプレ docx ただ一つ）。
  - フォントは reportlab 内蔵 CID フォント（HeiseiMin-W3）。フォントファイル
    への依存なし＝コンテナに追加インストール不要で和文を描画できる。
  - 体裁は「契約書として読める」水準（A4 縦・タイトル中央・docx の段落
    alignment を反映・文字単位の折返し）。Word と同一レイアウトは目標に
    しない（CONTRACT-GEN-2 裁定）。
  - 凍結検証は呼び出し側（contract_webhook.verify_frozen_pdf）が pdf_text()
    の抽出テキストに対して行う。
"""

import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfgen import canvas

FONT_NAME = "HeiseiMin-W3"


class PdfUnsupportedStructure(ValueError):
    """fix1[02]: 表など段落以外の本文構造は非対応（黙って本文を落とさない
    ための構造 pin。doc.paragraphs にも docx_to_pdf_bytes の描画にも表セル
    は含まれず、全文一致 pin では欠落を検知できないため入口で拒否する）。"""

_PAGE_W, _PAGE_H = A4
_MARGIN_L = 22 * mm
_MARGIN_R = 22 * mm
_MARGIN_TOP = 25 * mm
_MARGIN_BOTTOM = 22 * mm
_BODY_SIZE = 10.5
_TITLE_SIZE = 14.0
_LEADING = 16.0     # 行送り
_PARA_GAP = 3.0     # 段落間
_EMPTY_GAP = 8.0    # 原本の空段落（空行）ぶんの余白


def _ensure_font() -> None:
    if FONT_NAME not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(UnicodeCIDFont(FONT_NAME))


def _wrap(text: str, size: float) -> list[str]:
    """文字単位の貪欲折返し（和文前提・禁則処理は行わない）。"""
    limit = _PAGE_W - _MARGIN_L - _MARGIN_R
    lines, cur = [], ""
    for ch in text:
        if cur and pdfmetrics.stringWidth(cur + ch, FONT_NAME, size) > limit:
            lines.append(cur)
            cur = ch
        else:
            cur += ch
    lines.append(cur)
    return lines


def docx_to_pdf_bytes(docx_bytes: bytes) -> bytes:
    """fill_template 済み docx の段落テキストを A4 縦 PDF に描画して返す。

    先頭の非空段落（表題）は中央寄せ・大きめ、以降は docx の段落 alignment
    （中央/右寄せ）を反映し、その他は左寄せ。改ページは行単位。"""
    _ensure_font()
    doc = Document(io.BytesIO(docx_bytes))
    if doc.tables:
        raise PdfUnsupportedStructure(
            "docx contains tables (paragraphs only)")
    buf = io.BytesIO()
    pdf = canvas.Canvas(buf, pagesize=A4)
    y = _PAGE_H - _MARGIN_TOP
    title_done = False

    def _draw_line(text: str, size: float, align) -> None:
        nonlocal y
        if y - _LEADING < _MARGIN_BOTTOM:
            pdf.showPage()
            y = _PAGE_H - _MARGIN_TOP
        y -= _LEADING
        pdf.setFont(FONT_NAME, size)
        if align == "center":
            pdf.drawCentredString(_PAGE_W / 2, y, text)
        elif align == "right":
            pdf.drawRightString(_PAGE_W - _MARGIN_R, y, text)
        else:
            pdf.drawString(_MARGIN_L, y, text)

    for para in doc.paragraphs:
        text = para.text.rstrip()
        if not text.strip():
            y -= _EMPTY_GAP
            continue
        if not title_done:
            size, align = _TITLE_SIZE, "center"
            title_done = True
        else:
            size = _BODY_SIZE
            if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                align = "center"
            elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                align = "right"
            else:
                align = "left"
        for line in _wrap(text, size):
            _draw_line(line, size, align)
        y -= _PARA_GAP

    pdf.showPage()
    pdf.save()
    return buf.getvalue()


def pdf_text(pdf_bytes: bytes) -> str:
    """生成 PDF の全ページ抽出テキスト（凍結検証・テスト用）。"""
    import fitz
    with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
        return "".join(page.get_text() for page in doc)
