"""HOUKI-CARD-TEMPLATE: 相続放棄 相談カード（来所用・A4 紙）の項目定義と雛形。

固定する仕様:
- hub/houki_card.CARD_ITEMS（項目番号→App 40 欄コード）が単一の正。番号は 1〜N で
  一意・7 群の順（HEARING_ROUNDS の見出しと同じ）。HEARING_WRITABLE_FIELDS の全欄
  （未成年後見関与=NOT_ASKED を除く）を網羅し、書込集合外は 債権者一覧 のみ
- 選択肢の文言は HEARING_CHOICE_FIELDS の逐語。未成年後見関与は載せない
- 雛形 docx は scripts/make_houki_card.py が決定的に生成（同一定義→同一 SHA-256）。
  PDF は同じ定義から描画し 2 ページ以内・各ページ下部に版とページ番号
"""

import hashlib
import io
import os
import sys
import unittest
import zipfile

os.environ.setdefault("KINTONE_SUBDOMAIN", "testsub")

import fitz  # noqa: E402  (PyMuPDF)
from docx import Document  # noqa: E402

from hub import houki_card as hc  # noqa: E402
from hub import houki_case_store as store  # noqa: E402

sys.path.insert(0, "scripts")
from make_houki_card import build_docx, build_pdf  # noqa: E402

TEMPLATE_PATH = "docx_templates/houki/相談カード.docx"
TEMPLATE_SHA256 = "42b569597843e5898767c0bc52781e142507bb13b56da38c044151f2964a3c87"


class TestCardItems(unittest.TestCase):
    def test_numbers_unique_sequential_and_grouped_in_order(self):
        numbers = [it.number for it in hc.CARD_ITEMS]
        self.assertEqual(numbers, list(range(1, len(hc.CARD_ITEMS) + 1)))
        groups = [it.group for it in hc.CARD_ITEMS]
        self.assertEqual(groups, sorted(groups))                # 7 群の順
        self.assertEqual(sorted(set(groups)), [g for g, _t in hc.GROUPS])
        self.assertEqual([t for _g, t in hc.GROUPS],
                         [r[0] for r in store.HEARING_ROUNDS])   # 見出し=台本の見出し
        self.assertEqual(hc.ITEM_FIELDS, {it.number: it.fields for it in hc.CARD_ITEMS})

    def test_fields_cover_all_writable_except_not_asked(self):
        covered = {f for it in hc.CARD_ITEMS for f in it.fields}
        expected = set(store.HEARING_WRITABLE_FIELDS) - {"未成年後見関与"}
        self.assertEqual(expected - covered, set())               # 漏れなし
        self.assertEqual(covered - set(store.HEARING_WRITABLE_FIELDS), {"債権者一覧"})
        self.assertNotIn("未成年後見関与", covered)
        # 記録欄なしは第 5 群（戸籍）のみ
        self.assertEqual([it.number for it in hc.CARD_ITEMS if not it.fields], [20])
        self.assertEqual(hc.CARD_ITEMS[19].group, 5)

    def test_choices_verbatim_from_choice_fields(self):
        for it in hc.CARD_ITEMS:
            if it.kind == "choice":
                code = it.fields[0]
                self.assertEqual(it.choices, store.HEARING_CHOICE_FIELDS[code], code)
        choice_codes = {it.fields[0] for it in hc.CARD_ITEMS if it.kind == "choice"}
        self.assertEqual(choice_codes,
                         set(store.HEARING_CHOICE_FIELDS) - {"未成年後見関与"})
        self.assertEqual(hc.CARD_ITEMS[19].choices, ("あり", "なし"))   # 戸籍（記録欄なし）

    def test_kinds_and_dates(self):
        kinds = {it.kind for it in hc.CARD_ITEMS}
        self.assertEqual(kinds, {"text", "kana_text", "date", "choice", "free",
                                 "creditors", "check_only"})
        dates = [it.fields for it in hc.CARD_ITEMS if it.kind == "date"]
        self.assertEqual(dates, [("死亡日_申告",), ("死亡を知った日_申告",),
                                 ("相続人と知った日_申告",), ("生年月日",)])
        self.assertEqual(hc.DATE_BOX, "西暦［　　　　］年［　　］月［　　］日")
        self.assertEqual(hc.CARD_FOOTER, "相続放棄 相談カード v1")


class TestTemplate(unittest.TestCase):
    def test_docx_sha_pinned_and_reproducible(self):
        data = open(TEMPLATE_PATH, "rb").read()
        self.assertEqual(hashlib.sha256(data).hexdigest(), TEMPLATE_SHA256)
        built = build_docx()
        self.assertEqual(built, data)                              # 定義→雛形の再現性
        self.assertEqual(build_docx(), built)                      # 決定的
        for info in zipfile.ZipFile(io.BytesIO(built)).infolist():
            self.assertEqual(info.date_time, (1980, 1, 1, 0, 0, 0))

    def test_docx_contents(self):
        doc = Document(TEMPLATE_PATH)
        text = "\n".join(p.text for p in doc.paragraphs)
        cells = [c.text for t in doc.tables for r in t.rows for c in r.cells]
        all_text = text + "\n" + "\n".join(cells)
        for it in hc.CARD_ITEMS:
            self.assertIn(f"{it.number}. {it.label}", all_text, it.number)
            for c in it.choices:
                self.assertIn(f"□ {c}", all_text)
        for _g, title in hc.GROUPS:
            self.assertIn(title, all_text)
        self.assertIn(hc.CHECK_NOTE, all_text)
        self.assertIn(hc.PRIVACY_NOTE, all_text)
        self.assertIn("記入日", all_text)
        self.assertIn("受付番号（事務所記入）", all_text)
        self.assertNotIn("未成年", all_text)
        self.assertEqual(all_text.count(hc.DATE_BOX), 5)          # 記入日+日付 4 項目
        footer = doc.sections[0].footer._element.xml
        self.assertIn(hc.CARD_FOOTER, footer)
        self.assertIn('w:instr="PAGE"', footer)                     # ページ番号
        # 債権者表: ヘッダ+3 行
        ctable = next(t for t in doc.tables
                      if t.rows[0].cells[1].text == hc.CREDITOR_COLUMNS[0])
        self.assertEqual(len(ctable.rows), 1 + hc.CREDITOR_ROWS)

    def test_pdf_two_pages_with_footer(self):
        pdf = build_pdf()
        d = fitz.open(stream=pdf, filetype="pdf")
        self.assertLessEqual(d.page_count, 2)
        for i, page in enumerate(d, start=1):
            t = page.get_text()
            self.assertIn(hc.CARD_FOOTER, t)
            self.assertIn(f"ページ {i}", t)
        full = "".join(p.get_text() for p in d)
        for it in hc.CARD_ITEMS:
            self.assertIn(f"{it.number}.", full)                  # 折返しで改行が入り得る
        self.assertIn(hc.PRIVACY_NOTE, full)


if __name__ == "__main__":
    unittest.main()
