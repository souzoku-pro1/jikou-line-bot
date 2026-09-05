"""JIKOU-CONTRACT-TOKUYAKU: 委任契約書の特約（App 21 特約 → docx/PDF）。

固定する仕様:
- 雛形（新 SHA）: 最後の条文（第12条本文）の直後・締結文の前に「特約事項」見出し+
  {{特約}} 本文（既存の条見出し/条文本文の書式を複製・単一 run・表なし）。
  scripts/add_contract_tokuyaku.py を旧雛形 fixture に適用すると同一 SHA（再現性）
- 差し込み: 特約 空/空白のみ=2 段落とも削除／非空=改行ごとに段落化（本文段落の
  pPr・rPr を複製）／特約中の {{ }} は展開しない／600 字超は生成せず「要確認」
  （CAS 遷移+通知）／FROZEN_CLAUSE は不変／PDF 全文=docx 全文
- 再生成ガード: cloudsign_document_id 非空の「契約書作成」は生成せず「要確認」+固定文言
- 既存 test_contract_gen1/gen2 は pin 更新（TEMPLATE_SHA256・登録キー数）以外無変更
"""

import hashlib
import io
import os
import sys
import unittest
from unittest.mock import AsyncMock, MagicMock, patch

_ENV = {
    "ANTHROPIC_API_KEY": "dummy", "LINE_CHANNEL_SECRET": "dummy_secret",
    "LINE_CHANNEL_ACCESS_TOKEN": "dummy_token", "KINTONE_SUBDOMAIN": "testsub",
    "KINTONE_APP_ID": "21", "KINTONE_API_TOKEN": "dummy",
    "SOUZOKU_KINTONE_APP_ID": "26", "SOUZOKU_KINTONE_API_TOKEN": "dummy",
    "CLOUDSIGN_CLIENT_ID": "c", "CLOUDSIGN_WEBHOOK_SECRET": "cs",
    "KINTONE_WEBHOOK_TOKEN": "kintone-token",
    "DOCUMENT_WEBHOOK_SECRET": "doc-secret",
    "APP_APPROVAL": "29", "TOKEN_APPROVAL": "d", "HEALTHCHECK_DISABLED": "1",
    "STRIPE_WEBHOOK_SECRET": "w", "GOOGLE_VISION_API_KEY": "dummy_vision",
}
for _k, _v in _ENV.items():
    os.environ.setdefault(_k, _v)

from docx import Document  # noqa: E402
from fastapi.testclient import TestClient  # noqa: E402

import contract_pdf  # noqa: E402
import contract_webhook as cw  # noqa: E402
import main  # noqa: E402
from config import EXPECTED_DOCX_TEMPLATES, EXPECTED_KINTONE_SCHEMA  # noqa: E402

sys.path.insert(0, "scripts")
from add_contract_tokuyaku import build_tokuyaku_template  # noqa: E402

_client = TestClient(main.app)
_URL = "/contract/doc-secret"
_OLD_TEMPLATE = "scripts/fixtures/委任契約書_2026-08-22.docx"
_OLD_SHA256 = "7cc168a1bbce3ca183e9f4a3d46b6b8288c17d4d21954f07cdf038428c355334"
_NEW_SHA256 = "ead90bb8154f64318cc80ee7d6a2dda129192936ca29e32746348ac6145856fd"
_EMAIL = "hanako@example.com"


def _rec(**fields):
    return {k: {"value": v} for k, v in fields.items()}


def _record(**over):
    base = {"$revision": "5", "契約書ステータス": "契約書作成",
            "顧客名": "熊澤花子", "住所": "埼玉県川口市青木1-1-1",
            "問い合わせ業者名": "株式会社Aファイナンス",
            "対象債権者2": "", "対象債権者3": "",
            "委任契約書": [], "cloudsign_document_id": "", "特約": ""}
    base.update(over)
    return _rec(**base)


def _body(record_id="12", status="契約書作成", app_id="21"):
    return {"app": {"id": app_id},
            "record": {"$id": {"value": record_id},
                       "契約書ステータス": {"value": status}}}


def _paras(data: bytes):
    return Document(io.BytesIO(data)).paragraphs


def _texts(data: bytes) -> list[str]:
    return [p.text for p in _paras(data)]


def _ppr(p):
    return p._p.pPr.xml if p._p.pPr is not None else None


def _rpr(p):
    return p.runs[0]._r.rPr.xml if p.runs and p.runs[0]._r.rPr is not None else None


class _Base(unittest.TestCase):
    def post(self, *, record, body=None, update=None):
        upload = AsyncMock(return_value="fk-1")
        update = update or AsyncMock()
        notify = AsyncMock(return_value=True)
        get = AsyncMock(return_value=record)
        with patch.dict(os.environ, _ENV), \
             patch.object(cw.hub_kintone, "get_record", get), \
             patch.object(cw.hub_kintone, "upload_file", upload), \
             patch.object(cw.hub_kintone, "update_record", update), \
             patch("hub.notify.notify_admin_line", notify):
            r = _client.post(_URL, json=body or _body())
        return r, upload, update, notify

    def generated(self, **over) -> bytes:
        r, upload, update, notify = self.post(record=_record(**over))
        self.assertEqual(r.status_code, 200, r.text)
        self.assertEqual(r.json().get("record_id"), "12")
        upload.assert_awaited_once()
        notify.assert_not_awaited()
        return upload.await_args.args[2]


# ── 1〜4: 差し込み ────────────────────────────────────────────────────────────
class TestTokuyakuFill(_Base):
    def test_1_single_line(self):
        docx = self.generated(特約="例外的に、報酬は2回の分割払いとする。")
        texts = _texts(docx)
        i = texts.index("特約事項")
        self.assertEqual(texts[i + 1], "例外的に、報酬は2回の分割払いとする。")
        self.assertEqual(texts[i - 1][:15], "本契約に関して甲乙間に紛争が生")   # 第12条本文の直後
        self.assertEqual(texts[i + 2], "")
        self.assertTrue(texts[i + 3].startswith("本契約の成立を証するため"))      # 締結文の前
        self.assertNotIn("{{", "\n".join(texts))
        self.assertEqual(cw._clause_of(docx), cw.FROZEN_CLAUSE)          # 第 2 条不変
        pdf = contract_pdf.docx_to_pdf_bytes(docx)
        self.assertIn("特約事項", contract_pdf.pdf_text(pdf))
        self.assertIn("例外的に、報酬は2回の分割払いとする。",
                      "".join(contract_pdf.pdf_text(pdf).split()))
        cw.verify_frozen_pdf(pdf)
        cw.verify_pdf_full_text(docx, pdf)

    def test_2_three_lines_paragraphs_with_body_format(self):
        docx = self.generated(特約="1 分割払いを認める。\n2 期限は3か月。\n3 遅延時は一括。")
        paras = _paras(docx)
        texts = [p.text for p in paras]
        i = texts.index("特約事項")
        self.assertEqual(texts[i + 1:i + 4],
                         ["1 分割払いを認める。", "2 期限は3か月。", "3 遅延時は一括。"])
        self.assertEqual(texts[i + 4], "")
        tmpl = Document(cw.TEMPLATE_PATH).paragraphs
        body_tmpl = next(p for p in tmpl if p.text == cw.TOKUYAKU_KEY)
        head_tmpl = next(p for p in tmpl if p.text == cw.TOKUYAKU_HEADING)
        for p in paras[i + 1:i + 4]:
            self.assertEqual(_ppr(p), _ppr(body_tmpl))
            self.assertEqual(_rpr(p), _rpr(body_tmpl))
            self.assertEqual(len(p.runs), 1)
        self.assertEqual(_ppr(paras[i]), _ppr(head_tmpl))
        self.assertEqual(_rpr(paras[i]), _rpr(head_tmpl))
        pdf = contract_pdf.docx_to_pdf_bytes(docx)
        flat = "".join(contract_pdf.pdf_text(pdf).split())
        self.assertIn("1分割払いを認める。2期限は3か月。3遅延時は一括。", flat)
        cw.verify_pdf_full_text(docx, pdf)

    def test_3_empty_or_blank_removes_both_paragraphs(self):
        with_one = _texts(self.generated(特約="例外的に分割払いとする。"))
        for blank in ("", "   ", "\n \n"):
            with self.subTest(blank=repr(blank)):
                docx = self.generated(特約=blank)
                texts = _texts(docx)
                self.assertNotIn("特約事項", texts)
                self.assertNotIn(cw.TOKUYAKU_KEY, texts)
                self.assertEqual(len(texts), len(Document(cw.TEMPLATE_PATH).paragraphs) - 2)
                self.assertEqual(texts, [t for t in with_one
                                         if t not in ("特約事項", "例外的に分割払いとする。")])
                self.assertEqual(cw._clause_of(docx), cw.FROZEN_CLAUSE)
                cw.verify_pdf_full_text(docx, contract_pdf.docx_to_pdf_bytes(docx))

    def test_4_placeholder_syntax_in_tokuyaku_is_rejected(self):
        # fix1 CT-02（仕様変更）: {{ }} を含む特約は「展開せずそのまま出す」から
        # 「入力不正として要確認」へ。旧 test_4 は本票の訂正により置換
        for bad in ("{{依頼者氏名}}は…", "x}}y", "{{", "a{{b}}c"):
            with self.subTest(bad=bad):
                r, upload, update, notify = self.post(record=_record(特約=bad))
                self.assertEqual(r.json().get("skip"), "tokuyaku_invalid")
                upload.assert_not_awaited()                             # docx/PDF とも生成しない
                update.assert_awaited_once()
                self.assertEqual(update.await_args.args[2], {"契約書ステータス": "要確認"})
                self.assertNotIn("委任契約書", update.await_args.args[2])   # 添付不変
                notify.assert_awaited_once_with(
                    "【委任契約書・要確認】特約欄に使用できない記号（{{ }}）が含まれて"
                    "います（レコード番号 12）。特約欄を修正して再度お試しください。")
        # 二重でない「{」「}」1 個は正常に生成
        docx = self.generated(特約="{分割払い}は可とする。")
        texts = _texts(docx)
        self.assertIn("{分割払い}は可とする。", texts)
        self.assertEqual("\n".join(texts).count("熊澤花子"), 2)
        # CloudSign 経路でも同じ判定で 500 にならず要確認
        record = _record(契約書ステータス="クラウドサイン登録",
                         委任契約書=[{"fileKey": "fk-0"}], メールアドレス=_EMAIL,
                         特約="{{依頼者氏名}}")
        update, notify = AsyncMock(), AsyncMock(return_value=True)
        with patch.dict(os.environ, _ENV), \
             patch.object(cw.hub_kintone, "get_record", AsyncMock(return_value=record)), \
             patch.object(cw.hub_kintone, "update_record", update), \
             patch.object(cw, "_cs_create_document", MagicMock()) as create, \
             patch("hub.notify.notify_admin_line", notify):
            r = _client.post(_URL, json=_body(status="クラウドサイン登録"))
        self.assertEqual(r.status_code, 200)
        self.assertEqual(r.json().get("skip"), "tokuyaku_invalid")
        create.assert_not_called()
        self.assertEqual(update.await_args.args[2], {"契約書ステータス": "要確認"})
        notify.assert_awaited_once()


# ── 5〜6: 生成ガード ─────────────────────────────────────────────────────────
class TestGuards(_Base):
    def test_5_over_limit_to_review(self):
        r, upload, update, notify = self.post(record=_record(特約="あ" * (cw.TOKUYAKU_MAX_CHARS + 1)))
        self.assertEqual(r.json().get("skip"), "tokuyaku_too_long")
        upload.assert_not_awaited()
        update.assert_awaited_once()
        self.assertEqual(update.await_args.args[2], {"契約書ステータス": "要確認"})
        self.assertEqual(update.await_args.kwargs.get("revision"), "5")
        self.assertNotIn("委任契約書", update.await_args.args[2])          # 添付不変
        notify.assert_awaited_once()
        self.assertIn("600 字", notify.await_args.args[0])
        self.assertIn("レコード番号 12", notify.await_args.args[0])
        self.assertEqual(cw.TOKUYAKU_MAX_CHARS, 600)
        # ちょうど 600 字は生成する
        docx = self.generated(特約="あ" * cw.TOKUYAKU_MAX_CHARS)
        self.assertIn("あ" * cw.TOKUYAKU_MAX_CHARS, _texts(docx))

    def test_6_cloudsign_registered_blocks_regeneration(self):
        r, upload, update, notify = self.post(
            record=_record(cloudsign_document_id="doc-1", 特約="分割払い可",
                           委任契約書=[{"fileKey": "old"}]))
        self.assertEqual(r.json().get("skip"), "cs_registered")
        upload.assert_not_awaited()
        update.assert_awaited_once()
        self.assertEqual(update.await_args.args[2], {"契約書ステータス": "要確認"})
        self.assertNotIn("委任契約書", update.await_args.args[2])          # 添付不変
        notify.assert_awaited_once_with(
            "【委任契約書・要確認】CloudSign 登録済みのため再生成を中止しました"
            "（レコード番号 12）。特約を反映する場合は CloudSign の下書きを削除して"
            "から再度お試しください。")
        # 空なら従来どおり再生成（添付置換）
        r, upload, update, notify = self.post(
            record=_record(cloudsign_document_id="", 特約="分割払い可",
                           委任契約書=[{"fileKey": "old"}]))
        self.assertEqual(r.json().get("record_id"), "12")
        upload.assert_awaited_once()
        self.assertEqual(update.await_args_list[-1].args[2]["委任契約書"],
                         [{"fileKey": "fk-1"}])
        notify.assert_not_awaited()
        # 判定順: CloudSign ガード → 特約上限（両方該当は CloudSign の文言）
        r, upload, update, notify = self.post(
            record=_record(cloudsign_document_id="doc-1",
                           特約="あ" * (cw.TOKUYAKU_MAX_CHARS + 1)))
        self.assertEqual(r.json().get("skip"), "cs_registered")

    def test_6b_reconcile_path_also_guarded(self):
        r, upload, update, notify = self.post(
            record=_record(契約書ステータス="契約書作成中",
                           cloudsign_document_id="doc-1"))
        self.assertEqual(r.json().get("skip"), "cs_registered")
        upload.assert_not_awaited()
        self.assertEqual(update.await_args.args[2], {"契約書ステータス": "要確認"})


# ── 7: CloudSign 用 PDF に特約が入る ─────────────────────────────────────────────
class TestCloudSignPdf(_Base):
    def test_7_pdf_contains_tokuyaku(self):
        record = _record(契約書ステータス="クラウドサイン登録",
                         委任契約書=[{"fileKey": "fk-0"}], メールアドレス=_EMAIL,
                         特約="1 分割払いを認める。\n2 期限は3か月。")
        attach = MagicMock()
        update = AsyncMock()
        notify = AsyncMock(return_value=True)
        with patch.dict(os.environ, _ENV), \
             patch.object(cw.hub_kintone, "get_record", AsyncMock(return_value=record)), \
             patch.object(cw.hub_kintone, "update_record", update), \
             patch.object(cw, "_cs_create_document", MagicMock(return_value="doc-1")), \
             patch.object(cw, "_cs_attach_pdf", attach), \
             patch.object(cw, "_cs_add_participant", MagicMock()), \
             patch.object(cw, "_cs_delete_draft", MagicMock(return_value=True)), \
             patch("hub.notify.notify_admin_line", notify):
            r = _client.post(_URL, json=_body(status="クラウドサイン登録"))
        self.assertEqual(r.status_code, 200, r.text)
        self.assertTrue(r.json().get("cloudsign"))
        _doc_id, pdf_bytes = attach.call_args.args
        flat = "".join(contract_pdf.pdf_text(pdf_bytes).split())
        self.assertIn("特約事項1分割払いを認める。2期限は3か月。", flat)
        cw.verify_frozen_pdf(pdf_bytes)
        notify.assert_not_awaited()
        # fix1: 上限超は登録経路でも共通ガード（要確認へ CAS+通知・API 作用 0）
        record = _record(契約書ステータス="クラウドサイン登録",
                         委任契約書=[{"fileKey": "fk-0"}], メールアドレス=_EMAIL,
                         特約="あ" * (cw.TOKUYAKU_MAX_CHARS + 1))
        update, notify = AsyncMock(), AsyncMock(return_value=True)
        with patch.dict(os.environ, _ENV), \
             patch.object(cw.hub_kintone, "get_record", AsyncMock(return_value=record)), \
             patch.object(cw.hub_kintone, "update_record", update), \
             patch.object(cw, "_cs_create_document", MagicMock()) as create, \
             patch("hub.notify.notify_admin_line", notify):
            r = _client.post(_URL, json=_body(status="クラウドサイン登録"))
        self.assertEqual(r.json().get("skip"), "tokuyaku_too_long")
        create.assert_not_called()
        self.assertEqual(update.await_args.args[2], {"契約書ステータス": "要確認"})
        notify.assert_awaited_once()

    def test_ct01_cloudsign_registered_blocks_reregistration(self):
        record = _record(契約書ステータス="クラウドサイン登録",
                         委任契約書=[{"fileKey": "fk-0"}], メールアドレス=_EMAIL,
                         cloudsign_document_id="doc-old", 特約="分割払い可")
        update, notify = AsyncMock(), AsyncMock(return_value=True)
        create, attach, participant = MagicMock(), MagicMock(), MagicMock()
        with patch.dict(os.environ, _ENV), \
             patch.object(cw.hub_kintone, "get_record", AsyncMock(return_value=record)), \
             patch.object(cw.hub_kintone, "update_record", update), \
             patch.object(cw.hub_kintone, "upload_file", AsyncMock()) as upload, \
             patch.object(cw, "_cs_create_document", create), \
             patch.object(cw, "_cs_attach_pdf", attach), \
             patch.object(cw, "_cs_add_participant", participant), \
             patch("hub.notify.notify_admin_line", notify):
            r = _client.post(_URL, json=_body(status="クラウドサイン登録"))
        self.assertEqual(r.status_code, 200)
        self.assertEqual(r.json().get("skip"), "cs_registered")
        create.assert_not_called()
        attach.assert_not_called()
        participant.assert_not_called()
        upload.assert_not_awaited()
        update.assert_awaited_once()
        self.assertEqual(update.await_args.args[2], {"契約書ステータス": "要確認"})
        self.assertEqual(update.await_args.kwargs.get("revision"), "5")
        self.assertNotIn("委任契約書", update.await_args.args[2])          # 添付不変
        notify.assert_awaited_once_with(
            "【委任契約書・要確認】CloudSign 登録済みのため再登録を中止しました"
            "（レコード番号 12）。既存の下書きを確認してください。")
        # 判定順（共通）: cs_registered → too_long → invalid
        rec = _record(cloudsign_document_id="doc-1",
                      特約="{{x}}" + "あ" * cw.TOKUYAKU_MAX_CHARS)
        self.assertEqual(cw._regeneration_guard("12", rec, "register")[0], "cs_registered")
        self.assertEqual(cw._regeneration_guard("12", rec, "regenerate")[0], "cs_registered")
        rec = _record(特約="{{x}}" + "あ" * cw.TOKUYAKU_MAX_CHARS)
        self.assertEqual(cw._regeneration_guard("12", rec, "register")[0], "tokuyaku_too_long")
        rec = _record(特約="{{x}}")
        self.assertEqual(cw._regeneration_guard("12", rec, "regenerate")[0], "tokuyaku_invalid")
        self.assertIsNone(cw._regeneration_guard("12", _record(特約="通常"), "register"))


# ── 8: 雛形 pin と再現性 ───────────────────────────────────────────────────────
class TestTemplatePins(unittest.TestCase):
    def test_8_sha_paragraphs_and_reproducibility(self):
        new = open(cw.TEMPLATE_PATH, "rb").read()
        self.assertEqual(hashlib.sha256(new).hexdigest(), _NEW_SHA256)
        self.assertEqual(cw.TEMPLATE_SHA256, _NEW_SHA256)
        old = open(_OLD_TEMPLATE, "rb").read()
        self.assertEqual(hashlib.sha256(old).hexdigest(), _OLD_SHA256)      # 旧雛形 fixture
        rebuilt = build_tokuyaku_template(old)
        self.assertEqual(rebuilt, new)                                       # 再現性（同一 SHA）
        self.assertEqual(build_tokuyaku_template(old), rebuilt)              # 決定的
        paras = Document(io.BytesIO(new)).paragraphs
        self.assertEqual(len(paras), 64 + 2)
        texts = [p.text for p in paras]
        self.assertEqual(texts[48], "第12条（合意管轄）")
        self.assertTrue(texts[49].startswith("本契約に関して甲乙間に紛争が生じた場合は"))
        self.assertEqual(texts[50], cw.TOKUYAKU_HEADING)
        self.assertEqual(texts[51], cw.TOKUYAKU_KEY)
        self.assertEqual(texts[52], "")
        self.assertTrue(texts[53].startswith("本契約の成立を証するため"))
        # 書式は既存の条見出し/条文本文の複製・単一 run・表なし
        self.assertEqual(_ppr(paras[50]), _ppr(paras[48]))
        self.assertEqual(_rpr(paras[50]), _rpr(paras[48]))
        self.assertEqual(_ppr(paras[51]), _ppr(paras[49]))
        self.assertEqual(_rpr(paras[51]), _rpr(paras[49]))
        self.assertEqual(len(paras[51].runs), 1)
        self.assertEqual(Document(io.BytesIO(new)).tables, [])
        # 旧雛形との差分は 2 段落の挿入のみ（他の文面は一字も変えない）
        old_texts = [p.text for p in Document(io.BytesIO(old)).paragraphs]
        self.assertEqual(texts[:50] + texts[52:], old_texts)
        # 第 2 条は不変・登録・スキーマ
        self.assertEqual(cw._clause_of(new), cw.FROZEN_CLAUSE)
        self.assertIn("{{特約}}", EXPECTED_DOCX_TEMPLATES["docx_templates/jikou/委任契約書.docx"])
        app21 = next(v for v in EXPECTED_KINTONE_SCHEMA.values()
                     if isinstance(v, dict) and "特約" in v.get("fields", {}))
        self.assertEqual(app21["fields"]["特約"], {"type": "MULTI_LINE_TEXT"})

    def test_script_rejects_double_apply(self):
        new = open(cw.TEMPLATE_PATH, "rb").read()
        with self.assertRaises(ValueError):
            build_tokuyaku_template(new)


if __name__ == "__main__":
    unittest.main()
