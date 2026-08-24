"""JIKOU-NOTICE-1: 時効援用通知書の自動生成。

固定する仕様:
- テンプレの正 = docx_templates/jikou/時効援用通知書.docx（大野修正版から
  make_notice_template.py で個人情報をプレースホルダ化して収載・SHA-256
  pin・25 段落・プレースホルダ 6 種）。旧個人情報の文字列は残存しない
  （機械検査）。
- 大野裁定（改変禁止）: 宛先「債権者各位」のまま・冒頭の事務所住所ブロック
  （〒/住所/建物/通知代理人/TEL/FAX）は凍結本文・旧住所は空なら段落ごと
  削除/非空なら差し込み・対象債権者 1〜3 の非空ごとに 1 通生成し notice_file
  へ複数添付（単一レコード 1 PUT・宛先固定のため各通の本文は同一）。
- 本文凍結検証: 生成物の全段落がテンプレ由来の期待列と完全一致（差し込み
  箇所以外はテンプレ逐語）・プレースホルダ残存なし。不一致は添付せず 500。
- 書式 pin（fix2）: 記載ブロックの値 run の rPr が債務者氏名行と同一・
  ラベルの均等割り付け（fitText）は値へ及ばない（凍結一致検証はテキストの
  検査で run 書式の崩れを検出しないため、テンプレと生成物の両方で pin）。
- webhook は CONTRACT-GEN 確立構造の同型（/notice/・DOCUMENT_WEBHOOK_SECRET
  共用・契約書ステータス DROP_DOWN 共用で新 3 値・CAS $revision・409 のみ
  cas_lost・reconcile 添付あり=要確認・fail-closed=必須欠落/債権者 0 件は
  状態不変+フィールド名のみ通知）。
"""

import hashlib
import io
import os
import unittest
from unittest.mock import AsyncMock, patch

_ENV = {
    "ANTHROPIC_API_KEY": "dummy", "LINE_CHANNEL_SECRET": "dummy_secret",
    "LINE_CHANNEL_ACCESS_TOKEN": "dummy_token", "KINTONE_SUBDOMAIN": "testsub",
    "KINTONE_APP_ID": "21", "KINTONE_API_TOKEN": "dummy",
    "SOUZOKU_KINTONE_APP_ID": "26", "SOUZOKU_KINTONE_API_TOKEN": "dummy",
    "CLOUDSIGN_CLIENT_ID": "c", "CLOUDSIGN_WEBHOOK_SECRET": "cs",
    "KINTONE_WEBHOOK_TOKEN": "kintone-token",
    "DOCUMENT_WEBHOOK_SECRET": "doc-secret",
    "APP_APPROVAL": "29", "TOKEN_APPROVAL": "d", "HEALTHCHECK_DISABLED": "1",
    "STRIPE_WEBHOOK_SECRET": "w", "GOOGLE_VISION_API_KEY": "dummy_vision",
}
for _k, _v in _ENV.items():
    os.environ.setdefault(_k, _v)

from docx import Document  # noqa: E402
from fastapi.testclient import TestClient  # noqa: E402

import main  # noqa: E402
import notice_webhook as nw  # noqa: E402
from config import EXPECTED_KINTONE_SCHEMA  # noqa: E402
from hub.kintone import KintoneError  # noqa: E402

_client = TestClient(main.app)
_URL = "/notice/doc-secret"

# 収載時に置換した旧個人情報（値そのものはテストに持たない=氏名等の断片で
# 残存を機械検査する。事務所住所ブロックは裁定により凍結＝検査対象外）
_FORBIDDEN_RESIDUE = ("黒田", "達人", "くろだ", "たつひと", "たつと",
                      "蕨市", "赤田住宅")

PLACEHOLDERS = ("{{通知日付}}", "{{通知人氏名}}", "{{ふりがな}}",
                "{{生年月日}}", "{{通知人住所}}", "{{旧住所}}")


def _rec(**fields):
    return {k: {"value": v} for k, v in fields.items()}


def _full_record(**over):
    base = {"$revision": "5", "契約書ステータス": "時効援用通知作成",
            "顧客名": "試験太郎", "住所": "埼玉県川口市青木1-1-1",
            "生年月日": "昭和55年1月1日", "furigana": "しけんたろう",
            "old_address": "",
            "問い合わせ業者名": "株式会社Aファイナンス",
            "対象債権者2": "", "対象債権者3": "",
            "notice_file": []}
    base.update(over)
    return _rec(**base)


def _body(record_id="12", status="時効援用通知作成", app_id="21"):
    body = {"record": {"$id": {"value": record_id},
                       "契約書ステータス": {"value": status}}}
    if app_id is not None:
        body["app"] = {"id": app_id}
    return body


def _texts(data: bytes) -> list[str]:
    return [p.text for p in Document(io.BytesIO(data)).paragraphs]


# fix2: 記載ブロックの書式 pin（凍結一致検証は段落テキストの検査＝run 書式の
# 崩れを検出しないため、この穴を塞ぐ）。
_LABEL_ROWS = ("ふりがな", "債務者氏名", "生年月日", "住　　　所", "旧　住　所")


def _rpr_sig(run) -> tuple:
    """run の rPr を (要素名, 属性) の列に正規化（フォント名・サイズ・
    fitText・spacing 等の明示書式をすべて含む）。"""
    rpr = run._r.rPr
    if rpr is None:
        return ()
    return tuple(sorted(
        (c.tag.split("}")[1],
         tuple(sorted((k.split("}")[1], v) for k, v in c.attrib.items())))
        for c in rpr))


def _row(doc, label):
    return next(p for p in doc.paragraphs if p.text.startswith(label))


def _value_run(p, value):
    """差し込み値（またはプレースホルダ）を含む run（単一 run であること）。"""
    runs = [r for r in p.runs if value in r.text]
    assert len(runs) == 1, (p.text, [r.text for r in p.runs])
    return runs[0]


class TestTemplatePin(unittest.TestCase):
    def test_sha256_pinned(self):
        data = open(nw.TEMPLATE_PATH, "rb").read()
        self.assertEqual(hashlib.sha256(data).hexdigest(), nw.TEMPLATE_SHA256)

    def test_structure_and_placeholders(self):
        texts = [p.text for p in Document(nw.TEMPLATE_PATH).paragraphs]
        joined = "\n".join(texts)
        self.assertEqual(len(texts), 25)
        self.assertEqual(texts[1], "債権者各位")           # 宛先は裁定どおり
        self.assertEqual(texts[9], "時効援用通知書")
        # プレースホルダ 6 種（実テンプレ抽出・推測で増やさない）
        for ph in PLACEHOLDERS:
            self.assertIn(ph, joined)
        self.assertEqual(joined.count("{{通知人氏名}}"), 3)
        # 凍結本文の要点（信用情報文言・付言・代理人ブロック）
        self.assertIn("信用情報機関(CIC・JICC等)", joined)
        self.assertIn("債務の承認には該当しない", joined)
        self.assertIn("通知代理人　弁護士　大野太郎", joined)
        # 旧住所行は住所行の直後
        addr_idx = next(i for i, t in enumerate(texts)
                        if "{{通知人住所}}" in t)
        self.assertIn("{{旧住所}}", texts[addr_idx + 1])
        self.assertIn("旧　住　所", texts[addr_idx + 1])

    def test_no_personal_residue(self):
        # 収載後、旧個人情報の文字列が残っていないこと（機械検査）
        joined = "\n".join(p.text for p in Document(nw.TEMPLATE_PATH).paragraphs)
        for token in _FORBIDDEN_RESIDUE:
            self.assertNotIn(token, joined)

    def test_placeholders_single_run(self):
        # fill 互換: プレースホルダは run 跨ぎしない
        for p in Document(nw.TEMPLATE_PATH).paragraphs:
            if "{{" in p.text:
                self.assertTrue(any("{{" in r.text and "}}" in r.text
                                    for r in p.runs), p.text)

    def test_label_rows_run_format_preserved(self):
        # fix2: 元票の「ふりがな」「生年月日」ラベルは 5 文字幅の均等割り付け
        # （w:fitText val=1050）。収載時に行全体を fitText 付き run へ潰すと
        # 行全体が 1050 twips に圧縮され小さな崩れた字になる（実機事象）。
        # ラベル run の fitText は元票どおり残し、プレースホルダ run は
        # 債務者氏名行の値 run と同一書式（fitText なし）であること。
        doc = Document(nw.TEMPLATE_PATH)
        ref = _rpr_sig(_value_run(_row(doc, "債務者氏名"), "{{通知人氏名}}"))
        self.assertFalse(any(tag == "fitText" for tag, _ in ref))
        for label, ph in (("ふりがな", "{{ふりがな}}"),
                          ("生年月日", "{{生年月日}}"),
                          ("住　　　所", "{{通知人住所}}"),
                          ("旧　住　所", "{{旧住所}}")):
            with self.subTest(label=label):
                p = _row(doc, label)
                self.assertEqual(_rpr_sig(_value_run(p, ph)), ref)
        for label in ("ふりがな", "生年月日"):
            with self.subTest(label=label):
                p = _row(doc, label)
                first = p.runs[0]
                self.assertTrue(first.text and label.startswith(first.text))
                fit = [dict(attrs) for tag, attrs in _rpr_sig(first)
                       if tag == "fitText"]
                self.assertEqual([f.get("val") for f in fit], ["1050"])


class TestBuildAndFreeze(unittest.TestCase):
    _FILL = {"{{通知日付}}": "令和8年8月23日", "{{通知人氏名}}": "試験太郎",
             "{{ふりがな}}": "しけんたろう", "{{生年月日}}": "昭和55年1月1日",
             "{{通知人住所}}": "埼玉県川口市青木1-1-1"}

    def test_old_address_present_branch(self):
        data = nw.build_notice_docx(self._FILL, "東京都台東区旧町2-2-2")
        texts = _texts(data)
        self.assertEqual(len(texts), 25)
        row = next(t for t in texts if "旧　住　所" in t)
        self.assertIn("東京都台東区旧町2-2-2", row)
        nw.verify_notice_docx(data, self._FILL, "東京都台東区旧町2-2-2")

    def test_old_address_empty_branch_removes_paragraph(self):
        data = nw.build_notice_docx(self._FILL, "")
        texts = _texts(data)
        self.assertEqual(len(texts), 24)                 # 段落ごと削除
        self.assertFalse(any("旧　住　所" in t for t in texts))
        nw.verify_notice_docx(data, self._FILL, "")

    def test_frozen_body_verbatim(self):
        # 差し込み箇所以外の本文はテンプレと逐語一致（期待列=全段落一致）
        data = nw.build_notice_docx(self._FILL, "")
        texts = _texts(data)
        tpl = [p.text for p in Document(nw.TEMPLATE_PATH).paragraphs]
        for t in tpl:
            if "{{" not in t:                    # 凍結段落はそのまま残る
                self.assertIn(t, texts)
        self.assertIn("債権者各位", texts)
        self.assertEqual(sum(t.count("試験太郎") for t in texts), 3)
        self.assertNotIn("{{", "\n".join(texts))

    def test_filled_rows_font_consistent(self):
        # fix2: 生成物の記載ブロック 5 行で、差し込み値を持つ run の書式
        # （フォント名・サイズ・fitText 等の明示 rPr）が債務者氏名行と同一。
        # 差し込み時に段落全体を先頭 run へ潰すと、ふりがな・生年月日行だけ
        # ラベルの fitText(1050) に行全体が圧縮される（実機で発見）。
        old = "東京都台東区旧町2-2-2"
        doc = Document(io.BytesIO(nw.build_notice_docx(self._FILL, old)))
        ref_run = _value_run(_row(doc, "債務者氏名"), "試験太郎")
        ref = _rpr_sig(ref_run)
        self.assertFalse(any(tag == "fitText" for tag, _ in ref))
        for label, value in (("ふりがな", "しけんたろう"),
                             ("生年月日", "昭和55年1月1日"),
                             ("住　　　所", "埼玉県川口市青木1-1-1"),
                             ("旧　住　所", old)):
            with self.subTest(label=label):
                run = _value_run(_row(doc, label), value)
                self.assertEqual(_rpr_sig(run), ref)
                self.assertEqual(run.font.name, ref_run.font.name)
                self.assertEqual(run.font.size, ref_run.font.size)
                # ラベルは値と別 run（ラベル書式が値へ及ばない）
                self.assertFalse(run.text.startswith(label), run.text)

    def test_tampered_body_rejected(self):
        # 凍結本文の 1 字改変（44,000 と違い本文なので任意の語）を検知
        data = nw.build_notice_docx(self._FILL, "")
        doc = Document(io.BytesIO(data))
        for p in doc.paragraphs:
            if "時効が完成しております" in p.text and p.runs:
                p.runs[0].text = p.text.replace("完成", "未完成")
                for r in p.runs[1:]:
                    r.text = ""
        buf = io.BytesIO()
        doc.save(buf)
        with self.assertRaises(nw.NoticeIntegrityError):
            nw.verify_notice_docx(buf.getvalue(), self._FILL, "")

    def test_unfilled_placeholder_rejected(self):
        data = open(nw.TEMPLATE_PATH, "rb").read()   # 未差し込みテンプレ
        with self.assertRaises(nw.NoticeIntegrityError):
            nw.verify_notice_docx(data, self._FILL, "x")

    def test_template_hash_mismatch_rejected(self):
        with patch.object(nw, "TEMPLATE_SHA256", "0" * 64):
            with self.assertRaises(nw.NoticeIntegrityError):
                nw.verify_template_integrity()


class _WebhookBase(unittest.TestCase):
    def _post(self, *, record, body=None, url=_URL,
              upload=None, update=None):
        upload = upload or AsyncMock(side_effect=[f"fk-{i}"
                                                  for i in range(1, 10)])
        update = update or AsyncMock()
        notify = AsyncMock(return_value=True)
        get = AsyncMock(return_value=record)
        with patch.dict(os.environ, _ENV), \
             patch.object(nw.hub_kintone, "get_record", get), \
             patch.object(nw.hub_kintone, "upload_file", upload), \
             patch.object(nw.hub_kintone, "update_record", update), \
             patch("hub.notify.notify_admin_line", notify):
            r = _client.post(url, json=body or _body())
        return r, upload, update, notify, get


class TestWebhookEntry(_WebhookBase):
    def test_wrong_secret_403(self):
        r, *_ = self._post(record=_full_record(), url="/notice/wrong")
        self.assertEqual(r.status_code, 403)

    def test_app_mismatch_zero_effects(self):
        for app_id in ("26", None, "abc"):
            with self.subTest(app_id=app_id):
                r, upload, update, _n, get = self._post(
                    record=_full_record(), body=_body(app_id=app_id))
                self.assertEqual(r.json().get("skip"), "app_mismatch")
                get.assert_not_awaited()
                upload.assert_not_awaited()

    def test_not_triggered_body_gate(self):
        # 共用 DROP_DOWN の他フロー値・自 update echo は本文 gate で作用 0
        for status in ("契約書作成", "クラウドサイン登録",
                       "時効援用通知作成中", "時効援用通知作成済",
                       "要確認", ""):
            with self.subTest(status=status):
                r, upload, _u, _n, get = self._post(
                    record=_full_record(), body=_body(status=status))
                self.assertEqual(r.json().get("skip"), "not_triggered")
                get.assert_not_awaited()
                upload.assert_not_awaited()

    def test_stale_authoritative_states(self):
        cases = {"時効援用通知作成済": "already_done",
                 "契約書作成": "stale_status", "要確認": "stale_status",
                 "": "stale_status"}
        for state, skip in cases.items():
            with self.subTest(state=state):
                r, upload, update, _n, _g = self._post(
                    record=_full_record(契約書ステータス=state))
                self.assertEqual(r.json().get("skip"), skip)
                upload.assert_not_awaited()
                update.assert_not_awaited()


class TestWebhookStateMachine(_WebhookBase):
    def test_happy_path_single_creditor(self):
        r, upload, update, notify, _g = self._post(record=_full_record())
        self.assertEqual(r.status_code, 200)
        self.assertEqual(r.json().get("record_id"), "12")
        self.assertEqual(update.await_count, 2)
        cas_call, final_call = update.await_args_list
        self.assertEqual(cas_call.args[2],
                         {"契約書ステータス": "時効援用通知作成中"})
        self.assertEqual(cas_call.kwargs.get("revision"), "5")
        self.assertEqual(final_call.args[2]["契約書ステータス"],
                         "時効援用通知作成済")
        self.assertEqual(final_call.args[2]["notice_file"],
                         [{"fileKey": "fk-1"}])
        self.assertEqual(final_call.kwargs.get("revision"), "6")
        # 生成物: 凍結検証済み実体・氏名 3 箇所・ふりがな・旧住所行なし
        self.assertEqual(upload.await_count, 1)
        self.assertEqual(upload.await_args.args[1],
                         "時効援用通知書_対象債権者1.docx")
        texts = _texts(upload.await_args.args[2])
        self.assertEqual(sum(t.count("試験太郎") for t in texts), 3)
        self.assertTrue(any("しけんたろう" in t for t in texts))
        self.assertFalse(any("旧　住　所" in t for t in texts))
        notify.assert_not_awaited()

    def test_three_creditors_three_files(self):
        record = _full_record(対象債権者2="株式会社Bローン",
                              対象債権者3="C債権回収株式会社",
                              old_address="東京都台東区旧町2-2-2")
        r, upload, update, _n, _g = self._post(record=record)
        self.assertEqual(r.status_code, 200)
        self.assertEqual(upload.await_count, 3)          # 債権者ごとに 1 通
        names = [c.args[1] for c in upload.await_args_list]
        self.assertEqual(names, ["時効援用通知書_対象債権者1.docx",
                                 "時効援用通知書_対象債権者2.docx",
                                 "時効援用通知書_対象債権者3.docx"])
        final_call = update.await_args_list[1]
        self.assertEqual(final_call.args[2]["notice_file"],
                         [{"fileKey": "fk-1"}, {"fileKey": "fk-2"},
                          {"fileKey": "fk-3"}])
        # 旧住所行あり分岐
        texts = _texts(upload.await_args_list[0].args[2])
        self.assertTrue(any("旧　住　所" in t and "旧町2-2-2" in t
                            for t in texts))

    def test_sparse_slots_keep_original_numbers(self):
        # fix1[01]（Codex 指定 3 形）: 空欄除去後も元の枠番号を保持——
        # ファイル名の N は常に App21 の入力枠（監査対応）。生成・添付順は
        # 枠番号昇順で決定的
        cases = {
            "枠2のみ": ({"問い合わせ業者名": "", "対象債権者2": "株式会社B",
                         "対象債権者3": ""},
                        ["時効援用通知書_対象債権者2.docx"]),
            "枠1と3": ({"問い合わせ業者名": "株式会社A", "対象債権者2": "",
                        "対象債権者3": "C債権回収株式会社"},
                       ["時効援用通知書_対象債権者1.docx",
                        "時効援用通知書_対象債権者3.docx"]),
            "枠3のみ": ({"問い合わせ業者名": "", "対象債権者2": "",
                         "対象債権者3": "C債権回収株式会社"},
                        ["時効援用通知書_対象債権者3.docx"]),
        }
        for label, (fields, want_names) in cases.items():
            with self.subTest(case=label):
                r, upload, update, _n, _g = self._post(
                    record=_full_record(**fields))
                self.assertEqual(r.status_code, 200)
                names = [c.args[1] for c in upload.await_args_list]
                self.assertEqual(names, want_names)
                self.assertNotIn("時効援用通知書_対象債権者2.docx",
                                 names if label == "枠1と3" else [])
                final_call = update.await_args_list[1]
                self.assertEqual(
                    final_call.args[2]["notice_file"],
                    [{"fileKey": f"fk-{i + 1}"}
                     for i in range(len(want_names))])

    def test_partial_upload_failure_orphan_filekeys_by_spec(self):
        """fix1[02]（採用方式 (a)）: 3 通中 2 通目の upload 失敗 → 500・
        最終 PUT なし（CAS のみ=作成中維持→reconcile 回収）。成功済み
        fileKey は未添付のまま残るが、kintone 公式仕様「一時保管領域に
        保存されたファイルは、レコードやスペースなどに添付されない場合、
        3日間で削除されます」（cybozu developer network
        「ファイルをアップロードする」制限事項・
        https://cybozu.dev/ja/kintone/docs/rest-api/files/upload-file/ ・
        2026-08-23 確認）により孤立 fileKey は自動回収されるため、
        再実行が最初から再 upload しても蓄積しない。"""
        record = _full_record(対象債権者2="株式会社B",
                              対象債権者3="C債権回収株式会社")
        upload = AsyncMock(side_effect=["fk-1", RuntimeError("kintone down")])
        r, _u, update, _n, _g = self._post(record=record, upload=upload)
        self.assertEqual(r.status_code, 500)
        self.assertEqual(update.await_count, 1)          # CAS のみ・PUT なし
        self.assertEqual(update.await_args.args[2],
                         {"契約書ステータス": "時効援用通知作成中"})

    def test_zero_creditors_fails_closed(self):
        record = _full_record(問い合わせ業者名="", 対象債権者2="",
                              対象債権者3="")
        r, upload, update, notify, _g = self._post(record=record)
        self.assertEqual(r.json().get("skip"), "missing_fields")
        upload.assert_not_awaited()
        update.assert_not_awaited()              # 状態も動かさない
        notify.assert_awaited_once()
        self.assertIn("債権者", notify.await_args.args[0])

    def test_missing_fields_rejected_names_only(self):
        cases = {"ふりがな": _full_record(furigana=""),
                 "生年月日": _full_record(生年月日=""),
                 "住所": _full_record(住所="")}
        for label, record in cases.items():
            with self.subTest(case=label):
                r, upload, update, notify, _g = self._post(record=record)
                self.assertEqual(r.json().get("skip"), "missing_fields")
                upload.assert_not_awaited()
                update.assert_not_awaited()
                sent = notify.await_args.args[0]
                self.assertIn(label, sent)
                self.assertNotIn("試験", sent)   # 値・PII は通知に載せない

    def test_concurrent_second_request_loses_cas(self):
        loser = AsyncMock(side_effect=KintoneError(409, "GAIA_CO02"))
        r, upload, update, _n, _g = self._post(record=_full_record(),
                                               update=loser)
        self.assertEqual(r.json().get("skip"), "cas_lost")
        upload.assert_not_awaited()
        self.assertEqual(update.await_count, 1)

    def test_claim_failures_not_silenced(self):
        for err in (KintoneError(500, "GAIA_XX01", "x"),
                    KintoneError(0, "transport_error", "x")):
            with self.subTest(status=err.status):
                update = AsyncMock(side_effect=err)
                r, upload, _u, _n, _g = self._post(record=_full_record(),
                                                   update=update)
                self.assertEqual(r.status_code, 500)
                upload.assert_not_awaited()

    def test_recovery_no_attachment(self):
        record = _full_record(契約書ステータス="時効援用通知作成中")
        record["$revision"] = {"value": "7"}
        r, upload, update, notify, _g = self._post(record=record)
        self.assertTrue(r.json().get("recovered"))
        self.assertEqual(upload.await_count, 1)
        self.assertEqual(update.await_args_list[1].kwargs.get("revision"),
                         "8")
        notify.assert_not_awaited()

    def test_working_with_attachment_goes_review(self):
        record = _full_record(契約書ステータス="時効援用通知作成中")
        record["notice_file"] = {"value": [{"fileKey": "old"}]}
        r, upload, update, notify, _g = self._post(record=record)
        self.assertEqual(r.json().get("skip"), "needs_review")
        upload.assert_not_awaited()
        self.assertEqual(update.await_args.args[2],
                         {"契約書ステータス": "要確認"})
        notify.assert_awaited_once()
        self.assertIn("自動では上書きせず", notify.await_args.args[0])

    def test_upload_failure_500_stays_working(self):
        upload = AsyncMock(side_effect=RuntimeError("kintone down"))
        r, _u, update, _n, _g = self._post(record=_full_record(),
                                           upload=upload)
        self.assertEqual(r.status_code, 500)
        self.assertEqual(update.await_count, 1)  # CAS のみ

    def test_contract_gate_ignores_notice_values(self):
        # 共用 DROP_DOWN の相互不干渉: contract 側 gate は通知系の値を
        # not_triggered で落とす（トリガ 2 値のみ通過の既存仕様）
        import contract_webhook as cw
        get = AsyncMock()
        with patch.dict(os.environ, _ENV), \
             patch.object(cw.hub_kintone, "get_record", get):
            r = _client.post("/contract/doc-secret",
                             json=_body(status="時効援用通知作成"))
        self.assertEqual(r.json().get("skip"), "not_triggered")
        get.assert_not_awaited()


class TestSchemaPin(unittest.TestCase):
    def test_new_fields_registered(self):
        fields = EXPECTED_KINTONE_SCHEMA["App 21 (案件)"]["fields"]
        self.assertEqual(fields["furigana"]["type"], "SINGLE_LINE_TEXT")
        self.assertEqual(fields["old_address"]["type"], "SINGLE_LINE_TEXT")
        self.assertEqual(fields["notice_file"]["type"], "FILE")

    def test_status_options_include_notice_values(self):
        opts = (EXPECTED_KINTONE_SCHEMA["App 21 (案件)"]["fields"]
                ["契約書ステータス"]["required_options"])
        for v in ("時効援用通知作成", "時効援用通知作成中",
                  "時効援用通知作成済"):
            self.assertIn(v, opts)


if __name__ == "__main__":
    unittest.main()
