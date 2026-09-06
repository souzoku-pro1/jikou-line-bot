"""HOUKI-CONTRACT-GEN: 相続放棄 委任契約書の自動生成（App 40・ChannelConfig 化）。

固定する仕様:
- 雛形 docx_templates/houki/委任契約書.docx は人承認済み現物（SHA256 pin）・全 62 段落の
  テキスト sha256 pin・12 個の差し込み記号が各 1 回・表なし
- 費用定数（88,000 / 33,000 / 5,000 / 1,100 / 3 か所）と代表者定型文は弁護士凍結事項
- 申述人集合（被相続人グループID）・費用計算・申述人一覧の展開・契約署名 3 通り
- 状態機械・CAS・_regeneration_guard 3 経路は時効版と共通（cfg 経由）
- cloudsign_webhook は App 21 → App 40 の順に検索し、両方該当は fail-closed
- 時効側の既存テスト（test_contract_gen1/gen2/tokuyaku・test_cloudsign_*）は無変更で green
- fix1 HCG-01: 管理レコード（被相続人グループID == 自レコード番号）だけが生成/登録の経路に
  入れる。非管理レコードは 要確認（代表者番号を通知）。規則 4（他レコードの契約状態）。
  HCG-02: 登録直前に管理レコードを ID 指定で再取得して再検証。HCG-03: 指紋は JSON 直列化
- fix2 HCGF1-01: 再取得後に 契約書ステータス と revision で所有権を確認してから外部 API。
  HCGF1-02: 結果不明は cloudsign_document_id の印「結果不明:要手動確認」で永続的に遮断
"""

import asyncio
import datetime
import hashlib
import httpx
import io
import os
import re
import unittest
from unittest.mock import AsyncMock, MagicMock, patch

_ENV = {
    "ANTHROPIC_API_KEY": "dummy", "LINE_CHANNEL_SECRET": "dummy_secret",
    "LINE_CHANNEL_ACCESS_TOKEN": "dummy_token", "KINTONE_SUBDOMAIN": "testsub",
    "KINTONE_APP_ID": "21", "KINTONE_API_TOKEN": "dummy",
    "SOUZOKU_KINTONE_APP_ID": "26", "SOUZOKU_KINTONE_API_TOKEN": "dummy",
    "CLOUDSIGN_CLIENT_ID": "c", "CLOUDSIGN_WEBHOOK_SECRET": "cs",
    "KINTONE_WEBHOOK_TOKEN": "kintone-token",
    "DOCUMENT_WEBHOOK_SECRET": "doc-secret",
    "APP_APPROVAL": "29", "TOKEN_APPROVAL": "d", "HEALTHCHECK_DISABLED": "1",
    "STRIPE_WEBHOOK_SECRET": "w", "GOOGLE_VISION_API_KEY": "dummy_vision",
    "APP_CHATLOG": "28", "TOKEN_CHATLOG": "d",
    "APP_HOUKI": "40", "TOKEN_HOUKI": "d",
    "HOUKI_LINE_CHANNEL_SECRET": "houki_secret",
    "HOUKI_WEBHOOK_TOKEN": "houki-hook",
}
for _k, _v in _ENV.items():
    os.environ.setdefault(_k, _v)

from docx import Document  # noqa: E402
from fastapi.testclient import TestClient  # noqa: E402

import contract_pdf  # noqa: E402
import contract_webhook as cw  # noqa: E402
import main  # noqa: E402
from config import EXPECTED_DOCX_TEMPLATES  # noqa: E402
from hub import houki_contract as hc  # noqa: E402
from hub import kintone as hub_kintone  # noqa: E402
from hub import notify as hub_notify  # noqa: E402
from hub.kintone import KintoneError  # noqa: E402

_client = TestClient(main.app)
_URL = "/souzoku-houki/contract/houki-hook"
NOW = datetime.datetime(2026, 9, 6, 15, 0, tzinfo=hc.JST)
EMAIL_A, EMAIL_B, EMAIL_C = "a@example.com", "b@example.com", "c@example.com"


def _run(coro):
    return asyncio.run(coro)


def _docx_paras(data: bytes) -> list[str]:
    return [p.text for p in Document(io.BytesIO(data)).paragraphs]


def _creditors(*names):
    return [{"value": {"債権者名": {"value": n}, "通知要否": {"value": "未確認"}}} for n in names]


def _rec(rid, name, addr="架空県架空市1-1", email=EMAIL_A, decedent="架空太郎",
         group="", sign="", tokuyaku="", creditors=(), status="契約書作成",
         revision="5", attachment=None, cs_doc_id=""):
    base = {
        "$id": rid, "$revision": revision, "契約書ステータス": status,
        "顧客名": name, "住所": addr, "メールアドレス": email, "被相続人氏名": decedent,
        "被相続人グループID": group, "契約署名": sign, "特約": tokuyaku,
        "cloudsign_document_id": cs_doc_id, "委任契約書": attachment or [],
    }
    rec = {k: {"value": v} for k, v in base.items()}
    rec["債権者一覧"] = {"value": _creditors(*creditors)}
    return rec


def _body(rid="1", status="契約書作成", app="40"):
    return {"app": {"id": app}, "record": {"$id": {"value": rid},
                                           "契約書ステータス": {"value": status}}}


class _Base(unittest.TestCase):
    """App 40 の最小フェイク: get_record / search_records（被相続人グループID）/
    update_record（CAS）/ upload_file。"""

    def setUp(self):
        self.records: dict[str, dict] = {}
        self.uploads: list[tuple] = []
        self.updates: list[tuple] = []
        self.admin = AsyncMock(return_value=True)
        self.cs_create = MagicMock(return_value="doc-h1")
        self.cs_attach = MagicMock()
        self.cs_participant = MagicMock()
        self.cs_delete = MagicMock(return_value=True)

        async def get_record(app, rid):
            rec = self.records.get(str(rid))
            if rec is None:
                raise KintoneError(404, "GAIA_RE01", "nf")
            return {k: dict(v) for k, v in rec.items()}

        async def search_records(app, query, fields=None):
            m = re.search(r'被相続人グループID = "([^"]+)"', query)
            rows = [r for r in self.records.values()
                    if m and r["被相続人グループID"]["value"] == m.group(1)]
            rows.sort(key=lambda r: int(r["$id"]["value"]))
            return [{k: dict(v) for k, v in r.items()} for r in rows]

        async def update_record(app, rid, fields, revision=None):
            rec = self.records[str(rid)]
            cur = int(rec["$revision"]["value"])
            if revision is not None and int(revision) != cur:
                raise KintoneError(409, "GAIA_CO02", "c")
            self.updates.append((str(rid), dict(fields)))
            for k, v in fields.items():
                rec[k] = {"value": v}
            rec["$revision"] = {"value": str(cur + 1)}

        async def upload_file(app, filename, content, mime):
            self.uploads.append((filename, content, mime))
            return f"fk-{len(self.uploads)}"

        for p in (patch.dict(os.environ, _ENV),
                  patch.object(cw.hub_kintone, "get_record", get_record),
                  patch.object(cw.hub_kintone, "update_record", update_record),
                  patch.object(cw.hub_kintone, "upload_file", upload_file),
                  patch.object(hub_kintone, "search_records", search_records),
                  patch.object(cw, "_cs_create_document", self.cs_create),
                  patch.object(cw, "_cs_attach_pdf", self.cs_attach),
                  patch.object(cw, "_cs_add_participant", self.cs_participant),
                  patch.object(cw, "_cs_delete_draft", self.cs_delete),
                  patch.object(hub_notify, "notify_admin_line", self.admin),
                  patch("hub.notify.notify_admin_line", self.admin)):
            p.start()
            self.addCleanup(p.stop)

    def seed(self, *recs):
        for r in recs:
            self.records[r["$id"]["value"]] = r
        return recs[0]["$id"]["value"]

    def post(self, rid="1", status="契約書作成", url=_URL, app="40"):
        return _client.post(url, json=_body(rid, status, app))

    def field(self, rid, code):
        return self.records[str(rid)][code]["value"]

    def kinds(self):
        return [c.kwargs.get("throttle_key", "").split(":", 1)[0]
                for c in self.admin.await_args_list]

    def texts(self):
        return [c.args[0] for c in self.admin.await_args_list]

    def last_docx(self) -> bytes:
        return self.uploads[-1][1]


# ── 1. 雛形と凍結 pin ─────────────────────────────────────────────────────────
class TestTemplatePins(unittest.TestCase):
    def test_template_sha_and_paragraph_hash(self):
        data = open(hc.TEMPLATE_PATH, "rb").read()
        self.assertEqual(hashlib.sha256(data).hexdigest(), hc.TEMPLATE_SHA256)
        self.assertEqual(hc.TEMPLATE_SHA256[:8], "519d4bb6")
        doc = Document(hc.TEMPLATE_PATH)
        texts = [p.text for p in doc.paragraphs]
        self.assertEqual(len(texts), hc.TEMPLATE_PARAGRAPH_COUNT)
        self.assertEqual(hc.template_paragraphs_sha256(), hc.TEMPLATE_PARAGRAPHS_SHA256)
        self.assertEqual(len(doc.tables), 0)
        full = "\n".join(texts)
        for ph in hc.PLACEHOLDERS:
            self.assertEqual(full.count(ph), 1, ph)
            runs = [r for p in doc.paragraphs for r in p.runs if ph in r.text]
            self.assertEqual(len(runs), 1, ph)
        self.assertEqual(set(re.findall(r"\{\{[^{}]*\}\}", full)), set(hc.PLACEHOLDERS))
        self.assertEqual(len(hc.PLACEHOLDERS), 12)
        self.assertEqual(EXPECTED_DOCX_TEMPLATES[hc.TEMPLATE_PATH], list(hc.PLACEHOLDERS))
        # 凍結条項は雛形の段落と逐語一致
        self.assertEqual(hc.FROZEN_CLAUSE, tuple(texts[9:11]))

    def test_tampered_template_detected(self):
        with patch.object(hc, "TEMPLATE_SHA256", "0" * 64):
            with self.assertRaises(hc.HoukiContractIntegrityError):
                hc.verify_template_integrity()
        with patch.object(hc, "TEMPLATE_PARAGRAPHS_SHA256", "0" * 64):
            self.assertNotEqual(hc.template_paragraphs_sha256(), hc.TEMPLATE_PARAGRAPHS_SHA256)

    def test_frozen_constants_pinned(self):
        self.assertEqual((hc.FEE_BASE, hc.FEE_ADDITIONAL, hc.DEPOSIT_PER_APPLICANT,
                          hc.EXTRA_SEND_FEE, hc.INCLUDED_DESTINATIONS),
                         (88000, 33000, 5000, 1100, 3))
        self.assertEqual(hc.REPRESENTATIVE_CLAUSE,
                         "甲らは、申述人{name}を代表者と定め、代表者が甲ら全員のために本契約に電子署名する。")
        self.assertEqual(hc.TOKUYAKU_NONE, "特になし")
        self.assertIn("88,000", hc.FROZEN_CLAUSE[1])
        self.assertIn("33,000", hc.FROZEN_CLAUSE[1])
        self.assertIn("1,100", hc.FROZEN_CLAUSE[1])
        self.assertEqual(cw.TOKUYAKU_MAX_CHARS, 600)


# ── 2. 費用計算 ───────────────────────────────────────────────────────────────
class TestFees(unittest.TestCase):
    def test_fees_by_applicants_and_destinations(self):
        cases = {
            (1, 0): (88000, 5000, 0, 0, 93000),
            (1, 3): (88000, 5000, 0, 0, 93000),
            (1, 4): (88000, 5000, 1, 1100, 94100),
            (2, 0): (121000, 10000, 0, 0, 131000),
            (3, 5): (154000, 15000, 2, 2200, 171200),
        }
        for (n, d), (fee, dep, cnt, extra, total) in cases.items():
            with self.subTest(n=n, d=d):
                f = hc.compute_fees(n, d)
                self.assertEqual((f["報酬合計"], f["実費合計"], f["追加送付件数"],
                                  f["追加送付料合計"], f["支払総額"]), (fee, dep, cnt, extra, total))
        self.assertEqual(hc.fmt_yen(154000), "154,000")

    def test_destinations_union_normalized(self):
        a = _rec("1", "甲", creditors=("株式会社ＡＢＣ", "ｱｺﾑ", " アコム ", "", "アコム　"))
        b = _rec("2", "乙", creditors=("株式会社ABC", "プロミス", "プロ　ミス"))
        d = hc.destinations([a, b])
        self.assertEqual(d, {"株式会社ABC", "アコム", "プロミス"})
        self.assertEqual(hc.normalize_creditor("  "), "")


# ── 3. 申述人一覧の展開・特約・署名 ───────────────────────────────────────────
class TestRender(unittest.TestCase):
    def _render(self, records, mode):
        fees = hc.compute_fees(len(records), len(hc.destinations(records)))
        return hc.render(records, mode, fees, NOW)

    def test_one_applicant_expansion_and_fill(self):
        r = _rec("1", "甲野一郎", addr="架空県架空市1-1", creditors=("アコム",))
        paras = _docx_paras(self._render([r], hc.SIGN_ALL))
        i = paras.index("甲（申述人）")
        self.assertEqual(paras[i + 1:i + 3], ["住所　架空県架空市1-1", "氏名　甲野一郎"])
        self.assertEqual(paras[i + 3], "乙（受任者）")
        self.assertIn("弁護士報酬（申述人1名分）　金88,000円", paras)
        self.assertIn("受理通知書の追加送付料（送付先0か所分）　金0円", paras)
        self.assertIn("実費預託金（申述人1名につき5,000円）　金5,000円", paras)
        self.assertIn("支払総額　金93,000円", paras)
        self.assertIn("〔契約日〕　2026年9月6日", paras)
        self.assertTrue(any("被相続人架空太郎（以下" in p for p in paras))
        self.assertEqual(paras[paras.index("特約事項") + 1], "特になし")
        self.assertFalse(any("{{" in p for p in paras))
        self.assertEqual(len(Document(io.BytesIO(self._render([r], hc.SIGN_ALL))).tables), 0)

    def test_three_applicants_expansion_blank_between(self):
        recs = [_rec("3", "甲野一郎", addr="住所A"), _rec("1", "甲野二郎", addr="住所B"),
                _rec("2", "甲野三郎", addr="住所C")]
        paras = _docx_paras(self._render(recs, hc.SIGN_ALL))
        i = paras.index("甲（申述人）")
        self.assertEqual(paras[i + 1:i + 9],
                         ["住所　住所A", "氏名　甲野一郎", "", "住所　住所B", "氏名　甲野二郎",
                          "", "住所　住所C", "氏名　甲野三郎"])
        self.assertIn("弁護士報酬（申述人3名分）　金154,000円", paras)
        self.assertIn("支払総額　金169,000円", paras)

    def test_tokuyaku_by_sign_mode(self):
        rep = _rec("1", "甲野一郎", tokuyaku="分割払いを認める")
        other = _rec("2", "甲野二郎")
        # 全員: 欄そのまま
        self.assertEqual(hc.compose_tokuyaku([rep, other], hc.SIGN_ALL), "分割払いを認める")
        # 代表者のみ（2 名以上）: 定型文 + 改行 + 欄
        self.assertEqual(hc.compose_tokuyaku([rep, other], hc.SIGN_REPRESENTATIVE),
                         "甲らは、申述人甲野一郎を代表者と定め、代表者が甲ら全員のために本契約に"
                         "電子署名する。\n分割払いを認める")
        # 代表者のみ・欄が空: 定型文のみ
        rep2 = _rec("1", "甲野一郎")
        self.assertEqual(hc.compose_tokuyaku([rep2, other], hc.SIGN_REPRESENTATIVE),
                         "甲らは、申述人甲野一郎を代表者と定め、代表者が甲ら全員のために本契約に"
                         "電子署名する。")
        # 1 名: 定型文なし（空なら 特になし）
        self.assertEqual(hc.compose_tokuyaku([rep2], hc.SIGN_REPRESENTATIVE), "特になし")
        self.assertEqual(hc.compose_tokuyaku([rep], hc.SIGN_REPRESENTATIVE), "分割払いを認める")
        # docx では 2 段落（定型文・欄）に分かれる
        paras = _docx_paras(self._render([rep, other], hc.SIGN_REPRESENTATIVE))
        i = paras.index("特約事項")
        self.assertTrue(paras[i + 1].startswith("甲らは、申述人甲野一郎を代表者と定め"))
        self.assertEqual(paras[i + 2], "分割払いを認める")

    def test_sign_mode_and_participants(self):
        rep = _rec("1", "甲野一郎", email=EMAIL_A)
        other = _rec("2", "甲野二郎", email=EMAIL_B)
        self.assertEqual(hc.sign_mode(_rec("1", "x", sign="全員")), ("全員", False))
        self.assertEqual(hc.sign_mode(_rec("1", "x", sign="代表者のみ")), ("代表者のみ", False))
        self.assertEqual(hc.sign_mode(_rec("1", "x", sign="")), ("代表者のみ", True))
        self.assertEqual(hc.participants([rep, other], hc.SIGN_ALL),
                         ([(EMAIL_A, "甲野一郎"), (EMAIL_B, "甲野二郎")], []))
        self.assertEqual(hc.participants([rep, other], hc.SIGN_REPRESENTATIVE),
                         ([(EMAIL_A, "甲野一郎")], []))
        parts, problems = hc.participants([rep, _rec("2", "甲野二郎", email="")], hc.SIGN_ALL)
        self.assertEqual(parts, [(EMAIL_A, "甲野一郎")])
        self.assertEqual(problems, ["メールアドレス 未入力（レコード番号 2）"])
        _p, problems = hc.participants([_rec("1", "x", email="bad")], hc.SIGN_ALL)
        self.assertEqual(problems, ["メールアドレス 形式不正（レコード番号 1）"])

    def test_frozen_clause_verified_at_render(self):
        r = _rec("1", "甲野一郎")
        fees = hc.compute_fees(1, 0)
        with patch.object(hc, "FROZEN_CLAUSE", (hc.FROZEN_CLAUSE[0], "改変")):
            with self.assertRaises(cw.ContractIntegrityError):
                hc.render([r], hc.SIGN_ALL, fees, NOW)
        pdf = contract_pdf.docx_to_pdf_bytes(hc.render([r], hc.SIGN_ALL, fees, NOW))
        cw.verify_frozen_pdf(pdf, cw._houki_cfg())
        with self.assertRaises(cw.ContractIntegrityError):
            cw.verify_frozen_pdf(pdf)                       # 時効版の第2条は含まれない


# ── 4. 受け口・状態機械（生成） ──────────────────────────────────────────────
class TestGenerateFlow(_Base):
    def test_token_gates(self):
        self.seed(_rec("1", "甲野一郎"))
        with patch.dict(os.environ, {"HOUKI_WEBHOOK_TOKEN": ""}):
            self.assertEqual(self.post().status_code, 404)
        with patch.dict(os.environ, {"HOUKI_WEBHOOK_TOKEN": os.environ["DOCUMENT_WEBHOOK_SECRET"]}):
            self.assertEqual(self.post(url="/souzoku-houki/contract/doc-secret").status_code, 404)
        self.assertEqual(self.post(url="/souzoku-houki/contract/wrong").status_code, 403)
        self.assertEqual(self.post(app="21").json().get("skip"), "app_mismatch")
        self.assertEqual(self.post(status="契約書作成中").json().get("skip"), "not_triggered")
        self.assertEqual(self.uploads, [])
        self.assertEqual(self.updates, [])

    def test_single_applicant_group_blank(self):
        self.seed(_rec("1", "甲野一郎", creditors=("アコム", "プロミス", "レイク", "アイフル")))
        r = self.post()
        self.assertEqual(r.status_code, 200, r.text)
        self.assertEqual(self.field("1", "契約書ステータス"), "契約書作成済")
        self.assertEqual(self.field("1", "委任契約書"), [{"fileKey": "fk-1"}])
        self.assertEqual(self.uploads[0][0], "委任契約書_相続放棄.docx")
        paras = _docx_paras(self.last_docx())
        self.assertIn("弁護士報酬（申述人1名分）　金88,000円", paras)
        self.assertIn("受理通知書の追加送付料（送付先1か所分）　金1,100円", paras)
        self.assertIn("支払総額　金94,100円", paras)
        self.assertEqual(self.kinds(), ["houki_contract_created"])
        text = self.texts()[0]
        self.assertIn("・申述人数: 1名", text)
        self.assertIn("・送付先数: 4か所（追加送付 1か所）", text)
        self.assertIn("・支払総額: 94,100円", text)
        self.assertIn("・署名方式: 代表者のみ（契約署名 が空のため代表者のみとして作成）", text)
        self.assertIn("・被相続人グループID 空・1 名として作成", text)
        self.assertNotIn("甲野一郎", text)
        # CAS 遷移: 作成→作成中→作成済（起点レコードのみ更新）
        self.assertEqual([u[1].get("契約書ステータス") for u in self.updates],
                         ["契約書作成中", "契約書作成済"])

    def test_group_of_three_representative_first_then_id_order(self):
        self.seed(_rec("5", "代表花子", addr="住所5", group="5", sign="全員",
                       creditors=("アコム",)),
                  _rec("2", "二郎", addr="住所2", group="5", email=EMAIL_B,
                       creditors=("ｱｺﾑ", "プロミス"), status=""),
                  _rec("9", "九郎", addr="住所9", group="5", email=EMAIL_C, status=""),
                  _rec("7", "無関係", group="7"))
        r = self.post(rid="5")
        self.assertEqual(r.status_code, 200, r.text)
        paras = _docx_paras(self.last_docx())
        i = paras.index("甲（申述人）")
        self.assertEqual(paras[i + 1:i + 9],
                         ["住所　住所5", "氏名　代表花子", "", "住所　住所2", "氏名　二郎", "",
                          "住所　住所9", "氏名　九郎"])
        self.assertIn("弁護士報酬（申述人3名分）　金154,000円", paras)
        self.assertIn("受理通知書の追加送付料（送付先0か所分）　金0円", paras)   # 2 か所 ≤ 3
        self.assertIn("支払総額　金169,000円", paras)
        self.assertEqual(paras[paras.index("特約事項") + 1], "特になし")
        # 他の申述人レコードは更新しない
        self.assertTrue(all(rid == "5" for rid, _f in self.updates))
        self.assertEqual(self.field("2", "契約書ステータス"), "")

    def test_review_when_decedent_mismatch_or_missing_fields(self):
        for label, recs in {
            "decedent": [_rec("1", "甲", group="1", decedent="A"),
                         _rec("2", "乙", group="1", decedent="B", status="")],
            "name": [_rec("1", "甲", group="1"), _rec("2", "", group="1", status="")],
            "addr": [_rec("1", "甲", group="1"), _rec("2", "乙", group="1", addr="", status="")],
        }.items():
            with self.subTest(label=label):
                self.setUp()
                self.seed(*recs)
                r = self.post()
                self.assertEqual(r.json().get("skip"), "houki_preconditions")
                self.assertEqual(self.field("1", "契約書ステータス"), "要確認")
                self.assertEqual(self.uploads, [])
                self.assertEqual(self.kinds(), ["houki_contract_needs_review"])
                self.assertIn("レコード番号 2" if label != "decedent" else "被相続人氏名",
                              self.texts()[0])

    def test_regeneration_guard_three_paths(self):
        # ① 契約書作成トリガ: 特約に {{ }} → 要確認（生成 0）
        self.seed(_rec("1", "甲", tokuyaku="{{x}}"))
        self.assertEqual(self.post().json().get("skip"), "tokuyaku_invalid")
        self.assertEqual(self.field("1", "契約書ステータス"), "要確認")
        # ② 作成中 reconcile（添付なし）: 特約 601 字 → 要確認
        self.setUp()
        self.seed(_rec("1", "甲", status="契約書作成中", tokuyaku="あ" * 601))
        self.assertEqual(self.post().json().get("skip"), "tokuyaku_too_long")
        self.assertEqual(self.field("1", "契約書ステータス"), "要確認")
        # ③ クラウドサイン登録: cloudsign_document_id 非空 → 要確認・API 0 回
        self.setUp()
        self.seed(_rec("1", "甲", status="クラウドサイン登録", cs_doc_id="doc-old",
                       attachment=[{"fileKey": "k"}]))
        self.assertEqual(self.post(status="クラウドサイン登録").json().get("skip"), "cs_registered")
        self.assertEqual(self.field("1", "契約書ステータス"), "要確認")
        self.cs_create.assert_not_called()
        self.assertEqual(self.kinds(), ["houki_contract_needs_review"])
        # 定型文は上限に含めない（欄 600 字ちょうど+代表者定型文でも生成される）
        self.setUp()
        self.seed(_rec("1", "甲", group="1", tokuyaku="あ" * 600),
                  _rec("2", "乙", group="1", email=EMAIL_B, status=""))
        self.assertEqual(self.post().status_code, 200)
        self.assertEqual(self.field("1", "契約書ステータス"), "契約書作成済")

    def test_cas_lost_and_already_done(self):
        self.seed(_rec("1", "甲", revision="5"))
        real_update = cw.hub_kintone.update_record

        async def conflict(app, rid, fields, revision=None):
            raise KintoneError(409, "GAIA_CO02", "c")
        with patch.object(cw.hub_kintone, "update_record", conflict):
            self.assertEqual(self.post().json().get("skip"), "cas_lost")
        self.assertEqual(self.uploads, [])
        self.records["1"]["契約書ステータス"] = {"value": "契約書作成済"}
        self.assertEqual(self.post().json().get("skip"), "already_done")

    def test_internal_error_notifies_failed(self):
        self.seed(_rec("1", "甲"))
        with patch.object(hc, "render", side_effect=RuntimeError("boom")):
            r = self.post()
        self.assertEqual(r.status_code, 500)
        self.assertEqual(self.kinds(), ["houki_contract_failed"])
        self.assertEqual(self.field("1", "契約書ステータス"), "契約書作成中")   # reconcile 対象


# ── 5. CloudSign 登録 ─────────────────────────────────────────────────────────
class TestCloudSignFlow(_Base):
    def _seed_registered(self, sign="全員", emails=(EMAIL_A, EMAIL_B)):
        return self.seed(
            _rec("1", "代表花子", group="1", sign=sign, email=emails[0],
                 status="クラウドサイン登録", attachment=[{"fileKey": "k"}]),
            _rec("2", "二郎", group="1", email=emails[1], status=""))

    def test_all_signers(self):
        self._seed_registered("全員")
        r = self.post(status="クラウドサイン登録")
        self.assertEqual(r.status_code, 200, r.text)
        self.assertTrue(r.json().get("cloudsign"))
        self.cs_create.assert_called_once_with("1", "相続放棄_委任契約書_案件No.1")
        self.assertEqual(self.cs_attach.call_args.args[0], "doc-h1")
        self.assertEqual(self.cs_attach.call_args.args[2], "委任契約書_相続放棄.pdf")
        self.assertEqual([c.args for c in self.cs_participant.call_args_list],
                         [("doc-h1", EMAIL_A, "代表花子"), ("doc-h1", EMAIL_B, "二郎")])
        self.assertEqual(self.field("1", "cloudsign_document_id"), "doc-h1")
        self.assertEqual(self.field("1", "契約書ステータス"), "クラウドサイン登録済")
        self.assertEqual(self.kinds(), ["houki_contract_created"])
        self.assertIn("・署名方式: 全員", self.texts()[0])
        # PDF の凍結検証（第2条）と全文一致は共通経路で実行済み＝添付 PDF に {{ なし
        pdf = self.cs_attach.call_args.args[1]
        self.assertNotIn("{{", contract_pdf.pdf_text(pdf))

    def test_representative_only_and_blank(self):
        for sign in ("代表者のみ", ""):
            with self.subTest(sign=sign):
                self.setUp()
                self._seed_registered(sign, emails=(EMAIL_A, ""))     # 2 名目のメール空でも可
                r = self.post(status="クラウドサイン登録")
                self.assertEqual(r.status_code, 200, r.text)
                self.assertEqual([c.args for c in self.cs_participant.call_args_list],
                                 [("doc-h1", EMAIL_A, "代表花子")])
                text = self.texts()[0]
                self.assertIn("・署名方式: 代表者のみ", text)
                self.assertEqual("契約署名 が空のため" in text, sign == "")
                pdf_text = contract_pdf.pdf_text(self.cs_attach.call_args.args[1])
                self.assertIn("代表者と定め", "".join(pdf_text.split()))

    def test_email_missing_is_review(self):
        self._seed_registered("全員", emails=(EMAIL_A, ""))
        self.assertEqual(self.post(status="クラウドサイン登録").json().get("skip"),
                         "houki_preconditions")
        self.assertEqual(self.field("1", "契約書ステータス"), "要確認")
        self.cs_create.assert_not_called()
        self.assertIn("メールアドレス 未入力（レコード番号 2）", self.texts()[0])
        self.setUp()
        self._seed_registered("代表者のみ", emails=("", EMAIL_B))
        self.assertEqual(self.post(status="クラウドサイン登録").json().get("skip"),
                         "houki_preconditions")
        self.assertIn("レコード番号 1", self.texts()[0])

    def test_toctou_applicant_set_changed_before_send(self):
        self._seed_registered("全員")
        real_plan = hc.plan
        calls = {"n": 0}

        async def plan_then_change(record):
            calls["n"] += 1
            p = await real_plan(record)
            if calls["n"] == 1:                       # 1 回目の後に 3 人目が追加される
                self.records["3"] = _rec("3", "三郎", group="1", email=EMAIL_C, status="")
            return p
        with patch.object(hc, "plan", plan_then_change):
            r = self.post(status="クラウドサイン登録")
        self.assertEqual(r.json().get("skip"), "applicants_changed")
        self.assertEqual(self.field("1", "契約書ステータス"), "要確認")
        self.cs_create.assert_not_called()
        self.assertEqual(self.kinds(), ["houki_contract_needs_review"])

    def test_docx_missing_is_precondition_notice_only(self):
        self.seed(_rec("1", "代表花子", status="クラウドサイン登録"))
        self.assertEqual(self.post(status="クラウドサイン登録").json().get("skip"),
                         "cs_preconditions")
        self.assertEqual(self.field("1", "契約書ステータス"), "クラウドサイン登録")
        self.assertEqual(self.kinds(), ["houki_contract_needs_review"])


# ── 5b. fix1 HCG-01: 管理レコード方式 ───────────────────────────────────────────
class TestManagerRecord(_Base):
    """A=管理レコード（No.1・group "1"）・B=同グループの申述人（No.2・group "1"）。"""

    def _seed_pair(self, status="契約書作成"):
        att = [{"fileKey": "k"}] if status == "クラウドサイン登録" else None
        self.seed(_rec("1", "代表花子", group="1", sign="全員", status=status, attachment=att),
                  _rec("2", "二郎", group="1", email=EMAIL_B, status=status, attachment=att))

    def _counts(self):
        return len(self.uploads), self.cs_create.call_count

    def _expect_one(self, status):
        return (1, 0) if status == "契約書作成" else (0, 1)

    def test_sequential_a_then_b_and_b_then_a(self):
        for status in ("契約書作成", "クラウドサイン登録"):
            for order in (("1", "2"), ("2", "1")):
                with self.subTest(status=status, order=order):
                    self.setUp()
                    self._seed_pair(status)
                    results = {rid: self.post(rid=rid, status=status).json() for rid in order}
                    self.assertEqual(self._counts(), self._expect_one(status))
                    self.assertEqual(results["2"].get("skip"), hc.NOT_MANAGER)
                    self.assertEqual(self.field("2", "契約書ステータス"), "要確認")
                    self.assertEqual(self.field("1", "契約書ステータス"),
                                     "契約書作成済" if status == "契約書作成"
                                     else "クラウドサイン登録済")
                    review_text = [t for t in self.texts() if hc.NOT_MANAGER in t][0]
                    self.assertIn("代表者のレコード（No.1）で行ってください", review_text)
                    self.assertNotIn("二郎", review_text)

    def _parallel(self, first, second, status):
        """first を prepare（plan）内で停止させ、その間に second を完走させる。"""
        gate, at_gate = asyncio.Event(), asyncio.Event()
        real_plan = hc.plan
        calls = {"n": 0}

        async def plan_gate(record):
            calls["n"] += 1
            if calls["n"] == 1:
                at_gate.set()
                await gate.wait()
            return await real_plan(record)

        async def scenario():
            with patch.object(hc, "plan", plan_gate):
                async with httpx.AsyncClient(transport=httpx.ASGITransport(app=main.app),
                                             base_url="http://t") as client:
                    task_a = asyncio.create_task(
                        client.post(_URL, json=_body(first, status)))
                    await asyncio.wait_for(at_gate.wait(), timeout=5)
                    res_b = (await client.post(_URL, json=_body(second, status))).json()
                    gate.set()
                    res_a = (await task_a).json()
                    return res_a, res_b
        return _run(scenario())

    def test_parallel_a_and_b(self):
        for status in ("契約書作成", "クラウドサイン登録"):
            with self.subTest(status=status):
                self.setUp()
                self._seed_pair(status)
                res_a, res_b = self._parallel("1", "2", status)
                self.assertEqual(res_b.get("skip"), hc.NOT_MANAGER)
                self.assertEqual(res_a.get("record_id"), "1")
                self.assertEqual(self._counts(), self._expect_one(status))
                self.assertEqual(self.field("2", "契約書ステータス"), "要確認")

    def test_parallel_a_and_a(self):
        for status in ("契約書作成", "クラウドサイン登録"):
            with self.subTest(status=status):
                self.setUp()
                self._seed_pair(status)
                res_a1, res_a2 = self._parallel("1", "1", status)
                self.assertEqual(res_a2.get("record_id"), "1")          # 後続が勝つ
                self.assertEqual(res_a1.get("skip"), "cas_lost")        # 先行は CAS 敗北
                self.assertEqual(self._counts(), self._expect_one(status))

    def test_group_invalid_and_missing(self):
        self.seed(_rec("1", "甲", group="abc"))
        self.assertEqual(self.post().json().get("skip"), hc.GROUP_INVALID)
        self.assertEqual(self.field("1", "契約書ステータス"), "要確認")
        self.assertIn("数値ではありません", self.texts()[0])
        self.setUp()
        self.seed(_rec("1", "甲", group="999"))
        self.assertEqual(self.post().json().get("skip"), hc.GROUP_MISSING)
        self.assertEqual(self.field("1", "契約書ステータス"), "要確認")
        self.assertIn("No.999", self.texts()[0])
        self.assertEqual(self.uploads, [])

    def test_rule4_other_member_has_contract(self):
        for label, other in {
            "status": _rec("2", "乙", group="1", email=EMAIL_B, status="契約書作成済"),
            "doc_id": _rec("2", "乙", group="1", email=EMAIL_B, status="", cs_doc_id="doc-z"),
        }.items():
            with self.subTest(label=label):
                self.setUp()
                self.seed(_rec("1", "甲", group="1", sign="全員"), other)
                self.assertEqual(self.post().json().get("skip"), hc.OTHER_MEMBER_HAS_CONTRACT)
                self.assertEqual(self.field("1", "契約書ステータス"), "要確認")
                self.assertIn("レコード番号 2", self.texts()[0])
                self.assertEqual(self.uploads, [])
        # 登録直前の再検証でも規則 4 が効く
        self.setUp()
        self.seed(_rec("1", "甲", group="1", sign="全員", status="クラウドサイン登録",
                       attachment=[{"fileKey": "k"}]),
                  _rec("2", "乙", group="1", email=EMAIL_B, status=""))
        real_plan = hc.plan
        calls = {"n": 0}

        async def plan_then_taint(record):
            calls["n"] += 1
            p = await real_plan(record)
            if calls["n"] == 1:
                self.records["2"]["cloudsign_document_id"] = {"value": "doc-z"}
            return p
        with patch.object(hc, "plan", plan_then_taint):
            self.assertEqual(self.post(status="クラウドサイン登録").json().get("skip"),
                             "applicants_changed")
        self.cs_create.assert_not_called()
        self.assertEqual(self.field("1", "契約書ステータス"), "要確認")

    def test_manager_check_and_ordering(self):
        self.assertEqual(hc.manager_check(_rec("1", "x", group="")), (None, ""))
        self.assertEqual(hc.manager_check(_rec("1", "x", group="1")), (None, "1"))
        self.assertEqual(hc.manager_check(_rec("2", "x", group="1")), (hc.NOT_MANAGER, "1"))
        self.assertEqual(hc.manager_check(_rec("2", "x", group="x1")), (hc.GROUP_INVALID, "x1"))
        self.seed(_rec("7", "代表", group="7"), _rec("3", "a", group="7", status=""),
                  _rec("12", "b", group="7", status=""), _rec("8", "他", group="8"))
        recs, group, problems = _run(hc.gather_applicants(self.records["7"]))
        self.assertEqual(([r["$id"]["value"] for r in recs], group, problems), (["7", "3", "12"], "7", []))


# ── 5c. fix1 HCG-02/03: 登録直前の再検証と指紋 ──────────────────────────────────
class TestPreSendRevalidation(_Base):
    def _seed(self):
        self.seed(_rec("1", "代表花子", group="1", sign="全員", tokuyaku="",
                       status="クラウドサイン登録", attachment=[{"fileKey": "k"}],
                       creditors=("アコム",)),
                  _rec("2", "二郎", group="1", email=EMAIL_B, status=""))

    def _post_with_change(self, change):
        """1 回目の plan の後に change() を適用（登録直前の再取得が変化を見る）。"""
        real_plan = hc.plan
        calls = {"n": 0}

        async def plan_then_change(record):
            calls["n"] += 1
            p = await real_plan(record)
            if calls["n"] == 1:
                change()
            return p
        with patch.object(hc, "plan", plan_then_change):
            return self.post(status="クラウドサイン登録").json()

    def test_changes_detected_zero_cloudsign_calls(self):
        cases = {
            "rep_email": lambda: self.records["1"].update(メールアドレス={"value": "z@example.com"}),
            "rep_addr": lambda: self.records["1"].update(住所={"value": "別住所"}),
            "rep_sign": lambda: self.records["1"].update(契約署名={"value": "代表者のみ"}),
            "rep_tokuyaku": lambda: self.records["1"].update(特約={"value": "分割可"}),
            "member_email": lambda: self.records["2"].update(メールアドレス={"value": "y@example.com"}),
            "creditor_added": lambda: self.records["2"].update(
                債権者一覧={"value": _creditors("プロミス")}),
        }
        for label, change in cases.items():
            with self.subTest(label=label):
                self.setUp()
                self._seed()
                self.assertEqual(self._post_with_change(change).get("skip"), "applicants_changed")
                self.cs_create.assert_not_called()
                self.cs_participant.assert_not_called()
                self.assertEqual(self.field("1", "契約書ステータス"), "要確認")
                self.assertEqual(self.kinds(), ["houki_contract_needs_review"])

    def test_unchanged_proceeds_using_refetched_record(self):
        self._seed()
        real_get = cw.hub_kintone.get_record
        seen = []

        async def spy(app, rid):
            seen.append(str(rid))
            return await real_get(app, rid)
        with patch.object(cw.hub_kintone, "get_record", spy):
            r = self.post(status="クラウドサイン登録")
        self.assertTrue(r.json().get("cloudsign"), r.text)
        self.assertGreaterEqual(seen.count("1"), 2)                     # webhook + 登録直前の再取得
        self.cs_create.assert_called_once()

    def test_refetch_failure_is_review(self):
        self._seed()
        real_get = cw.hub_kintone.get_record
        calls = {"n": 0}

        async def flaky(app, rid):
            calls["n"] += 1
            if str(rid) == "1" and calls["n"] >= 2:
                raise KintoneError(500, "x", "y")
            return await real_get(app, rid)
        with patch.object(cw.hub_kintone, "get_record", flaky):
            self.assertEqual(self.post(status="クラウドサイン登録").json().get("skip"),
                             "refetch_failed")
        self.cs_create.assert_not_called()
        self.assertEqual(self.field("1", "契約書ステータス"), "要確認")

    def test_fingerprint_is_json_serialized(self):
        def plan_for(creditors_a, creditors_b, name="甲", addr="架空市"):
            a = _rec("1", name, addr=addr, group="1", creditors=creditors_a)
            b = _rec("2", "乙", group="1", email=EMAIL_B, creditors=creditors_b, status="")
            fees = hc.compute_fees(2, len(hc.destinations([a, b])))
            return hc.HoukiPlan([a, b], "1", hc.SIGN_ALL, False, fees, [], [],
                                [(EMAIL_A, name), (EMAIL_B, "乙")])
        p1 = plan_for(("A,B", "C"), ("D",))
        p2 = plan_for(("A", "B", "C"), ("D",))
        self.assertNotEqual(p1.fingerprint, p2.fingerprint)          # 配列境界を保持
        p3 = plan_for(("A",), ("B",), name="甲|乙", addr="a,b")
        p4 = plan_for(("A",), ("B",), name="甲", addr="乙|a,b")
        self.assertNotEqual(p3.fingerprint, p4.fingerprint)          # 項目境界を保持
        src = p1.fingerprint_source()
        self.assertEqual(src["members"][0]["creditors"], ["A,B", "C"])
        self.assertEqual(set(src["fees"]), {"報酬合計", "実費合計", "追加送付件数",
                                             "追加送付料合計", "支払総額"})
        self.assertEqual(src["signers"], [EMAIL_A, EMAIL_B])
        self.assertEqual(p1.fingerprint, plan_for(("A,B", "C"), ("D",)).fingerprint)  # 決定的


# ── 5d. fix2 HCGF1-01: 再取得後の所有権確認 ─────────────────────────────────────
class TestOwnershipAfterClaim(_Base):
    def _seed(self):
        self.seed(_rec("1", "代表花子", group="1", sign="全員",
                       status="クラウドサイン登録", attachment=[{"fileKey": "k"}]),
                  _rec("2", "二郎", group="1", email=EMAIL_B, status=""))

    def _post_with_edit_between_claim_and_refetch(self, edit):
        """claim 後・再取得前（record 1 の 2 回目の get_record）に edit() を挟む。"""
        real_get = cw.hub_kintone.get_record
        calls = {"n": 0}

        async def get(app, rid):
            if str(rid) == "1":
                calls["n"] += 1
                if calls["n"] == 2:
                    edit()
            return await real_get(app, rid)
        with patch.object(cw.hub_kintone, "get_record", get):
            return self.post(status="クラウドサイン登録").json()

    def _human(self, rid, **fields):
        rec = self.records[rid]
        for k, v in fields.items():
            rec[k] = {"value": v}
        rec["$revision"] = {"value": str(int(rec["$revision"]["value"]) + 1)}

    def test_a_status_changed_by_human_no_write_no_api(self):
        self._seed()
        res = self._post_with_edit_between_claim_and_refetch(
            lambda: self._human("1", 契約書ステータス="要確認"))
        self.assertEqual(res.get("skip"), "state_changed")
        self.cs_create.assert_not_called()
        self.cs_attach.assert_not_called()
        self.assertEqual(len(self.updates), 1)                      # claim のみ
        self.assertEqual(self.field("1", "契約書ステータス"), "要確認")
        self.assertEqual(self.kinds(), ["houki_contract_needs_review"])   # 情報通知 1 通
        self.assertIn("レコードは変更していません", self.texts()[0])

    def test_b_non_fingerprint_edit_bumps_revision(self):
        self._seed()
        res = self._post_with_edit_between_claim_and_refetch(
            lambda: self._human("1", 電話番号="090-0000-0000"))
        self.assertEqual(res.get("skip"), "edited_during_claim")
        self.cs_create.assert_not_called()
        self.assertEqual(self.field("1", "契約書ステータス"), "要確認")
        self.assertEqual(self.field("1", "cloudsign_document_id"), "")    # 印なし
        self.assertEqual(self.updates[-1][1], {"契約書ステータス": "要確認"})

    def test_c_email_edit_and_member_email_edit(self):
        # 管理レコードのメール変更（revision が進む）→ edited_during_claim
        self._seed()
        res = self._post_with_edit_between_claim_and_refetch(
            lambda: self._human("1", メールアドレス="z@example.com"))
        self.assertEqual(res.get("skip"), "edited_during_claim")
        self.cs_create.assert_not_called()
        self.assertEqual(self.field("1", "契約書ステータス"), "要確認")   # 登録中が残らない
        # 他申述人のメール変更（管理レコードの revision は不変）→ applicants_changed
        self.setUp()
        self._seed()
        res = self._post_with_edit_between_claim_and_refetch(
            lambda: self._human("2", メールアドレス="y@example.com"))
        self.assertEqual(res.get("skip"), "applicants_changed")
        self.cs_create.assert_not_called()
        self.assertEqual(self.field("1", "契約書ステータス"), "要確認")
        self.assertEqual(self.field("1", "cloudsign_document_id"), "")

    def test_d_unchanged_registers_with_claim_revision(self):
        self._seed()
        res = self.post(status="クラウドサイン登録").json()
        self.assertTrue(res.get("cloudsign"), res)
        self.cs_create.assert_called_once()
        self.assertEqual(self.field("1", "契約書ステータス"), "クラウドサイン登録済")
        self.assertEqual(self.field("1", "cloudsign_document_id"), "doc-h1")
        self.assertEqual([u[1].get("契約書ステータス") for u in self.updates],
                         ["クラウドサイン登録中", "クラウドサイン登録済"])

    def test_edited_during_claim_cas_lost_leaves_reconcile(self):
        self._seed()
        real_update = cw.hub_kintone.update_record
        calls = {"n": 0}

        async def update(app, rid, fields, revision=None):
            calls["n"] += 1
            if calls["n"] == 2:                                  # 要確認への CAS が 409
                raise KintoneError(409, "GAIA_CO02", "c")
            return await real_update(app, rid, fields, revision)
        with patch.object(cw.hub_kintone, "update_record", update):
            res = self._post_with_edit_between_claim_and_refetch(
                lambda: self._human("1", 電話番号="1"))
        self.assertEqual(res.get("skip"), "cas_lost")
        self.assertEqual(self.field("1", "契約書ステータス"), "クラウドサイン登録中")
        self.cs_create.assert_not_called()


# ── 5e. fix2 HCGF1-02: 結果不明の永続印 ─────────────────────────────────────────
class TestResultUnknownMark(_Base):
    MARK = hc.CLOUDSIGN_RESULT_UNKNOWN_MARK

    def _seed_group(self, a_status="クラウドサイン登録"):
        self.seed(_rec("1", "代表花子", group="1", sign="全員", status=a_status,
                       attachment=[{"fileKey": "k"}]),
                  _rec("2", "二郎", group="1", email=EMAIL_B, status="",
                       attachment=[{"fileKey": "k2"}]))

    def test_mark_constant_and_text(self):
        self.assertEqual(self.MARK, "結果不明:要手動確認")
        self.assertIn("本物の document ID", hc.MANUAL_RESOLUTION_TEXT)
        self.assertIn("クラウドサイン登録済", hc.MANUAL_RESOLUTION_TEXT)

    def test_unknown_result_marks_and_blocks(self):
        for label, exc in {"unknown": cw.CloudSignResultUnknown("lost"),
                           "attach_fail": RuntimeError("attach")}.items():
            with self.subTest(label=label):
                self.setUp()
                self._seed_group()
                if label == "unknown":
                    self.cs_create.side_effect = exc
                else:
                    self.cs_attach.side_effect = exc
                res = self.post(status="クラウドサイン登録").json()
                self.assertIn(res.get("skip"), ("cs_result_unknown", "cs_partial_failure"))
                self.assertEqual(self.field("1", "契約書ステータス"), "要確認")
                self.assertEqual(self.field("1", "cloudsign_document_id"), self.MARK)
                self.assertEqual(self.kinds(), ["houki_contract_needs_review"])
                self.assertIn("結果不明:要手動確認", self.texts()[0])
                self.assertIn("(i) 下書きがあれば", self.texts()[0])

    def test_docid_save_failure_marks(self):
        self._seed_group()
        real_update = cw.hub_kintone.update_record

        async def update(app, rid, fields, revision=None):
            if fields.get("cloudsign_document_id") == "doc-h1":
                raise KintoneError(500, "x", "y")
            return await real_update(app, rid, fields, revision)
        with patch.object(cw.hub_kintone, "update_record", update):
            res = self.post(status="クラウドサイン登録").json()
        self.assertEqual(res.get("skip"), "cs_docid_save_failed")
        self.assertEqual(self.field("1", "cloudsign_document_id"), self.MARK)
        self.assertEqual(self.field("1", "契約書ステータス"), "要確認")

    def test_reconcile_stale_working_marks(self):
        self._seed_group(a_status="クラウドサイン登録中")
        res = self.post(status="クラウドサイン登録").json()
        self.assertEqual(res.get("skip"), "cs_needs_review")
        self.assertEqual(self.field("1", "契約書ステータス"), "要確認")
        self.assertEqual(self.field("1", "cloudsign_document_id"), self.MARK)
        self.assertIn("(ii) 無ければ", self.texts()[0])
        # 既に非空（本物の ID）なら上書きしない
        self.setUp()
        self._seed_group(a_status="クラウドサイン登録中")
        self.records["1"]["cloudsign_document_id"] = {"value": "doc-real"}
        self.post(status="クラウドサイン登録")
        self.assertEqual(self.field("1", "cloudsign_document_id"), "doc-real")

    def test_codex_repro_regroup_after_unknown_is_blocked(self):
        # A 作成 → 応答喪失 → 要確認+印
        self._seed_group()
        self.cs_create.side_effect = cw.CloudSignResultUnknown("lost")
        self.post(status="クラウドサイン登録")
        self.assertEqual(self.field("1", "cloudsign_document_id"), self.MARK)
        self.cs_create.side_effect = None
        self.cs_create.reset_mock()
        self.admin.reset_mock()
        # 全員のグループ ID を B（No.2）へ付け替え → B が管理レコード
        self.records["1"]["被相続人グループID"] = {"value": "2"}
        self.records["1"]["契約書ステータス"] = {"value": "要確認"}
        self.records["2"]["被相続人グループID"] = {"value": "2"}
        self.records["2"]["契約書ステータス"] = {"value": "契約書作成"}
        res = self.post(rid="2", status="契約書作成").json()
        self.assertEqual(res.get("skip"), hc.OTHER_MEMBER_HAS_CONTRACT)
        self.assertEqual(self.uploads, [])
        self.records["2"]["契約書ステータス"] = {"value": "クラウドサイン登録"}
        res = self.post(rid="2", status="クラウドサイン登録").json()
        self.assertEqual(res.get("skip"), hc.OTHER_MEMBER_HAS_CONTRACT)
        self.cs_create.assert_not_called()
        self.assertEqual(self.field("1", "cloudsign_document_id"), self.MARK)   # 上書きなし
        # A 自身の再登録も cs_registered で遮断（解除手順つき）
        self.records["1"]["被相続人グループID"] = {"value": "1"}
        self.records["2"]["被相続人グループID"] = {"value": "1"}
        self.records["2"]["契約書ステータス"] = {"value": ""}
        self.records["1"]["契約書ステータス"] = {"value": "クラウドサイン登録"}
        self.admin.reset_mock()
        res = self.post(rid="1", status="クラウドサイン登録").json()
        self.assertEqual(res.get("skip"), "cs_registered")
        self.cs_create.assert_not_called()
        self.assertIn("結果が不明のため", self.texts()[0])
        self.assertIn("(i) 下書きがあれば", self.texts()[0])
        self.assertEqual(self.field("1", "cloudsign_document_id"), self.MARK)
        # 印なしの 要確認（誤操作）は B→A の通常操作を妨げない
        self.setUp()
        self.seed(_rec("1", "代表花子", group="1", sign="全員"),
                  _rec("2", "二郎", group="1", email=EMAIL_B, status="契約書作成"))
        self.assertEqual(self.post(rid="2").json().get("skip"), hc.NOT_MANAGER)
        self.assertEqual(self.field("2", "契約書ステータス"), "要確認")
        self.assertEqual(self.field("2", "cloudsign_document_id"), "")
        self.assertEqual(self.post(rid="1").json().get("record_id"), "1")
        self.assertEqual(len(self.uploads), 1)

    def test_no_mark_paths(self):
        # 外部呼出し前に止まる要確認は印を書かない
        self.seed(_rec("1", "代表花子", group="1", sign="全員", status="クラウドサイン登録",
                       attachment=[{"fileKey": "k"}], email=""))
        self.assertEqual(self.post(status="クラウドサイン登録").json().get("skip"),
                         "houki_preconditions")
        self.assertEqual(self.field("1", "cloudsign_document_id"), "")
        self.setUp()
        self.seed(_rec("2", "乙", group="1"), _rec("1", "甲", group="1"))
        self.assertEqual(self.post(rid="2").json().get("skip"), hc.NOT_MANAGER)
        self.assertEqual(self.field("2", "cloudsign_document_id"), "")

    def test_cloudsign_webhook_ignores_mark(self):
        import cloudsign_webhook as mod
        with patch.object(mod.requests, "get") as g, patch.object(mod.requests, "put") as p:
            self.assertIsNone(mod.update_kintone_status(self.MARK, "受任"))
            self.assertIsNone(mod.update_kintone_status("", "受任"))
        g.assert_not_called()
        p.assert_not_called()


# ── 6. cloudsign_webhook の App 40 対応 ───────────────────────────────────────
class TestCloudSignWebhookApp40(unittest.TestCase):
    def setUp(self):
        import cloudsign_webhook
        self.mod = cloudsign_webhook

    def _resp(self, records):
        m = MagicMock()
        m.raise_for_status = MagicMock()
        m.json.return_value = {"records": records}
        return m

    def test_found_in_app40_only(self):
        put = MagicMock()
        put.raise_for_status = MagicMock()
        with patch.dict(os.environ, {"APP_HOUKI": "40", "TOKEN_HOUKI": "t40"}), \
                patch.object(self.mod.requests, "get",
                             side_effect=[self._resp([]), self._resp([{"$id": {"value": "8"}}])]) as g, \
                patch.object(self.mod.requests, "put", return_value=put) as p:
            self.assertEqual(self.mod.update_kintone_status("doc-h", "受任"), "8")
        self.assertEqual([c.kwargs["params"]["app"] for c in g.call_args_list],
                         [self.mod.KINTONE_APP_ID, "40"])
        self.assertEqual(p.call_args.kwargs["json"]["app"], "40")
        self.assertEqual(p.call_args.kwargs["headers"]["X-Cybozu-API-Token"], "t40")
        self.assertEqual(p.call_args.kwargs["json"]["record"], {"status": {"value": "受任"}})

    def test_both_apps_match_is_fail_closed(self):
        with patch.dict(os.environ, {"APP_HOUKI": "40", "TOKEN_HOUKI": "t40"}), \
                patch.object(self.mod.requests, "get",
                             side_effect=[self._resp([{"$id": {"value": "1"}}]),
                                          self._resp([{"$id": {"value": "8"}}])]), \
                patch.object(self.mod.requests, "put") as p:
            with self.assertRaises(self.mod.AmbiguousDocumentMatch):
                self.mod.update_kintone_status("doc-x", "受任")
        p.assert_not_called()
        # handle_webhook: 処理せず封筒+通知・200
        with patch.object(self.mod, "verify_completed_document",
                          return_value=({"id": "doc-x", "status": 2}, "")), \
                patch.object(self.mod, "update_kintone_status",
                             side_effect=self.mod.AmbiguousDocumentMatch("x")), \
                patch.object(self.mod, "file_mismatch_envelope", return_value="77") as env, \
                patch.object(self.mod, "notify_business_line") as nb:
            code, body = self.mod.handle_webhook(
                self.mod.WEBHOOK_SECRET, {"documentID": "doc-x", "status": 2})
        self.assertEqual((code, body["state"]), (200, "ambiguous_match"))
        env.assert_called_once_with("doc-x", "kintone_ambiguous_match")
        self.assertIn("kintone_ambiguous_match", nb.call_args.args[0])
        self.assertNotIn("doc-x", nb.call_args.args[0])

    def test_neither_app_is_none(self):
        with patch.dict(os.environ, {"APP_HOUKI": "40", "TOKEN_HOUKI": "t40"}), \
                patch.object(self.mod.requests, "get", return_value=self._resp([])), \
                patch.object(self.mod.requests, "put") as p:
            self.assertIsNone(self.mod.update_kintone_status("doc-none", "受任"))
        p.assert_not_called()


# ── 7. 通知 kind の登録・時効側の cfg 不変 ───────────────────────────────────
class TestConfigAndKinds(unittest.TestCase):
    def test_notify_kinds_registered(self):
        for kind in ("houki_contract_created", "houki_contract_needs_review",
                     "houki_contract_failed"):
            with self.subTest(kind=kind), \
                    self.assertLogs(hub_notify.logger, level="INFO") as cm:
                hub_notify._log_throttled(f"{kind}:12")
                self.assertIn(f"kind={kind}", "\n".join(cm.output))
                self.assertNotIn("unknown_kind", "\n".join(cm.output))

    def test_channel_configs(self):
        j, h = cw._jikou_cfg(), cw._houki_cfg()
        self.assertEqual((j.name, j.template_path, j.frozen_clause, j.output_pdf_name),
                         ("jikou", cw.TEMPLATE_PATH, cw.FROZEN_CLAUSE, cw.OUTPUT_PDF_NAME))
        self.assertEqual((h.name, h.template_path, h.template_sha256, h.output_filename,
                          h.output_pdf_name),
                         ("houki", hc.TEMPLATE_PATH, hc.TEMPLATE_SHA256,
                          "委任契約書_相続放棄.docx", "委任契約書_相続放棄.pdf"))
        self.assertEqual(h.app.app_id(), "40")
        self.assertEqual(h.cs_title("7"), "相続放棄_委任契約書_案件No.7")
        self.assertEqual(j.cs_title("7"), "委任契約書_案件No.7")
        # 時効側の定数を patch すると jikou cfg は追随する（既存テスト互換）
        with patch.object(cw, "TEMPLATE_SHA256", "0" * 64):
            self.assertEqual(cw._jikou_cfg().template_sha256, "0" * 64)


if __name__ == "__main__":
    unittest.main()
